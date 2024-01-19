"""Step implementations for footnote-related features."""

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.footnotes import Footnote
from docx.text.paragraph import Paragraph

from helpers import test_docx

# given ====================================================


@given("a document with 3 footnotes and 2 default footnotes")
def given_a_document_with_3_footnotes_and_2_default_footnotes(context: Context):
    document = Document(test_docx("footnotes"))
    context.footnotes = document.footnotes


@given("a document with footnotes and with all footnotes properties")
def given_a_document_with_footnotes_and_with_all_footnotes_properties(context: Context):
    document = Document(test_docx("footnotes"))
    context.section = document.sections[0]


@given("a document with footnotes")
def given_a_document_with_footnotes(context: Context):
    document = Document(test_docx("footnotes"))
    context.footnotes = document.footnotes


@given("a document without footnotes")
def given_a_document_without_footnotes(context: Context):
    document = Document(test_docx("doc-default"))
    context.footnotes = document.footnotes
    context.section = document.sections[0]


@given("a paragraph in a document without footnotes")
def given_a_paragraph_in_a_document_without_footnotes(context: Context):
    document = Document(test_docx("par-known-paragraphs"))
    context.paragraphs = document.paragraphs
    context.footnotes = document.footnotes


@given(
    "a document with paragraphs[0] containing one, paragraphs[1] containing none, and paragraphs[2] containing two footnotes"
)
def given_a_document_with_3_footnotes(context: Context):
    document = Document(test_docx("footnotes"))
    context.paragraphs = document.paragraphs
    context.footnotes = document.footnotes


# when ====================================================


@when("I try to access a footnote with invalid reference id")
def when_I_try_to_access_a_footnote_with_invalid_reference_id(context: Context):
    context.exc = None
    try:
        context.footnotes[10]
    except IndexError as e:
        context.exc = e


@when("I add a footnote to the paragraphs[{parId}] with text '{footnoteText}'")
def when_I_add_a_footnote_to_the_paragraph_with_text_text(
    context: Context, parId: str, footnoteText: str
):
    par = context.paragraphs[int(parId)]
    new_footnote = par.add_footnote()
    new_footnote.add_paragraph(footnoteText)


@when("I change footnote property {propName} to {value}")
def when_I_change_footnote_property_propName_to_value(
    context: Context, propName: str, value: str
):
    context.section.__setattr__(propName, eval(value))


# then =====================================================


@then("len(footnotes) is {expectedLen}")
def then_len_footnotes_is_len(context: Context, expectedLen: str):
    footnotes = context.footnotes
    assert len(footnotes) == int(
        expectedLen
    ), f"expected len(footnotes) of {expectedLen}, got {len(footnotes)}"


@then("I can access a footnote by footnote reference id")
def then_I_can_access_a_footnote_by_footnote_reference_id(context: Context):
    footnotes = context.footnotes
    for refId in range(-1, 3):
        footnote = footnotes[refId]
        assert isinstance(footnote, Footnote)


@then("I can access a paragraph in a specific footnote")
def then_I_can_access_a_paragraph_in_a_specific_footnote(context: Context):
    footnotes = context.footnotes
    for refId in range(1, 3):
        footnote = footnotes[refId]
        assert isinstance(footnote.paragraphs[0], Paragraph)


@then("it trows an {exceptionType}")
def then_it_trows_an_IndexError(context: Context, exceptionType: str):
    exc = context.exc
    assert isinstance(exc, eval(exceptionType)), f"expected IndexError, got {type(exc)}"


@then("I can access footnote property {propName} with value {value}")
def then_I_can_access_footnote_propery_name_with_value_value(
    context: Context, propName: str, value: str
):
    actual_value = context.section.__getattribute__(propName)
    expected = eval(value)
    assert (
        actual_value == expected
    ), f"expected section.{propName} {value}, got {expected}"


@then(
    "the document contains a footnote with footnote reference id of {refId} with text '{footnoteText}'"
)
def then_the_document_contains_a_footnote_with_footnote_reference_id_of_refId_with_text_text(
    context: Context, refId: str, footnoteText: str
):
    par = context.paragraphs[1]
    f = par.footnotes[0]
    assert f.id == int(refId), f"expected {refId}, got {f.id}"
    assert (
        f.paragraphs[0].text == footnoteText
    ), f"expected {footnoteText}, got {f.paragraphs[0].text}"


@then(
    "paragraphs[{parId}] has footnote reference ids of {refIds}, with footnote text {fText}"
)
def then_paragraph_has_footnote_reference_ids_of_refIds_with_footnote_text_text(
    context: Context, parId: str, refIds: str, fText: str
):
    par = context.paragraphs[int(parId)]
    refIds = eval(refIds)
    fText = eval(fText)
    if refIds is not None:
        if type(refIds) is list:
            for i in range(len(refIds)):
                f = par.footnotes[i]
                assert isinstance(
                    f, Footnote
                ), f"expected to be instance of Footnote, got {type(f)}"
                assert f.id == refIds[i], f"expected {refIds[i]}, got {f.id}"
                assert (
                    f.paragraphs[0].text == fText[i]
                ), f"expected '{fText[i]}', got '{f.paragraphs[0].text}'"
        else:
            f = par.footnotes[0]
            assert f.id == int(refIds), f"expected {refIds}, got {f.id}"
            assert (
                f.paragraphs[0].text == fText
            ), f"expected '{fText}', got '{f.paragraphs[0].text}'"
    else:
        assert (
            len(par.footnotes) == 0
        ), f"expected an empty list, got {len(par.footnotes)} elements"
