# encoding: utf-8

"""
Step implementations for graphical object (shape) related features
"""

from __future__ import absolute_import, print_function, unicode_literals

import hashlib

from behave import given, then, when

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shape import InlineShape
from docx.shared import Inches

from helpers import test_docx


# given ===================================================

@given('an inline shape collection containing five shapes')
def given_an_inline_shape_collection_containing_five_shapes(context):
    docx_path = test_docx('shp-inline-shape-access')
    document = Document(docx_path)
    context.inline_shapes = document.inline_shapes


@given('an inline shape of known dimensions')
def given_inline_shape_of_known_dimensions(context):
    document = Document(test_docx('shp-inline-shape-access'))
    context.inline_shape = document.inline_shapes[0]


@given('an inline shape known to be {shp_of_type}')
def given_inline_shape_known_to_be_shape_of_type(context, shp_of_type):
    inline_shape_idx = {
        'an embedded picture':  0,
        'a linked picture':     1,
        'a link+embed picture': 2,
        'a smart art diagram':  3,
        'a chart':              4,
    }[shp_of_type]
    docx_path = test_docx('shp-inline-shape-access')
    document = Document(docx_path)
    context.inline_shape = document.inline_shapes[inline_shape_idx]


# when =====================================================

@when('I change the dimensions of the inline shape')
def when_change_dimensions_of_inline_shape(context):
    inline_shape = context.inline_shape
    inline_shape.width = Inches(1)
    inline_shape.height = Inches(0.5)


# then =====================================================

@then('I can access each inline shape by index')
def then_can_access_each_inline_shape_by_index(context):
    inline_shapes = context.inline_shapes
    for idx in range(2):
        inline_shape = inline_shapes[idx]
        assert isinstance(inline_shape, InlineShape)


@then('I can iterate over the inline shape collection')
def then_can_iterate_over_inline_shape_collection(context):
    inline_shapes = context.inline_shapes
    shape_count = 0
    for inline_shape in inline_shapes:
        shape_count += 1
        assert isinstance(inline_shape, InlineShape)
    expected_count = 5
    assert shape_count == expected_count, (
        'expected %d, got %d' % (expected_count, shape_count)
    )


@then('its inline shape type is {shape_type}')
def then_inline_shape_type_is_shape_type(context, shape_type):
    expected_value = {
        'WD_INLINE_SHAPE.CHART':          WD_INLINE_SHAPE.CHART,
        'WD_INLINE_SHAPE.LINKED_PICTURE': WD_INLINE_SHAPE.LINKED_PICTURE,
        'WD_INLINE_SHAPE.PICTURE':        WD_INLINE_SHAPE.PICTURE,
        'WD_INLINE_SHAPE.SMART_ART':      WD_INLINE_SHAPE.SMART_ART,
    }[shape_type]
    inline_shape = context.inline_shape
    assert inline_shape.type == expected_value


@then('the dimensions of the inline shape match the known values')
def then_dimensions_of_inline_shape_match_known_values(context):
    inline_shape = context.inline_shape
    assert inline_shape.width == 1778000, 'got %s' % inline_shape.width
    assert inline_shape.height == 711200, 'got %s' % inline_shape.height


@then('the dimensions of the inline shape match the new values')
def then_dimensions_of_inline_shape_match_new_values(context):
    inline_shape = context.inline_shape
    assert inline_shape.width == 914400, 'got %s' % inline_shape.width
    assert inline_shape.height == 457200, 'got %s' % inline_shape.height


@then('the document contains the inline picture')
def then_the_document_contains_the_inline_picture(context):
    document = context.document
    picture_shape = document.inline_shapes[0]
    blip = picture_shape._inline.graphic.graphicData.pic.blipFill.blip
    rId = blip.embed
    image_part = document.part.related_parts[rId]
    image_sha1 = hashlib.sha1(image_part.blob).hexdigest()
    expected_sha1 = '79769f1e202add2e963158b532e36c2c0f76a70c'
    assert image_sha1 == expected_sha1, (
        "image SHA1 doesn't match, expected %s, got %s" %
        (expected_sha1, image_sha1)
    )


@then('the length of the inline shape collection is 5')
def then_len_of_inline_shape_collection_is_5(context):
    inline_shapes = context.inline_shapes
    shape_count = len(inline_shapes)
    assert shape_count == 5, 'got %s' % shape_count


@then('the picture has its native width and height')
def then_picture_has_native_width_and_height(context):
    picture = context.picture
    assert picture.width == 1905000, 'got %d' % picture.width
    assert picture.height == 2717800, 'got %d' % picture.height


@then('picture.height is {inches} inches')
def then_picture_height_is_value(context, inches):
    expected_value = {
        '2.14': 1956816,
        '2.5':  2286000,
    }[inches]
    picture = context.picture
    assert picture.height == expected_value, 'got %d' % picture.height


@then('picture.width is {inches} inches')
def then_picture_width_is_value(context, inches):
    expected_value = {
        '1.05':  961402,
        '1.75': 1600200,
    }[inches]
    picture = context.picture
    assert picture.width == expected_value, 'got %d' % picture.width
