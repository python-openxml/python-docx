# encoding: utf-8

"""
Step implementations for font-related features.
"""

from __future__ import absolute_import, division, print_function, unicode_literals

from behave import given, then, when

from docx import Document
from docx.dml.color import ColorFormat
from docx.enum.dml import MSO_COLOR_TYPE, MSO_THEME_COLOR
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
from docx.shared import RGBColor

from helpers import test_docx


# given ===================================================


@given("a font")
def given_a_font(context):
    document = Document(test_docx("txt-font-props"))
    context.font = document.paragraphs[0].runs[0].font


@given("a font having {color} highlighting")
def given_a_font_having_color_highlighting(context, color):
    paragraph_index = {
        "no": 0,
        "yellow": 1,
        "bright green": 2,
    }[color]
    document = Document(test_docx("txt-font-highlight-color"))
    context.font = document.paragraphs[paragraph_index].runs[0].font


@given("a font having {type} color")
def given_a_font_having_type_color(context, type):
    run_idx = ["no", "auto", "an RGB", "a theme"].index(type)
    document = Document(test_docx("fnt-color"))
    context.font = document.paragraphs[0].runs[run_idx].font


@given("a font having typeface name {name}")
def given_a_font_having_typeface_name(context, name):
    document = Document(test_docx("txt-font-props"))
    style_name = {
        "not specified": "Normal",
        "Avenir Black": "Having Typeface",
    }[name]
    context.font = document.styles[style_name].font


@given("a font having {underline_type} underline")
def given_a_font_having_type_underline(context, underline_type):
    style_name = {
        "inherited": "Normal",
        "no": "None Underlined",
        "single": "Underlined",
        "double": "Double Underlined",
    }[underline_type]
    document = Document(test_docx("txt-font-props"))
    context.font = document.styles[style_name].font


@given("a font having {vertAlign_state} vertical alignment")
def given_a_font_having_vertAlign_state(context, vertAlign_state):
    style_name = {
        "inherited": "Normal",
        "subscript": "Subscript",
        "superscript": "Superscript",
    }[vertAlign_state]
    document = Document(test_docx("txt-font-props"))
    context.font = document.styles[style_name].font


@given("a font of size {size}")
def given_a_font_of_size(context, size):
    document = Document(test_docx("txt-font-props"))
    style_name = {
        "unspecified": "Normal",
        "14 pt": "Having Typeface",
        "18 pt": "Large Size",
    }[size]
    context.font = document.styles[style_name].font


# when ====================================================


@when("I assign {value} to font.color.rgb")
def when_I_assign_value_to_font_color_rgb(context, value):
    font = context.font
    new_value = None if value == "None" else RGBColor.from_string(value)
    font.color.rgb = new_value


@when("I assign {value} to font.color.theme_color")
def when_I_assign_value_to_font_color_theme_color(context, value):
    font = context.font
    new_value = None if value == "None" else getattr(MSO_THEME_COLOR, value)
    font.color.theme_color = new_value


@when("I assign {value} to font.highlight_color")
def when_I_assign_value_to_font_highlight_color(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(WD_COLOR_INDEX, value)
    font.highlight_color = expected_value


@when("I assign {value} to font.name")
def when_I_assign_value_to_font_name(context, value):
    font = context.font
    value = None if value == "None" else value
    font.name = value


@when("I assign {value} to font.size")
def when_I_assign_value_str_to_font_size(context, value):
    value = None if value == "None" else int(value)
    font = context.font
    font.size = value


@when("I assign {value} to font.underline")
def when_I_assign_value_to_font_underline(context, value):
    new_value = {
        "True": True,
        "False": False,
        "None": None,
        "WD_UNDERLINE.SINGLE": WD_UNDERLINE.SINGLE,
        "WD_UNDERLINE.DOUBLE": WD_UNDERLINE.DOUBLE,
    }[value]
    font = context.font
    font.underline = new_value


@when("I assign {value} to font.{sub_super}script")
def when_I_assign_value_to_font_sub_super(context, value, sub_super):
    font = context.font
    name = {
        "sub": "subscript",
        "super": "superscript",
    }[sub_super]
    new_value = {
        "None": None,
        "True": True,
        "False": False,
    }[value]

    setattr(font, name, new_value)


# then =====================================================


@then("font.color is a ColorFormat object")
def then_font_color_is_a_ColorFormat_object(context):
    font = context.font
    assert isinstance(font.color, ColorFormat)


@then("font.color.rgb is {value}")
def then_font_color_rgb_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else RGBColor.from_string(value)
    assert font.color.rgb == expected_value


@then("font.color.theme_color is {value}")
def then_font_color_theme_color_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(MSO_THEME_COLOR, value)
    assert font.color.theme_color == expected_value


@then("font.color.type is {value}")
def then_font_color_type_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(MSO_COLOR_TYPE, value)
    assert font.color.type == expected_value


@then("font.highlight_color is {value}")
def then_font_highlight_color_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(WD_COLOR_INDEX, value)
    assert font.highlight_color == expected_value


@then("font.name is {value}")
def then_font_name_is_value(context, value):
    font = context.font
    value = None if value == "None" else value
    assert font.name == value


@then("font.size is {value}")
def then_font_size_is_value(context, value):
    value = None if value == "None" else int(value)
    font = context.font
    assert font.size == value


@then("font.underline is {value}")
def then_font_underline_is_value(context, value):
    expected_value = {
        "None": None,
        "True": True,
        "False": False,
        "WD_UNDERLINE.DOUBLE": WD_UNDERLINE.DOUBLE,
    }[value]
    font = context.font
    assert font.underline == expected_value


@then("font.{sub_super}script is {value}")
def then_font_sub_super_is_value(context, sub_super, value):
    name = {
        "sub": "subscript",
        "super": "superscript",
    }[sub_super]
    expected_value = {
        "None": None,
        "True": True,
        "False": False,
    }[value]
    font = context.font
    actual_value = getattr(font, name)
    assert actual_value == expected_value
