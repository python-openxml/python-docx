"""Step implementations for text-related features."""

import hashlib

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.text import WD_BREAK, WD_UNDERLINE
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.text.font import Font
from docx.text.run import Run

from helpers import test_docx, test_file, test_text

# given ===================================================


@given("a run")
def given_a_run(context):
    document = Document()
    run = document.add_paragraph().add_run()
    context.document = document
    context.run = run


@given("a run having {bool_prop_name} set on")
def given_a_run_having_bool_prop_set_on(context, bool_prop_name):
    run = Document().add_paragraph().add_run()
    setattr(run, bool_prop_name, True)
    context.run = run


@given("a run having known text and formatting")
def given_a_run_having_known_text_and_formatting(context):
    run = Document().add_paragraph().add_run("foobar")
    run.bold = True
    run.italic = True
    context.run = run


@given("a run having mixed text content")
def given_a_run_having_mixed_text_content(context):
    """
    Mixed here meaning it contains ``<w:tab/>``, ``<w:cr/>``, etc. elements.
    """
    r_xml = """\
        <w:r %s>
          <w:t>abc</w:t>
          <w:br/>
          <w:t>def</w:t>
          <w:cr/>
          <w:t>ghi</w:t>
          <w:drawing/>
          <w:t>jkl</w:t>
          <w:tab/>
          <w:t>mno</w:t>
          <w:noBreakHyphen/>
          <w:t>pqr</w:t>
          <w:ptab/>
          <w:t>stu</w:t>
        </w:r>""" % nsdecls(
        "w"
    )
    r = parse_xml(r_xml)
    context.run = Run(r, None)


@given("a run having {underline_type} underline")
def given_a_run_having_underline_type(context, underline_type):
    run_idx = {"inherited": 0, "no": 1, "single": 2, "double": 3}[underline_type]
    document = Document(test_docx("run-enumerated-props"))
    context.run = document.paragraphs[0].runs[run_idx]


@given("a run having {style} style")
def given_a_run_having_style(context, style):
    run_idx = {
        "no explicit": 0,
        "Emphasis": 1,
        "Strong": 2,
    }[style]
    context.document = document = Document(test_docx("run-char-style"))
    context.run = document.paragraphs[0].runs[run_idx]


@given("a run having {zero_or_more} rendered page breaks")
def given_a_run_having_rendered_page_breaks(context: Context, zero_or_more: str):
    paragraph_idx = {"no": 0, "one": 1, "two": 3}[zero_or_more]
    document = Document(test_docx("par-rendered-page-breaks"))
    paragraph = document.paragraphs[paragraph_idx]
    context.run = paragraph.runs[0]


@given("a run inside a table cell retrieved from {cell_source}")
def given_a_run_inside_a_table_cell_from_source(context, cell_source):
    document = Document()
    table = document.add_table(rows=2, cols=2)
    if cell_source == "Table.cell":
        cell = table.cell(0, 0)
    elif cell_source == "Table.row.cells":
        cell = table.rows[0].cells[1]
    elif cell_source == "Table.column.cells":
        cell = table.columns[1].cells[0]
    run = cell.paragraphs[0].add_run()
    context.document = document
    context.run = run


# when ====================================================


@when("I add a column break")
def when_add_column_break(context):
    run = context.run
    run.add_break(WD_BREAK.COLUMN)


@when("I add a line break")
def when_add_line_break(context):
    run = context.run
    run.add_break()


@when("I add a page break")
def when_add_page_break(context):
    run = context.run
    run.add_break(WD_BREAK.PAGE)


@when("I add a picture to the run")
def when_I_add_a_picture_to_the_run(context):
    run = context.run
    run.add_picture(test_file("monty-truth.png"))


@when("I add a run specifying its text")
def when_I_add_a_run_specifying_its_text(context):
    context.run = context.paragraph.add_run(test_text)


@when("I add a run specifying the character style Emphasis")
def when_I_add_a_run_specifying_the_character_style_Emphasis(context):
    context.run = context.paragraph.add_run(test_text, "Emphasis")


@when("I add a tab")
def when_I_add_a_tab(context):
    context.run.add_tab()


@when("I add text to the run")
def when_I_add_text_to_the_run(context):
    context.run.add_text(test_text)


@when("I assign mixed text to the text property")
def when_I_assign_mixed_text_to_the_text_property(context):
    context.run.text = "abc\ndef\rghijkl\tmno-pqr\tstu"


@when("I assign {value_str} to its {bool_prop_name} property")
def when_assign_true_to_bool_run_prop(context, value_str, bool_prop_name):
    value = {"True": True, "False": False, "None": None}[value_str]
    run = context.run
    setattr(run, bool_prop_name, value)


@when("I assign {value} to run.style")
def when_I_assign_value_to_run_style(context, value):
    if value == "None":
        new_value = None
    elif value.startswith("styles["):
        new_value = context.document.styles[value.split("'")[1]]
    else:
        new_value = context.document.styles[value]

    context.run.style = new_value


@when("I clear the run")
def when_I_clear_the_run(context):
    context.run.clear()


@when("I set the run underline to {underline_value}")
def when_I_set_the_run_underline_to_value(context, underline_value):
    new_value = {
        "True": True,
        "False": False,
        "None": None,
        "WD_UNDERLINE.SINGLE": WD_UNDERLINE.SINGLE,
        "WD_UNDERLINE.DOUBLE": WD_UNDERLINE.DOUBLE,
    }[underline_value]
    context.run.underline = new_value


# then =====================================================


@then("it is a column break")
def then_type_is_column_break(context):
    attrib = context.last_child.attrib
    assert attrib == {qn("w:type"): "column"}


@then("it is a line break")
def then_type_is_line_break(context):
    attrib = context.last_child.attrib
    assert attrib == {}


@then("it is a page break")
def then_type_is_page_break(context):
    attrib = context.last_child.attrib
    assert attrib == {qn("w:type"): "page"}


@then("run.contains_page_break is {value}")
def then_run_contains_page_break_is_value(context: Context, value: str):
    actual = context.run.contains_page_break
    expected = {"True": True, "False": False}[value]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("run.font is the Font object for the run")
def then_run_font_is_the_Font_object_for_the_run(context):
    run, font = context.run, context.run.font
    assert isinstance(font, Font)
    assert font.element is run.element


@then("run.iter_inner_content() generates the run text and rendered page-breaks")
def then_run_iter_inner_content_generates_text_and_page_breaks(context: Context):
    actual_value = [type(item).__name__ for item in context.run.iter_inner_content()]
    expected_value = ["str", "RenderedPageBreak", "str", "RenderedPageBreak", "str"]
    assert (
        actual_value == expected_value
    ), f"expected: {expected_value}, got: {actual_value}"


@then("run.style is styles['{style_name}']")
def then_run_style_is_style(context, style_name):
    expected_value = context.document.styles[style_name]
    run = context.run
    assert run.style == expected_value, "got %s" % run.style


@then("run.text contains the text content of the run")
def then_run_text_contains_the_text_content_of_the_run(context):
    actual = context.run.text
    expected = "abc\ndef\nghijkl\tmno-pqr\tstu"
    assert actual == expected, f"expected:\n'{expected}'\n\ngot:\n'{actual}'"


@then("the last item in the run is a break")
def then_last_item_in_run_is_a_break(context):
    run = context.run
    context.last_child = run._r[-1]
    expected_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
    assert context.last_child.tag == expected_tag


@then("the picture appears at the end of the run")
def then_the_picture_appears_at_the_end_of_the_run(context):
    run = context.run
    r = run._r
    blip_rId = r.xpath(
        "./w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic/pic:blipFill/"
        "a:blip/@r:embed"
    )[0]
    image_part = run.part.related_parts[blip_rId]
    image_sha1 = hashlib.sha1(image_part.blob).hexdigest()
    expected_sha1 = "79769f1e202add2e963158b532e36c2c0f76a70c"
    assert (
        image_sha1 == expected_sha1
    ), "image SHA1 doesn't match, expected %s, got %s" % (expected_sha1, image_sha1)


@then("the run appears in {boolean_prop_name} unconditionally")
def then_run_appears_in_boolean_prop_name(context, boolean_prop_name):
    run = context.run
    assert getattr(run, boolean_prop_name) is True


@then("the run appears with its inherited {boolean_prop_name} setting")
def then_run_inherits_bool_prop_value(context, boolean_prop_name):
    run = context.run
    assert getattr(run, boolean_prop_name) is None


@then("the run appears without {boolean_prop_name} unconditionally")
def then_run_appears_without_bool_prop(context, boolean_prop_name):
    run = context.run
    assert getattr(run, boolean_prop_name) is False


@then("the run contains no text")
def then_the_run_contains_no_text(context):
    assert context.run.text == ""


@then("the run contains the text I specified")
def then_the_run_contains_the_text_I_specified(context):
    assert context.run.text == test_text


@then("the run formatting is preserved")
def then_the_run_formatting_is_preserved(context):
    assert context.run.bold is True
    assert context.run.italic is True


@then("the run underline property value is {underline_value}")
def then_the_run_underline_property_value_is(context, underline_value):
    expected_value = {
        "None": None,
        "False": False,
        "True": True,
        "WD_UNDERLINE.DOUBLE": WD_UNDERLINE.DOUBLE,
    }[underline_value]
    assert context.run.underline == expected_value


@then("the tab appears at the end of the run")
def then_the_tab_appears_at_the_end_of_the_run(context):
    r = context.run._r
    tab = r.find(qn("w:tab"))
    assert tab is not None
