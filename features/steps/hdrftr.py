"""Step implementations for header and footer-related features."""

from behave import given, then, when

from docx import Document

from helpers import test_docx, test_file

# given ====================================================


@given("a _Footer object {with_or_no} footer definition as footer")
def given_a_Footer_object_with_or_no_footer_definition(context, with_or_no):
    section_idx = {"with a": 0, "with no": 1}[with_or_no]
    context.sections = Document(test_docx("hdr-header-footer")).sections
    context.footer = context.sections[section_idx].footer


@given("a _Header object {with_or_no} header definition as header")
def given_a_Header_object_with_or_no_header_definition(context, with_or_no):
    section_idx = {"with a": 0, "with no": 1}[with_or_no]
    context.sections = Document(test_docx("hdr-header-footer")).sections
    context.header = context.sections[section_idx].header


@given("a _Run object from a footer as run")
def given_a_Run_object_from_a_footer_as_run(context):
    footer = Document(test_docx("hdr-header-footer")).sections[0].footer
    context.run = footer.paragraphs[0].add_run()


@given("a _Run object from a header as run")
def given_a_Run_object_from_a_header_as_run(context):
    header = Document(test_docx("hdr-header-footer")).sections[0].header
    context.run = header.paragraphs[0].add_run()


@given("the next _Footer object with no footer definition as footer_2")
def given_the_next_Footer_object_with_no_footer_definition(context):
    context.footer_2 = context.sections[1].footer


@given("the next _Header object with no header definition as header_2")
def given_the_next_Header_object_with_no_header_definition(context):
    context.header_2 = context.sections[1].header


# when =====================================================


@when('I assign "Normal" to footer.paragraphs[0].style')
def when_I_assign_Body_Text_to_footer_style(context):
    context.footer.paragraphs[0].style = "Normal"


@when('I assign "Normal" to header.paragraphs[0].style')
def when_I_assign_Body_Text_to_header_style(context):
    context.header.paragraphs[0].style = "Normal"


@when("I assign {value} to header.is_linked_to_previous")
def when_I_assign_value_to_header_is_linked_to_previous(context, value):
    context.header.is_linked_to_previous = eval(value)


@when("I assign {value} to footer.is_linked_to_previous")
def when_I_assign_value_to_footer_is_linked_to_previous(context, value):
    context.footer.is_linked_to_previous = eval(value)


@when("I call run.add_picture()")
def when_I_call_run_add_picture(context):
    context.run.add_picture(test_file("test.png"))


# then =====================================================


@then("footer.is_linked_to_previous is {value}")
def then_footer_is_linked_to_previous_is_value(context, value):
    actual = context.footer.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "footer.is_linked_to_previous is %s" % actual


@then('footer.paragraphs[0].style.name == "Normal"')
def then_footer_paragraphs_0_style_name_eq_Normal(context):
    actual = context.footer.paragraphs[0].style.name
    expected = "Normal"
    assert actual == expected, "footer.paragraphs[0].style.name is %s" % actual


@then("footer_2.is_linked_to_previous is {value}")
def then_footer_2_is_linked_to_previous_is_value(context, value):
    actual = context.footer_2.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "footer_2.is_linked_to_previous is %s" % actual


@then("footer_2.paragraphs[0].text == footer.paragraphs[0].text")
def then_footer_2_text_eq_footer_text(context):
    actual = context.footer_2.paragraphs[0].text
    expected = context.footer.paragraphs[0].text
    assert actual == expected, "footer_2.paragraphs[0].text == %s" % actual


@then("header.is_linked_to_previous is {value}")
def then_header_is_linked_to_previous_is_value(context, value):
    actual = context.header.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "header.is_linked_to_previous is %s" % actual


@then('header.paragraphs[0].style.name == "Normal"')
def then_header_paragraphs_0_style_name_eq_Normal(context):
    actual = context.header.paragraphs[0].style.name
    expected = "Normal"
    assert actual == expected, "header.paragraphs[0].style.name is %s" % actual


@then("header_2.is_linked_to_previous is {value}")
def then_header_2_is_linked_to_previous_is_value(context, value):
    actual = context.header_2.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "header_2.is_linked_to_previous is %s" % actual


@then("header_2.paragraphs[0].text == header.paragraphs[0].text")
def then_header_2_text_eq_header_text(context):
    actual = context.header_2.paragraphs[0].text
    expected = context.header.paragraphs[0].text
    assert actual == expected, "header_2.paragraphs[0].text == %s" % actual


@then("I can't detect the image but no exception is raised")
def then_I_cant_detect_the_image_but_no_exception_is_raised(context):
    pass
