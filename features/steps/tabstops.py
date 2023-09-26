"""Step implementations for paragraph-related features."""

from behave import given, then, when

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches
from docx.text.tabstops import TabStop

from helpers import test_docx

# given ===================================================


@given("a tab_stops having {count} tab stops")
def given_a_tab_stops_having_count_tab_stops(context, count):
    paragraph_idx = {"0": 0, "3": 1}[count]
    document = Document(test_docx("tab-stops"))
    paragraph_format = document.paragraphs[paragraph_idx].paragraph_format
    context.tab_stops = paragraph_format.tab_stops


@given("a tab stop 0.5 inches {in_or_out} from the paragraph left edge")
def given_a_tab_stop_inches_from_paragraph_left_edge(context, in_or_out):
    tab_idx = {"out": 0, "in": 1}[in_or_out]
    document = Document(test_docx("tab-stops"))
    paragraph_format = document.paragraphs[2].paragraph_format
    context.tab_stops = paragraph_format.tab_stops
    context.tab_stop = paragraph_format.tab_stops[tab_idx]


@given("a tab stop having {alignment} alignment")
def given_a_tab_stop_having_alignment_alignment(context, alignment):
    tab_idx = {"LEFT": 0, "CENTER": 1, "RIGHT": 2}[alignment]
    document = Document(test_docx("tab-stops"))
    paragraph_format = document.paragraphs[1].paragraph_format
    context.tab_stop = paragraph_format.tab_stops[tab_idx]


@given("a tab stop having {leader} leader")
def given_a_tab_stop_having_leader_leader(context, leader):
    tab_idx = {"no specified": 0, "a dotted": 2}[leader]
    document = Document(test_docx("tab-stops"))
    paragraph_format = document.paragraphs[1].paragraph_format
    context.tab_stop = paragraph_format.tab_stops[tab_idx]


# when ====================================================


@when("I add a tab stop")
def when_I_add_a_tab_stop(context):
    tab_stops = context.tab_stops
    tab_stops.add_tab_stop(Inches(1.75))


@when("I assign {member} to tab_stop.alignment")
def when_I_assign_member_to_tab_stop_alignment(context, member):
    value = getattr(WD_TAB_ALIGNMENT, member)
    context.tab_stop.alignment = value


@when("I assign {member} to tab_stop.leader")
def when_I_assign_member_to_tab_stop_leader(context, member):
    value = getattr(WD_TAB_LEADER, member)
    context.tab_stop.leader = value


@when("I assign {value} to tab_stop.position")
def when_I_assign_value_to_tab_stop_value(context, value):
    context.tab_stop.position = int(value)


@when("I call tab_stops.clear_all()")
def when_I_call_tab_stops_clear_all(context):
    tab_stops = context.tab_stops
    tab_stops.clear_all()


@when("I remove a tab stop")
def when_I_remove_a_tab_stop(context):
    tab_stops = context.tab_stops
    del tab_stops[1]


# then =====================================================


@then("I can access a tab stop by index")
def then_I_can_access_a_tab_stop_by_index(context):
    tab_stops = context.tab_stops
    for idx in range(3):
        tab_stop = tab_stops[idx]
        assert isinstance(tab_stop, TabStop)


@then("I can iterate the TabStops object")
def then_I_can_iterate_the_TabStops_object(context):
    items = list(context.tab_stops)
    assert len(items) == 3
    assert all(isinstance(item, TabStop) for item in items)


@then("len(tab_stops) is {count}")
def then_len_tab_stops_is_count(context, count):
    tab_stops = context.tab_stops
    assert len(tab_stops) == int(count)


@then("tab_stop.alignment is {alignment}")
def then_tab_stop_alignment_is_alignment(context, alignment):
    expected_value = getattr(WD_TAB_ALIGNMENT, alignment)
    tab_stop = context.tab_stop
    assert tab_stop.alignment == expected_value


@then("tab_stop.leader is {leader}")
def then_tab_stop_leader_is_leader(context, leader):
    expected_value = getattr(WD_TAB_LEADER, leader)
    tab_stop = context.tab_stop
    assert tab_stop.leader == expected_value


@then("tab_stop.position is {position}")
def then_tab_stop_position_is_position(context, position):
    tab_stop = context.tab_stop
    assert tab_stop.position == int(position)


@then("the removed tab stop is no longer present in tab_stops")
def then_the_removed_tab_stop_is_no_longer_present_in_tab_stops(context):
    tab_stops = context.tab_stops
    assert tab_stops[0].position == Inches(1)
    assert tab_stops[1].position == Inches(3)


@then("the tab stops are sequenced in position order")
def then_the_tab_stops_are_sequenced_in_position_order(context):
    tab_stops = context.tab_stops
    for idx in range(len(tab_stops) - 1):
        assert tab_stops[idx].position < tab_stops[idx + 1].position
