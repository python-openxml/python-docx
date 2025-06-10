"""Step implementations for document comments-related features."""

import datetime as dt

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.comments import Comment, Comments
from docx.drawing import Drawing

from helpers import test_docx

# given ====================================================


@given("a Comment object")
def given_a_comment_object(context: Context):
    context.comment = Document(test_docx("comments-rich-para")).comments.get(0)


@given("a Comment object containing an embedded image")
def given_a_comment_object_containing_an_embedded_image(context: Context):
    context.comment = Document(test_docx("comments-rich-para")).comments.get(1)


@given("a Comments object with {count} comments")
def given_a_comments_object_with_count_comments(context: Context, count: str):
    testfile_name = {"0": "doc-default", "4": "comments-rich-para"}[count]
    context.comments = Document(test_docx(testfile_name)).comments


@given("a default Comment object")
def given_a_default_comment_object(context: Context):
    context.comment = Document(test_docx("comments-rich-para")).comments.add_comment()


@given("a document having a comments part")
def given_a_document_having_a_comments_part(context: Context):
    context.document = Document(test_docx("comments-rich-para"))


@given("a document having no comments part")
def given_a_document_having_no_comments_part(context: Context):
    context.document = Document(test_docx("doc-default"))


# when =====================================================


@when('I assign "{author}" to comment.author')
def when_I_assign_author_to_comment_author(context: Context, author: str):
    context.comment.author = author


@when("I assign comment = comments.add_comment()")
def when_I_assign_comment_eq_add_comment(context: Context):
    context.comment = context.comments.add_comment()


@when('I assign comment = comments.add_comment(author="John Doe", initials="JD")')
def when_I_assign_comment_eq_comments_add_comment_with_author_and_initials(context: Context):
    context.comment = context.comments.add_comment(author="John Doe", initials="JD")


@when('I assign comment = document.add_comment(runs, "A comment", "John Doe", "JD")')
def when_I_assign_comment_eq_document_add_comment(context: Context):
    runs = list(context.document.paragraphs[0].runs)
    context.comment = context.document.add_comment(
        runs=runs,
        text="A comment",
        author="John Doe",
        initials="JD",
    )


@when('I assign "{initials}" to comment.initials')
def when_I_assign_initials(context: Context, initials: str):
    context.comment.initials = initials


@when("I assign para_text = comment.paragraphs[0].text")
def when_I_assign_para_text(context: Context):
    context.para_text = context.comment.paragraphs[0].text


@when("I assign paragraph = comment.add_paragraph()")
def when_I_assign_default_add_paragraph(context: Context):
    context.paragraph = context.comment.add_paragraph()


@when("I assign paragraph = comment.add_paragraph(text, style)")
def when_I_assign_add_paragraph_with_text_and_style(context: Context):
    context.para_text = text = "Comment text"
    context.para_style = style = "Normal"
    context.paragraph = context.comment.add_paragraph(text, style)


@when("I assign run = paragraph.add_run()")
def when_I_assign_paragraph_add_run(context: Context):
    context.run = context.paragraph.add_run()


@when("I call comments.get(2)")
def when_I_call_comments_get_2(context: Context):
    context.comment = context.comments.get(2)


# then =====================================================


@then("comment is a Comment object")
def then_comment_is_a_Comment_object(context: Context):
    assert type(context.comment) is Comment


@then('comment.author == "{author}"')
def then_comment_author_eq_author(context: Context, author: str):
    actual = context.comment.author
    assert actual == author, f"expected author '{author}', got '{actual}'"


@then("comment.author is the author of the comment")
def then_comment_author_is_the_author_of_the_comment(context: Context):
    actual = context.comment.author
    assert actual == "Steve Canny", f"expected author 'Steve Canny', got '{actual}'"


@then("comment.comment_id == 0")
def then_comment_id_is_0(context: Context):
    assert context.comment.comment_id == 0


@then("comment.comment_id is the comment identifier")
def then_comment_comment_id_is_the_comment_identifier(context: Context):
    assert context.comment.comment_id == 0


@then("comment.initials is the initials of the comment author")
def then_comment_initials_is_the_initials_of_the_comment_author(context: Context):
    initials = context.comment.initials
    assert initials == "SJC", f"expected initials 'SJC', got '{initials}'"


@then('comment.initials == "{initials}"')
def then_comment_initials_eq_initials(context: Context, initials: str):
    actual = context.comment.initials
    assert actual == initials, f"expected initials '{initials}', got '{actual}'"


@then("comment.paragraphs[{idx}] == paragraph")
def then_comment_paragraphs_idx_eq_paragraph(context: Context, idx: str):
    actual = context.comment.paragraphs[int(idx)]._p
    expected = context.paragraph._p
    assert actual == expected, "paragraphs do not compare equal"


@then('comment.paragraphs[{idx}].style.name == "{style}"')
def then_comment_paragraphs_idx_style_name_eq_style(context: Context, idx: str, style: str):
    actual = context.comment.paragraphs[int(idx)]._p.style
    expected = style
    assert actual == expected, f"expected style name '{expected}', got '{actual}'"


@then('comment.text == "{text}"')
def then_comment_text_eq_text(context: Context, text: str):
    actual = context.comment.text
    expected = text
    assert actual == expected, f"expected text '{expected}', got '{actual}'"


@then("comment.timestamp is the date and time the comment was authored")
def then_comment_timestamp_is_the_date_and_time_the_comment_was_authored(context: Context):
    assert context.comment.timestamp == dt.datetime(2025, 6, 7, 11, 20, 0, tzinfo=dt.timezone.utc)


@then("comments.get({id}) == comment")
def then_comments_get_comment_id_eq_comment(context: Context, id: str):
    comment_id = int(id)
    comment = context.comments.get(comment_id)

    assert type(comment) is Comment, f"expected a Comment object, got {type(comment)}"
    assert comment.comment_id == comment_id, (
        f"expected comment_id '{comment_id}', got '{comment.comment_id}'"
    )


@then("document.comments is a Comments object")
def then_document_comments_is_a_Comments_object(context: Context):
    document = context.document
    assert type(document.comments) is Comments


@then("I can extract the image from the comment")
def then_I_can_extract_the_image_from_the_comment(context: Context):
    paragraph = context.comment.paragraphs[0]
    run = paragraph.runs[2]
    drawing = next(d for d in run.iter_inner_content() if isinstance(d, Drawing))
    assert drawing.has_picture

    image = drawing.image

    assert image.content_type == "image/jpeg", f"got {image.content_type}"
    assert image.filename == "image.jpg", f"got {image.filename}"
    assert image.sha1 == "1be010ea47803b00e140b852765cdf84f491da47", f"got {image.sha1}"


@then("iterating comments yields {count} Comment objects")
def then_iterating_comments_yields_count_comments(context: Context, count: str):
    comment_iter = iter(context.comments)

    comment = next(comment_iter)
    assert type(comment) is Comment, f"expected a Comment object, got {type(comment)}"

    remaining = list(comment_iter)
    assert len(remaining) == int(count) - 1, "iterating comments did not yield the expected count"


@then("len(comment.paragraphs) == {count}")
def then_len_comment_paragraphs_eq_count(context: Context, count: str):
    actual = len(context.comment.paragraphs)
    expected = int(count)
    assert actual == expected, f"expected len(comment.paragraphs) of {expected}, got {actual}"


@then("len(comments) == {count}")
def then_len_comments_eq_count(context: Context, count: str):
    actual = len(context.comments)
    expected = int(count)
    assert actual == expected, f"expected len(comments) of {expected}, got {actual}"


@then("para_text is the text of the first paragraph in the comment")
def then_para_text_is_the_text_of_the_first_paragraph_in_the_comment(context: Context):
    actual = context.para_text
    expected = "Text with hyperlink https://google.com embedded."
    assert actual == expected, f"expected para_text '{expected}', got '{actual}'"


@then("paragraph.style == style")
def then_paragraph_style_eq_known_style(context: Context):
    actual = context.paragraph.style.name
    expected = context.para_style
    assert actual == expected, f"expected paragraph.style '{expected}', got '{actual}'"


@then('paragraph.style == "{style}"')
def then_paragraph_style_eq_style(context: Context, style: str):
    actual = context.paragraph._p.style
    expected = style
    assert actual == expected, f"expected paragraph.style '{expected}', got '{actual}'"


@then("paragraph.text == text")
def then_paragraph_text_eq_known_text(context: Context):
    actual = context.paragraph.text
    expected = context.para_text
    assert actual == expected, f"expected paragraph.text '{expected}', got '{actual}'"


@then('paragraph.text == ""')
def then_paragraph_text_eq_text(context: Context):
    actual = context.paragraph.text
    expected = ""
    assert actual == expected, f"expected paragraph.text '{expected}', got '{actual}'"


@then("run.iter_inner_content() yields a single Picture drawing")
def then_run_iter_inner_content_yields_a_single_picture_drawing(context: Context):
    inner_content = list(context.run.iter_inner_content())

    assert len(inner_content) == 1, (
        f"expected a single inner content element, got {len(inner_content)}"
    )
    inner_content_item = inner_content[0]
    assert isinstance(inner_content_item, Drawing)
    assert inner_content_item.has_picture


@then("the result is a Comment object with id 2")
def then_the_result_is_a_comment_object_with_id_2(context: Context):
    comment = context.comment
    assert type(comment) is Comment, f"expected a Comment object, got {type(comment)}"
    assert comment.comment_id == 2, f"expected comment_id `2`, got '{comment.comment_id}'"
