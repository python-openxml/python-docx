"""Step implementations for styles-related features."""

from behave import given, then, when

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.styles.latent import LatentStyles, _LatentStyle
from docx.styles.style import BaseStyle
from docx.text.font import Font
from docx.text.parfmt import ParagraphFormat

from helpers import bool_vals, test_docx, tri_state_vals

style_types = {
    "WD_STYLE_TYPE.CHARACTER": WD_STYLE_TYPE.CHARACTER,
    "WD_STYLE_TYPE.PARAGRAPH": WD_STYLE_TYPE.PARAGRAPH,
    "WD_STYLE_TYPE.LIST": WD_STYLE_TYPE.LIST,
    "WD_STYLE_TYPE.TABLE": WD_STYLE_TYPE.TABLE,
}


# given ===================================================


@given("a document having a styles part")
def given_a_document_having_a_styles_part(context):
    docx_path = test_docx("sty-having-styles-part")
    context.document = Document(docx_path)


@given("a document having known styles")
def given_a_document_having_known_styles(context):
    docx_path = test_docx("sty-known-styles")
    document = Document(docx_path)
    context.document = document
    context.style_count = len(document.styles)


@given("a document having no styles part")
def given_a_document_having_no_styles_part(context):
    docx_path = test_docx("sty-having-no-styles-part")
    context.document = Document(docx_path)


@given("a latent style collection")
def given_a_latent_style_collection(context):
    document = Document(test_docx("sty-known-styles"))
    context.latent_styles = document.styles.latent_styles


@given("a latent style having a known name")
def given_a_latent_style_having_a_known_name(context):
    document = Document(test_docx("sty-known-styles"))
    latent_styles = list(document.styles.latent_styles)
    context.latent_style = latent_styles[0]  # should be 'Normal'


@given("a latent style having priority of {setting}")
def given_a_latent_style_having_priority_of_setting(context, setting):
    latent_style_name = {
        "42": "Normal",
        "no setting": "Subtitle",
    }[setting]
    document = Document(test_docx("sty-known-styles"))
    latent_styles = document.styles.latent_styles
    context.latent_style = latent_styles[latent_style_name]


@given("a latent style having {prop_name} set {setting}")
def given_a_latent_style_having_prop_setting(context, prop_name, setting):
    latent_style_name = {
        "on": "Normal",
        "off": "Title",
        "no setting": "Subtitle",
    }[setting]
    document = Document(test_docx("sty-known-styles"))
    latent_styles = document.styles.latent_styles
    context.latent_style = latent_styles[latent_style_name]


@given("a latent styles object with known defaults")
def given_a_latent_styles_object_with_known_defaults(context):
    document = Document(test_docx("sty-known-styles"))
    context.latent_styles = document.styles.latent_styles


@given("a style based on {base_style}")
def given_a_style_based_on_setting(context, base_style):
    style_name = {
        "no style": "Base",
        "Normal": "Sub Normal",
        "Base": "Citation",
    }[base_style]
    document = Document(test_docx("sty-known-styles"))
    context.styles = document.styles
    context.style = document.styles[style_name]


@given("a style having a known {attr_name}")
def given_a_style_having_a_known_attr_name(context, attr_name):
    docx_path = test_docx("sty-having-styles-part")
    document = Document(docx_path)
    context.style = document.styles["Normal"]


@given("a style having hidden set {setting}")
def given_a_style_having_hidden_set_setting(context, setting):
    document = Document(test_docx("sty-behav-props"))
    style_name = {
        "on": "Foo",
        "off": "Bar",
        "no setting": "Baz",
    }[setting]
    context.style = document.styles[style_name]


@given("a style having locked set {setting}")
def given_a_style_having_locked_setting(context, setting):
    document = Document(test_docx("sty-behav-props"))
    style_name = {
        "on": "Foo",
        "off": "Bar",
        "no setting": "Baz",
    }[setting]
    context.style = document.styles[style_name]


@given("a style having next paragraph style set to {setting}")
def given_a_style_having_next_paragraph_style_setting(context, setting):
    document = Document(test_docx("sty-known-styles"))
    style_name = {
        "Sub Normal": "Citation",
        "Foobar": "Sub Normal",
        "Base": "Foo",
        "no setting": "Base",
    }[setting]
    context.styles = document.styles
    context.style = document.styles[style_name]


@given("a style having priority of {setting}")
def given_a_style_having_priority_of_setting(context, setting):
    document = Document(test_docx("sty-behav-props"))
    style_name = {
        "no setting": "Baz",
        "42": "Foo",
    }[setting]
    context.style = document.styles[style_name]


@given("a style having quick-style set {setting}")
def given_a_style_having_quick_style_setting(context, setting):
    document = Document(test_docx("sty-behav-props"))
    style_name = {
        "on": "Foo",
        "off": "Bar",
        "no setting": "Baz",
    }[setting]
    context.style = document.styles[style_name]


@given("a style having unhide-when-used set {setting}")
def given_a_style_having_unhide_when_used_setting(context, setting):
    document = Document(test_docx("sty-behav-props"))
    style_name = {
        "on": "Foo",
        "off": "Bar",
        "no setting": "Baz",
    }[setting]
    context.style = document.styles[style_name]


@given("a style of type {style_type}")
def given_a_style_of_type(context, style_type):
    document = Document(test_docx("sty-known-styles"))
    name = {
        "WD_STYLE_TYPE.CHARACTER": "Default Paragraph Font",
        "WD_STYLE_TYPE.LIST": "No List",
        "WD_STYLE_TYPE.PARAGRAPH": "Normal",
        "WD_STYLE_TYPE.TABLE": "Normal Table",
    }[style_type]
    context.style = document.styles[name]


@given("the style collection of a document")
def given_the_style_collection_of_a_document(context):
    document = Document(test_docx("sty-known-styles"))
    context.styles = document.styles


# when =====================================================


@when("I add a latent style named 'Foobar'")
def when_I_add_a_latent_style_named_Foobar(context):
    latent_styles = context.document.styles.latent_styles
    context.latent_styles = latent_styles
    context.latent_style_count = len(latent_styles)
    latent_styles.add_latent_style("Foobar")


@when("I assign a new name to the style")
def when_I_assign_a_new_name_to_the_style(context):
    context.style.name = "Foobar"


@when("I assign a new value to style.style_id")
def when_I_assign_a_new_value_to_style_style_id(context):
    context.style.style_id = "Foo42"


@when("I assign {value} to latent_style.{prop_name}")
def when_I_assign_value_to_latent_style_prop(context, value, prop_name):
    latent_style = context.latent_style
    new_value = tri_state_vals[value] if value in tri_state_vals else int(value)
    setattr(latent_style, prop_name, new_value)


@when("I assign {value} to latent_styles.{prop_name}")
def when_I_assign_value_to_latent_styles_prop(context, value, prop_name):
    latent_styles = context.latent_styles
    new_value = bool_vals[value] if value in bool_vals else int(value)
    setattr(latent_styles, prop_name, new_value)


@when("I assign {value_key} to style.base_style")
def when_I_assign_value_to_style_base_style(context, value_key):
    value = {
        "None": None,
        "styles['Normal']": context.styles["Normal"],
        "styles['Base']": context.styles["Base"],
    }[value_key]
    context.style.base_style = value


@when("I assign {value} to style.hidden")
def when_I_assign_value_to_style_hidden(context, value):
    style, new_value = context.style, tri_state_vals[value]
    style.hidden = new_value


@when("I assign {value} to style.locked")
def when_I_assign_value_to_style_locked(context, value):
    style, new_value = context.style, bool_vals[value]
    style.locked = new_value


@when("I assign {value} to style.next_paragraph_style")
def when_I_assign_value_to_style_next_paragraph_style(context, value):
    styles, style = context.styles, context.style
    new_value = None if value == "None" else styles[value]
    style.next_paragraph_style = new_value


@when("I assign {value} to style.priority")
def when_I_assign_value_to_style_priority(context, value):
    style = context.style
    new_value = None if value == "None" else int(value)
    style.priority = new_value


@when("I assign {value} to style.quick_style")
def when_I_assign_value_to_style_quick_style(context, value):
    style, new_value = context.style, bool_vals[value]
    style.quick_style = new_value


@when("I assign {value} to style.unhide_when_used")
def when_I_assign_value_to_style_unhide_when_used(context, value):
    style, new_value = context.style, bool_vals[value]
    style.unhide_when_used = new_value


@when("I call add_style('{name}', {type_str}, builtin={builtin_str})")
def when_I_call_add_style(context, name, type_str, builtin_str):
    styles = context.document.styles
    type = style_types[type_str]
    builtin = bool_vals[builtin_str]
    styles.add_style(name, type, builtin=builtin)


@when("I delete a latent style")
def when_I_delete_a_latent_style(context):
    latent_styles = context.document.styles.latent_styles
    context.latent_styles = latent_styles
    context.latent_style_count = len(latent_styles)
    latent_styles["Normal"].delete()


@when("I delete a style")
def when_I_delete_a_style(context):
    context.document.styles["No List"].delete()


# then =====================================================


@then("I can access a latent style by name")
def then_I_can_access_a_latent_style_by_name(context):
    latent_styles = context.latent_styles
    latent_style = latent_styles["Colorful Shading"]
    assert isinstance(latent_style, _LatentStyle)


@then("I can access a style by its UI name")
def then_I_can_access_a_style_by_its_UI_name(context):
    styles = context.document.styles
    style = styles["Default Paragraph Font"]
    assert isinstance(style, BaseStyle)


@then("I can access a style by style id")
def then_I_can_access_a_style_by_style_id(context):
    styles = context.document.styles
    style = styles["DefaultParagraphFont"]
    assert isinstance(style, BaseStyle)


@then("I can iterate over its styles")
def then_I_can_iterate_over_its_styles(context):
    styles = list(context.document.styles)
    assert len(styles) > 0
    assert all(isinstance(s, BaseStyle) for s in styles)


@then("I can iterate over the latent styles")
def then_I_can_iterate_over_the_latent_styles(context):
    latent_styles = list(context.latent_styles)
    assert len(latent_styles) == 137
    assert all(isinstance(ls, _LatentStyle) for ls in latent_styles)


@then("latent_style.name is the known name")
def then_latent_style_name_is_the_known_name(context):
    latent_style = context.latent_style
    assert latent_style.name == "Normal"


@then("latent_style.priority is {value}")
def then_latent_style_priority_is_value(context, value):
    latent_style = context.latent_style
    expected_value = None if value == "None" else int(value)
    assert latent_style.priority == expected_value


@then("latent_style.{prop_name} is {value}")
def then_latent_style_prop_name_is_value(context, prop_name, value):
    latent_style = context.latent_style
    actual_value = getattr(latent_style, prop_name)
    expected_value = tri_state_vals[value]
    assert actual_value == expected_value


@then("latent_styles['Foobar'] is a latent style")
def then_latentStyles_Foobar_is_a_latent_style(context):
    latent_styles = context.latent_styles
    latent_style = latent_styles["Foobar"]
    assert isinstance(latent_style, _LatentStyle)


@then("latent_styles.{prop_name} is {value}")
def then_latent_styles_prop_name_is_value(context, prop_name, value):
    latent_styles = context.latent_styles
    expected_value = bool_vals[value] if value in bool_vals else int(value)
    actual_value = getattr(latent_styles, prop_name)
    assert actual_value == expected_value


@then("len(latent_styles) is 137")
def then_len_latent_styles_is_137(context):
    assert len(context.latent_styles) == 137


@then("len(styles) is {style_count_str}")
def then_len_styles_is_style_count(context, style_count_str):
    assert len(context.document.styles) == int(style_count_str)


@then("style.base_style is {value_key}")
def then_style_base_style_is_value(context, value_key):
    expected_value = {
        "None": None,
        "styles['Normal']": context.styles["Normal"],
        "styles['Base']": context.styles["Base"],
    }[value_key]
    style = context.style
    assert style.base_style == expected_value


@then("style.builtin is {builtin_str}")
def then_style_builtin_is_builtin(context, builtin_str):
    style = context.style
    builtin = bool_vals[builtin_str]
    assert style.builtin == builtin


@then("style.font is the Font object for the style")
def then_style_font_is_the_Font_object_for_the_style(context):
    style = context.style
    font = style.font
    assert isinstance(font, Font)
    assert font.element is style.element


@then("style.hidden is {value}")
def then_style_hidden_is_value(context, value):
    style, expected_value = context.style, tri_state_vals[value]
    assert style.hidden is expected_value


@then("style.locked is {value}")
def then_style_locked_is_value(context, value):
    style, expected_value = context.style, bool_vals[value]
    assert style.locked is expected_value


@then("style.name is the {which} name")
def then_style_name_is_the_which_name(context, which):
    expected_name = {
        "known": "Normal",
        "new": "Foobar",
    }[which]
    style = context.style
    assert style.name == expected_name


@then("style.next_paragraph_style is {value}")
def then_style_next_paragraph_style_is_value(context, value):
    style, styles = context.style, context.styles
    actual_value = style.next_paragraph_style
    expected_value = styles[value]
    assert actual_value == expected_value, "got %s" % actual_value


@then("style.paragraph_format is the ParagraphFormat object for the style")
def then_style_paragraph_format_is_the_ParagraphFormat_object(context):
    style = context.style
    paragraph_format = style.paragraph_format
    assert isinstance(paragraph_format, ParagraphFormat)
    assert paragraph_format.element is style.element


@then("style.priority is {value}")
def then_style_priority_is_value(context, value):
    style = context.style
    expected_value = None if value == "None" else int(value)
    assert style.priority == expected_value


@then("style.quick_style is {value}")
def then_style_quick_style_is_value(context, value):
    style, expected_value = context.style, bool_vals[value]
    assert style.quick_style is expected_value


@then("style.style_id is the {which} style id")
def then_style_style_id_is_the_which_style_id(context, which):
    expected_style_id = {
        "known": "Normal",
        "new": "Foo42",
    }[which]
    style = context.style
    assert style.style_id == expected_style_id


@then("style.type is the known type")
def then_style_type_is_the_known_type(context):
    style = context.style
    assert style.type == WD_STYLE_TYPE.PARAGRAPH


@then("style.type is {type_str}")
def then_style_type_is_type(context, type_str):
    style = context.style
    style_type = style_types[type_str]
    assert style.type == style_type


@then("style.unhide_when_used is {value}")
def then_style_unhide_when_used_is_value(context, value):
    style, expected_value = context.style, bool_vals[value]
    assert style.unhide_when_used is expected_value


@then("styles.latent_styles is the LatentStyles object for the document")
def then_styles_latent_styles_is_the_LatentStyles_object(context):
    styles = context.styles
    context.latent_styles = latent_styles = styles.latent_styles
    assert isinstance(latent_styles, LatentStyles)
    assert latent_styles.element is styles.element.latentStyles


@then("styles['{name}'] is a style")
def then_styles_name_is_a_style(context, name):
    styles = context.document.styles
    style = context.style = styles[name]
    assert isinstance(style, BaseStyle)


@then("the deleted latent style is not in the latent styles collection")
def then_the_deleted_latent_style_is_not_in_the_collection(context):
    latent_styles = context.latent_styles
    try:
        latent_styles["Normal"]
    except KeyError:
        return
    raise AssertionError("Latent style not deleted")


@then("the deleted style is not in the styles collection")
def then_the_deleted_style_is_not_in_the_styles_collection(context):
    document = context.document
    try:
        document.styles["No List"]
    except KeyError:
        return
    raise AssertionError("Style not deleted")


@then("the document has one additional latent style")
def then_the_document_has_one_additional_latent_style(context):
    latent_styles = context.document.styles.latent_styles
    latent_style_count = len(latent_styles)
    expected_count = context.latent_style_count + 1
    assert latent_style_count == expected_count


@then("the document has one additional style")
def then_the_document_has_one_additional_style(context):
    document = context.document
    style_count = len(document.styles)
    expected_style_count = context.style_count + 1
    assert style_count == expected_style_count


@then("the document has one fewer latent styles")
def then_the_document_has_one_fewer_latent_styles(context):
    latent_style_count = len(context.latent_styles)
    expected_count = context.latent_style_count - 1
    assert latent_style_count == expected_count


@then("the document has one fewer styles")
def then_the_document_has_one_fewer_styles(context):
    document = context.document
    style_count = len(document.styles)
    expected_style_count = context.style_count - 1
    assert style_count == expected_style_count
