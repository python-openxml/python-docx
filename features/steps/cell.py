# encoding: utf-8

"""
Step implementations for table cell-related features
"""

from __future__ import absolute_import, print_function, unicode_literals

from behave import given, then, when

from docx import Document

from helpers import test_docx


# given ===================================================

@given('a table cell')
def given_a_table_cell(context):
    table = Document(test_docx('tbl-2x2-table')).tables[0]
    context.cell = table.cell(0, 0)


# when =====================================================

@when('I add a 2 x 2 table into the first cell')
def when_I_add_a_2x2_table_into_the_first_cell(context):
    context.table_ = context.cell.add_table(2, 2)


@when('I assign a string to the cell text attribute')
def when_assign_string_to_cell_text_attribute(context):
    cell = context.cell
    text = 'foobar'
    cell.text = text
    context.expected_text = text


# then =====================================================

@then('cell.tables[0] is a 2 x 2 table')
def then_cell_tables_0_is_a_2x2_table(context):
    cell = context.cell
    table = cell.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2


@then('the cell contains the string I assigned')
def then_cell_contains_string_assigned(context):
    cell, expected_text = context.cell, context.expected_text
    text = cell.paragraphs[0].runs[0].text
    msg = "expected '%s', got '%s'" % (expected_text, text)
    assert text == expected_text, msg
