"""Step implementations for block content containers."""

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.table import Table

from helpers import test_docx

# given ===================================================


@given("a _Cell object with paragraphs and tables")
def given_a_cell_with_paragraphs_and_tables(context: Context):
    context.cell = (
        Document(test_docx("blk-paras-and-tables")).tables[1].rows[0].cells[0]
    )


@given("a Document object with paragraphs and tables")
def given_a_document_with_paragraphs_and_tables(context: Context):
    context.document = Document(test_docx("blk-paras-and-tables"))


@given("a document containing a table")
def given_a_document_containing_a_table(context: Context):
    context.document = Document(test_docx("blk-containing-table"))


@given("a Footer object with paragraphs and tables")
def given_a_footer_with_paragraphs_and_tables(context: Context):
    context.footer = Document(test_docx("blk-paras-and-tables")).sections[0].footer


@given("a Header object with paragraphs and tables")
def given_a_header_with_paragraphs_and_tables(context: Context):
    context.header = Document(test_docx("blk-paras-and-tables")).sections[0].header


@given("a paragraph")
def given_a_paragraph(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph()


# when ====================================================


@when("I add a paragraph")
def when_add_paragraph(context: Context):
    document = context.document
    context.p = document.add_paragraph()


@when("I add a table")
def when_add_table(context: Context):
    rows, cols = 2, 2
    context.document.add_table(rows, cols)


# then =====================================================


@then("cell.iter_inner_content() produces the block-items in document order")
def then_cell_iter_inner_content_produces_the_block_items(context: Context):
    actual = [type(item).__name__ for item in context.cell.iter_inner_content()]
    expected = ["Paragraph", "Table", "Paragraph"]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("document.iter_inner_content() produces the block-items in document order")
def then_document_iter_inner_content_produces_the_block_items(context: Context):
    actual = [type(item).__name__ for item in context.document.iter_inner_content()]
    expected = ["Table", "Paragraph", "Table", "Paragraph", "Table", "Paragraph"]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("footer.iter_inner_content() produces the block-items in document order")
def then_footer_iter_inner_content_produces_the_block_items(context: Context):
    actual = [type(item).__name__ for item in context.footer.iter_inner_content()]
    expected = ["Paragraph", "Table", "Paragraph"]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("header.iter_inner_content() produces the block-items in document order")
def then_header_iter_inner_content_produces_the_block_items(context: Context):
    actual = [type(item).__name__ for item in context.header.iter_inner_content()]
    expected = ["Table", "Paragraph"]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("I can access the table")
def then_can_access_table(context: Context):
    table = context.document.tables[-1]
    assert isinstance(table, Table)


@then("the new table appears in the document")
def then_new_table_appears_in_document(context: Context):
    table = context.document.tables[-1]
    assert isinstance(table, Table)
