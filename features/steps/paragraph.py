"""Step implementations for paragraph-related features."""

from __future__ import annotations

from typing import Any

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.parfmt import ParagraphFormat

from helpers import saved_docx_path, test_docx, test_text

# given ===================================================


@given("a document containing three paragraphs")
def given_a_document_containing_three_paragraphs(context: Context):
    document = Document()
    document.add_paragraph("foo")
    document.add_paragraph("bar")
    document.add_paragraph("baz")
    context.document = document


@given("a paragraph having {align_type} alignment")
def given_a_paragraph_align_type_alignment(context: Context, align_type: str):
    paragraph_idx = {
        "inherited": 0,
        "left": 1,
        "center": 2,
        "right": 3,
        "justified": 4,
    }[align_type]
    document = Document(test_docx("par-alignment"))
    context.paragraph = document.paragraphs[paragraph_idx]


@given("a paragraph having {style_state} style")
def given_a_paragraph_having_style(context: Context, style_state: str):
    paragraph_idx = {
        "no specified": 0,
        "a missing": 1,
        "Heading 1": 2,
        "Body Text": 3,
    }[style_state]
    document = context.document = Document(test_docx("par-known-styles"))
    context.paragraph = document.paragraphs[paragraph_idx]


@given("a paragraph having {zero_or_more} hyperlinks")
def given_a_paragraph_having_hyperlinks(context: Context, zero_or_more: str):
    paragraph_idx = {
        "no": 0,
        "one": 1,
        "three": 2,
    }[zero_or_more]
    document = context.document = Document(test_docx("par-hyperlinks"))
    context.paragraph = document.paragraphs[paragraph_idx]


@given("a paragraph having {zero_or_more} rendered page breaks")
def given_a_paragraph_having_rendered_page_breaks(context: Context, zero_or_more: str):
    paragraph_idx = {
        "no": 0,
        "one": 1,
        "two": 2,
    }[zero_or_more]
    document = Document(test_docx("par-rendered-page-breaks"))
    context.paragraph = document.paragraphs[paragraph_idx]


@given("a paragraph with content and formatting")
def given_a_paragraph_with_content_and_formatting(context: Context):
    document = Document(test_docx("par-known-paragraphs"))
    context.paragraph = document.paragraphs[0]


# when ====================================================


@when("I add a run to the paragraph")
def when_add_new_run_to_paragraph(context: Context):
    context.run = context.p.add_run()


@when("I assign a {style_type} to paragraph.style")
def when_I_assign_a_style_type_to_paragraph_style(context: Context, style_type: str):
    paragraph = context.paragraph
    style = context.style = context.document.styles["Heading 1"]
    style_spec = {
        "style object": style,
        "style name": "Heading 1",
    }[style_type]
    paragraph.style = style_spec


@when("I clear the paragraph content")
def when_I_clear_the_paragraph_content(context: Context):
    context.paragraph.clear()


@when("I insert a paragraph above the second paragraph")
def when_I_insert_a_paragraph_above_the_second_paragraph(context: Context):
    paragraph = context.document.paragraphs[1]
    paragraph.insert_paragraph_before("foobar", "Heading1")


@when("I set the paragraph text")
def when_I_set_the_paragraph_text(context: Context):
    context.paragraph.text = "bar\tfoo\r"


# then =====================================================


@then("paragraph.contains_page_break is {value}")
def then_paragraph_contains_page_break_is_value(context: Context, value: str):
    actual_value = context.paragraph.contains_page_break
    expected_value = {"True": True, "False": False}[value]
    assert (
        actual_value == expected_value
    ), f"expected: {expected_value}, got: {actual_value}"


@then("paragraph.hyperlinks contains only Hyperlink instances")
def then_paragraph_hyperlinks_contains_only_Hyperlink_instances(context: Context):
    assert all(
        type(item).__name__ == "Hyperlink" for item in context.paragraph.hyperlinks
    )


@then("paragraph.hyperlinks has length {value}")
def then_paragraph_hyperlinks_has_length(context: Context, value: str):
    expected_value = int(value)
    assert len(context.paragraph.hyperlinks) == expected_value


@then("paragraph.iter_inner_content() generates the paragraph runs and hyperlinks")
def then_paragraph_iter_inner_content_generates_runs_and_hyperlinks(context: Context):
    assert [type(item).__name__ for item in context.paragraph.iter_inner_content()] == [
        "Run",
        "Hyperlink",
        "Run",
        "Hyperlink",
        "Run",
        "Hyperlink",
        "Run",
    ]


@then("paragraph.paragraph_format is its ParagraphFormat object")
def then_paragraph_paragraph_format_is_its_parfmt_object(context: Context):
    paragraph = context.paragraph
    paragraph_format = paragraph.paragraph_format
    assert isinstance(paragraph_format, ParagraphFormat)
    assert paragraph_format.element is paragraph._element


@then("paragraph.rendered_page_breaks has length {value}")
def then_paragraph_rendered_page_breaks_has_length(context: Context, value: str):
    actual_value = len(context.paragraph.rendered_page_breaks)
    expected_value = int(value)
    assert (
        actual_value == expected_value
    ), f"got: {actual_value}, expected: {expected_value}"


@then("paragraph.rendered_page_breaks contains only RenderedPageBreak instances")
def then_paragraph_rendered_page_breaks_contains_only_RenderedPageBreak_instances(
    context: Context,
):
    assert all(
        type(item).__name__ == "RenderedPageBreak"
        for item in context.paragraph.rendered_page_breaks
    )


@then("paragraph.style is {value_key}")
def then_paragraph_style_is_value(context: Context, value_key: str):
    styles = context.document.styles
    expected_value = {
        "Normal": styles["Normal"],
        "Heading 1": styles["Heading 1"],
        "Body Text": styles["Body Text"],
    }[value_key]
    paragraph = context.paragraph
    assert paragraph.style == expected_value


@then("paragraph.text contains the text of both the runs and the hyperlinks")
def then_paragraph_text_contains_the_text_of_both_the_runs_and_the_hyperlinks(
    context: Context,
):
    actual = context.paragraph.text
    expected = "Three hyperlinks: the first one here, the second one, and the third."
    assert actual == expected, f"expected:\n'{expected}'\n\ngot:\n'{actual}'"


@then("the document contains four paragraphs")
def then_the_document_contains_four_paragraphs(context: Context):
    assert len(context.document.paragraphs) == 4


@then("the document contains the text I added")
def then_document_contains_text_I_added(context: Context):
    document = Document(saved_docx_path)
    paragraphs = document.paragraphs
    paragraph = paragraphs[-1]
    run = paragraph.runs[0]
    actual = run.text
    expected = test_text
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("the paragraph alignment property value is {align_value}")
def then_the_paragraph_alignment_prop_value_is_value(
    context: Context, align_value: str
):
    expected_value: Any = {
        "None": None,
        "WD_ALIGN_PARAGRAPH.LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,  # pyright: ignore
        "WD_ALIGN_PARAGRAPH.CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,  # pyright: ignore
        "WD_ALIGN_PARAGRAPH.RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,  # pyright: ignore
    }[align_value]
    assert context.paragraph.alignment == expected_value


@then("the paragraph formatting is preserved")
def then_the_paragraph_formatting_is_preserved(context: Context):
    paragraph = context.paragraph
    assert paragraph.style.name == "Heading 1"


@then("the paragraph has no content")
def then_the_paragraph_has_no_content(context: Context):
    assert context.paragraph.text == ""


@then("the paragraph has the style I set")
def then_the_paragraph_has_the_style_I_set(context: Context):
    paragraph, expected_style = context.paragraph, context.style
    assert paragraph.style == expected_style


@then("the paragraph has the text I set")
def then_the_paragraph_has_the_text_I_set(context: Context):
    actual = context.paragraph.text
    expected = "bar\tfoo\n"
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("the style of the second paragraph matches the style I set")
def then_the_style_of_the_second_paragraph_matches_the_style_I_set(context: Context):
    second_paragraph = context.document.paragraphs[1]
    assert second_paragraph.style.name == "Heading 1"


@then("the text of the second paragraph matches the text I set")
def then_the_text_of_the_second_paragraph_matches_the_text_I_set(context: Context):
    second_paragraph = context.document.paragraphs[1]
    assert second_paragraph.text == "foobar"
