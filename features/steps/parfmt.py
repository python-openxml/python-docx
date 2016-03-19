# encoding: utf-8

"""
Step implementations for paragraph format-related features.
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from behave import given, then, when

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from docx.text.tabstops import TabStops

from helpers import test_docx


# given ===================================================

@given('a paragraph format')
def given_a_paragraph_format(context):
    document = Document(test_docx('tab-stops'))
    context.paragraph_format = document.paragraphs[0].paragraph_format


@given('a paragraph format having {prop_name} set {setting}')
def given_a_paragraph_format_having_prop_set(context, prop_name, setting):
    style_name = {
        'to inherit': 'Normal',
        'On':         'Base',
        'Off':        'Citation',
    }[setting]
    document = Document(test_docx('sty-known-styles'))
    context.paragraph_format = document.styles[style_name].paragraph_format


@given('a paragraph format having {setting} line spacing')
def given_a_paragraph_format_having_setting_line_spacing(context, setting):
    style_name = {
        'inherited': 'Normal',
        '14 pt':     'Base',
        'double':    'Citation',
    }[setting]
    document = Document(test_docx('sty-known-styles'))
    context.paragraph_format = document.styles[style_name].paragraph_format


@given('a paragraph format having {setting} space {side}')
def given_a_paragraph_format_having_setting_spacing(context, setting, side):
    style_name = 'Normal' if setting == 'inherited' else 'Base'
    document = Document(test_docx('sty-known-styles'))
    context.paragraph_format = document.styles[style_name].paragraph_format


@given('a paragraph format having {type} alignment')
def given_a_paragraph_format_having_align_type_alignment(context, type):
    style_name = {
        'inherited': 'Normal',
        'center':    'Base',
        'right':     'Citation',
    }[type]
    document = Document(test_docx('sty-known-styles'))
    context.paragraph_format = document.styles[style_name].paragraph_format


@given('a paragraph format having {type} indent of {value}')
def given_a_paragraph_format_having_type_indent_value(context, type, value):
    style_name = {
        'inherit':  'Normal',
        '18 pt':    'Base',
        '17.3 pt':  'Base',
        '-17.3 pt': 'Citation',
        '46.1 pt':  'Citation',
    }[value]
    document = Document(test_docx('sty-known-styles'))
    context.paragraph_format = document.styles[style_name].paragraph_format


# when ====================================================

@when('I assign {value} to paragraph_format.line_spacing')
def when_I_assign_value_to_paragraph_format_line_spacing(context, value):
    new_value = {
        'Pt(14)': Pt(14),
        '2':      2,
    }.get(value)
    new_value = float(value) if new_value is None else new_value
    context.paragraph_format.line_spacing = new_value


@when('I assign {value} to paragraph_format.line_spacing_rule')
def when_I_assign_value_to_paragraph_format_line_rule(context, value):
    new_value = {
        'None':                           None,
        'WD_LINE_SPACING.EXACTLY':        WD_LINE_SPACING.EXACTLY,
        'WD_LINE_SPACING.MULTIPLE':       WD_LINE_SPACING.MULTIPLE,
        'WD_LINE_SPACING.SINGLE':         WD_LINE_SPACING.SINGLE,
        'WD_LINE_SPACING.DOUBLE':         WD_LINE_SPACING.DOUBLE,
        'WD_LINE_SPACING.AT_LEAST':       WD_LINE_SPACING.AT_LEAST,
        'WD_LINE_SPACING.ONE_POINT_FIVE': WD_LINE_SPACING.ONE_POINT_FIVE,
    }[value]
    paragraph_format = context.paragraph_format
    paragraph_format.line_spacing_rule = new_value


@when('I assign {value} to paragraph_format.alignment')
def when_I_assign_value_to_paragraph_format_alignment(context, value):
    new_value = {
        'None':                      None,
        'WD_ALIGN_PARAGRAPH.CENTER': WD_ALIGN_PARAGRAPH.CENTER,
        'WD_ALIGN_PARAGRAPH.RIGHT':  WD_ALIGN_PARAGRAPH.RIGHT,
    }[value]
    paragraph_format = context.paragraph_format
    paragraph_format.alignment = new_value


@when('I assign {value} to paragraph_format.space_{side}')
def when_I_assign_value_to_paragraph_format_space(context, value, side):
    paragraph_format = context.paragraph_format
    prop_name = 'space_%s' % side
    new_value = {
        'None':   None,
        'Pt(12)': Pt(12),
        'Pt(18)': Pt(18),
    }[value]
    setattr(paragraph_format, prop_name, new_value)


@when('I assign {value} to paragraph_format.{type_}_indent')
def when_I_assign_value_to_paragraph_format_indent(context, value, type_):
    paragraph_format = context.paragraph_format
    prop_name = '%s_indent' % type_
    value = None if value == 'None' else Pt(float(value.split()[0]))
    setattr(paragraph_format, prop_name, value)


@when('I assign {value} to paragraph_format.{prop_name}')
def when_I_assign_value_to_paragraph_format_prop(context, value, prop_name):
    paragraph_format = context.paragraph_format
    value = {'None': None, 'True': True, 'False': False}[value]
    setattr(paragraph_format, prop_name, value)


# then =====================================================

@then('paragraph_format.tab_stops is a TabStops object')
def then_paragraph_format_tab_stops_is_a_tabstops_object(context):
    tab_stops = context.paragraph_format.tab_stops
    assert isinstance(tab_stops, TabStops)


@then('paragraph_format.alignment is {value}')
def then_paragraph_format_alignment_is_value(context, value):
    expected_value = {
        'None':                      None,
        'WD_ALIGN_PARAGRAPH.LEFT':   WD_ALIGN_PARAGRAPH.LEFT,
        'WD_ALIGN_PARAGRAPH.CENTER': WD_ALIGN_PARAGRAPH.CENTER,
        'WD_ALIGN_PARAGRAPH.RIGHT':  WD_ALIGN_PARAGRAPH.RIGHT,
    }[value]
    paragraph_format = context.paragraph_format
    assert paragraph_format.alignment == expected_value


@then('paragraph_format.line_spacing is {value}')
def then_paragraph_format_line_spacing_is_value(context, value):
    expected_value = (
        None if value == 'None' else
        float(value) if '.' in value else
        int(value)
    )
    paragraph_format = context.paragraph_format

    if expected_value is None or isinstance(expected_value, int):
        assert paragraph_format.line_spacing == expected_value
    else:
        assert abs(paragraph_format.line_spacing - expected_value) < 0.001


@then('paragraph_format.line_spacing_rule is {value}')
def then_paragraph_format_line_spacing_rule_is_value(context, value):
    expected_value = {
        'None':                           None,
        'WD_LINE_SPACING.EXACTLY':        WD_LINE_SPACING.EXACTLY,
        'WD_LINE_SPACING.MULTIPLE':       WD_LINE_SPACING.MULTIPLE,
        'WD_LINE_SPACING.SINGLE':         WD_LINE_SPACING.SINGLE,
        'WD_LINE_SPACING.DOUBLE':         WD_LINE_SPACING.DOUBLE,
        'WD_LINE_SPACING.AT_LEAST':       WD_LINE_SPACING.AT_LEAST,
        'WD_LINE_SPACING.ONE_POINT_FIVE': WD_LINE_SPACING.ONE_POINT_FIVE,
    }[value]
    paragraph_format = context.paragraph_format
    assert paragraph_format.line_spacing_rule == expected_value


@then('paragraph_format.space_{side} is {value}')
def then_paragraph_format_space_side_is_value(context, side, value):
    expected_value = None if value == 'None' else int(value)
    prop_name = 'space_%s' % side
    paragraph_format = context.paragraph_format
    actual_value = getattr(paragraph_format, prop_name)
    assert actual_value == expected_value


@then('paragraph_format.{type_}_indent is {value}')
def then_paragraph_format_type_indent_is_value(context, type_, value):
    expected_value = None if value == 'None' else int(value)
    prop_name = '%s_indent' % type_
    paragraph_format = context.paragraph_format
    actual_value = getattr(paragraph_format, prop_name)
    assert actual_value == expected_value


@then('paragraph_format.{prop_name} is {value}')
def then_paragraph_format_prop_name_is_value(context, prop_name, value):
    expected_value = {'None': None, 'True': True, 'False': False}[value]
    paragraph_format = context.paragraph_format
    actual_value = getattr(paragraph_format, prop_name)
    assert actual_value == expected_value
