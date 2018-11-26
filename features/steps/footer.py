# encoding: utf-8

"""
Step implementations for document settings-related features
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from behave import given, then, when

from docx import Document
from docx.settings import Settings

from helpers import test_docx, tri_state_vals


# given ====================================================

@given('a footer having a is_linked_to_previous property')
def given_a_document_having_a_footer_linked_to_previous(context):
    context.document = Document(test_docx('doc-word-linked_footer'))


@given('a footer having no is_linked_to_previous property')
def given_a_document_having_no_footer_linked_to_previous(context):
    context.document = Document(test_docx('doc-word-no-linked_footer'))


@given('a document section having footer')
def given_a_document_section_having_footer(context):
    context.document = Document(test_docx('a_footer_of_all_types'))

# when =====================================================


@when('I set footer to {value}')
def when_i_link_footer_to_previous_section(context, value):
    document, value = context.document, tri_state_vals[value]
    document.sections[-1].footer.is_linked_to_previous = value

# then =====================================================

@then('document.sections[-1].footer.is_linked_to_previous is {value}')
def then_document_settings_even_and_odd_footer_is_value(context, value):
    document, expected_value = context.document, tri_state_vals[value]
    assert document.sections[-1].footer.is_linked_to_previous is expected_value

@then('a section footer is_linked to previous is {value}')
def then_section_is_linked_to_previous_with_value(context, value):
    document, expected_value = context.document, tri_state_vals[value]
    assert document.sections[-1].footer.is_linked_to_previous is expected_value