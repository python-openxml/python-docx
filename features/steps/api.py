# encoding: utf-8

"""
Step implementations for basic API features
"""

from behave import given, then, when

import docx

from docx import Document

from helpers import test_docx


# given ====================================================


@given("I have python-docx installed")
def given_I_have_python_docx_installed(context):
    pass


# when =====================================================


@when("I call docx.Document() with no arguments")
def when_I_call_docx_Document_with_no_arguments(context):
    context.document = Document()


@when("I call docx.Document() with the path of a .docx file")
def when_I_call_docx_Document_with_the_path_of_a_docx_file(context):
    context.document = Document(test_docx("doc-default"))


# then =====================================================


@then("document is a Document object")
def then_document_is_a_Document_object(context):
    document = context.document
    assert isinstance(document, docx.document.Document)


@then("the last paragraph contains the text I specified")
def then_last_p_contains_specified_text(context):
    document = context.document
    text = context.paragraph_text
    p = document.paragraphs[-1]
    assert p.text == text


@then("the last paragraph has the style I specified")
def then_the_last_paragraph_has_the_style_I_specified(context):
    document, expected_style = context.document, context.style
    paragraph = document.paragraphs[-1]
    assert paragraph.style == expected_style


@then("the last paragraph is the empty paragraph I added")
def then_last_p_is_empty_paragraph_added(context):
    document = context.document
    p = document.paragraphs[-1]
    assert p.text == ""
