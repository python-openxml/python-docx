"""Step implementations for document settings-related features."""

from behave import given, then, when

from docx import Document
from docx.settings import Settings

from helpers import test_docx

# given ====================================================


@given("a document having a settings part")
def given_a_document_having_a_settings_part(context):
    context.document = Document(test_docx("doc-word-default-blank"))


@given("a document having no settings part")
def given_a_document_having_no_settings_part(context):
    context.document = Document(test_docx("set-no-settings-part"))


@given("a Settings object {with_or_without} odd and even page headers as settings")
def given_a_Settings_object_with_or_without_odd_and_even_hdrs(context, with_or_without):
    testfile_name = {"with": "doc-odd-even-hdrs", "without": "sct-section-props"}[
        with_or_without
    ]
    context.settings = Document(test_docx(testfile_name)).settings


# when =====================================================


@when("I assign {bool_val} to settings.odd_and_even_pages_header_footer")
def when_I_assign_value_to_settings_odd_and_even_pages_header_footer(context, bool_val):
    context.settings.odd_and_even_pages_header_footer = eval(bool_val)


# then =====================================================


@then("document.settings is a Settings object")
def then_document_settings_is_a_Settings_object(context):
    document = context.document
    assert type(document.settings) is Settings


@then("settings.odd_and_even_pages_header_footer is {bool_val}")
def then_settings_odd_and_even_pages_header_footer_is(context, bool_val):
    actual = context.settings.odd_and_even_pages_header_footer
    expected = eval(bool_val)
    assert actual == expected, (
        "settings.odd_and_even_pages_header_footer is %s" % actual
    )
