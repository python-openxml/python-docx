# encoding: utf-8

"""
Step implementations for document-related features
"""

from __future__ import absolute_import, print_function, unicode_literals

from behave import given, then, when

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shape import InlineShapes
from docx.shared import Inches
from docx.section import Sections
from docx.styles.styles import Styles
from docx.table import Table
from docx.text.paragraph import Paragraph

from helpers import test_docx, test_file


# given ===================================================

@given('a blank document')
def given_a_blank_document(context):
    context.document = Document(test_docx('doc-word-default-blank'))


@given('a document having built-in styles')
def given_a_document_having_builtin_styles(context):
    context.document = Document()


@given('a document having inline shapes')
def given_a_document_having_inline_shapes(context):
    context.document = Document(test_docx('shp-inline-shape-access'))


@given('a document having sections')
def given_a_document_having_sections(context):
    context.document = Document(test_docx('doc-access-sections'))


@given('a document having styles')
def given_a_document_having_styles(context):
    context.document = Document(test_docx('sty-having-styles-part'))


@given('a document having three tables')
def given_a_document_having_three_tables(context):
    context.document = Document(test_docx('tbl-having-tables'))


@given('a single-section document having portrait layout')
def given_a_single_section_document_having_portrait_layout(context):
    context.document = Document(test_docx('doc-add-section'))
    section = context.document.sections[-1]
    context.original_dimensions = (section.page_width, section.page_height)


@given("a single-section Document object with headers and footers as document")
def given_a_single_section_Document_object_with_headers_and_footers(context):
    context.document = Document(test_docx("doc-add-section"))


# when ====================================================

@when('I add a 2 x 2 table specifying only row and column count')
def when_add_2x2_table_specifying_only_row_and_col_count(context):
    document = context.document
    document.add_table(rows=2, cols=2)


@when('I add a 2 x 2 table specifying style \'{style_name}\'')
def when_add_2x2_table_specifying_style_name(context, style_name):
    document = context.document
    document.add_table(rows=2, cols=2, style=style_name)


@when('I add a heading specifying level={level}')
def when_add_heading_specifying_level(context, level):
    context.document.add_heading(level=int(level))


@when('I add a heading specifying only its text')
def when_add_heading_specifying_only_its_text(context):
    document = context.document
    context.heading_text = text = 'Spam vs. Eggs'
    document.add_heading(text)


@when('I add a page break to the document')
def when_add_page_break_to_document(context):
    document = context.document
    document.add_page_break()


@when('I add a paragraph specifying its style as a {kind}')
def when_I_add_a_paragraph_specifying_its_style_as_a(context, kind):
    document = context.document
    style = context.style = document.styles['Heading 1']
    style_spec = {
        'style object': style,
        'style name':   'Heading 1',
    }[kind]
    document.add_paragraph(style=style_spec)


@when('I add a paragraph specifying its text')
def when_add_paragraph_specifying_text(context):
    document = context.document
    context.paragraph_text = 'foobar'
    document.add_paragraph(context.paragraph_text)


@when('I add a paragraph without specifying text or style')
def when_add_paragraph_without_specifying_text_or_style(context):
    document = context.document
    document.add_paragraph()


@when('I add a picture specifying 1.75" width and 2.5" height')
def when_add_picture_specifying_width_and_height(context):
    document = context.document
    context.picture = document.add_picture(
        test_file('monty-truth.png'),
        width=Inches(1.75), height=Inches(2.5)
    )


@when('I add a picture specifying a height of 1.5 inches')
def when_add_picture_specifying_height(context):
    document = context.document
    context.picture = document.add_picture(
        test_file('monty-truth.png'), height=Inches(1.5)
    )


@when('I add a picture specifying a width of 1.5 inches')
def when_add_picture_specifying_width(context):
    document = context.document
    context.picture = document.add_picture(
        test_file('monty-truth.png'), width=Inches(1.5)
    )


@when('I add a picture specifying only the image file')
def when_add_picture_specifying_only_image_file(context):
    document = context.document
    context.picture = document.add_picture(test_file('monty-truth.png'))


@when('I add an even-page section to the document')
def when_I_add_an_even_page_section_to_the_document(context):
    context.section = context.document.add_section(WD_SECTION.EVEN_PAGE)


@when('I change the new section layout to landscape')
def when_I_change_the_new_section_layout_to_landscape(context):
    new_height, new_width = context.original_dimensions
    section = context.section
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height


@when("I execute section = document.add_section()")
def when_I_execute_section_eq_document_add_section(context):
    context.section = context.document.add_section()


# then ====================================================

@then('document.inline_shapes is an InlineShapes object')
def then_document_inline_shapes_is_an_InlineShapes_object(context):
    document = context.document
    inline_shapes = document.inline_shapes
    assert isinstance(inline_shapes, InlineShapes)


@then('document.paragraphs is a list containing three paragraphs')
def then_document_paragraphs_is_a_list_containing_three_paragraphs(context):
    document = context.document
    paragraphs = document.paragraphs
    assert isinstance(paragraphs, list)
    assert len(paragraphs) == 3
    for paragraph in paragraphs:
        assert isinstance(paragraph, Paragraph)


@then('document.sections is a Sections object')
def then_document_sections_is_a_Sections_object(context):
    sections = context.document.sections
    msg = 'document.sections not instance of Sections'
    assert isinstance(sections, Sections), msg


@then('document.styles is a Styles object')
def then_document_styles_is_a_Styles_object(context):
    styles = context.document.styles
    assert isinstance(styles, Styles)


@then('document.tables is a list containing three tables')
def then_document_tables_is_a_list_containing_three_tables(context):
    document = context.document
    tables = document.tables
    assert isinstance(tables, list)
    assert len(tables) == 3
    for table in tables:
        assert isinstance(table, Table)


@then('the document contains a 2 x 2 table')
def then_the_document_contains_a_2x2_table(context):
    table = context.document.tables[-1]
    assert isinstance(table, Table)
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    context.table_ = table


@then('the document has two sections')
def then_the_document_has_two_sections(context):
    assert len(context.document.sections) == 2


@then('the first section is portrait')
def then_the_first_section_is_portrait(context):
    first_section = context.document.sections[0]
    expected_width, expected_height = context.original_dimensions
    assert first_section.orientation == WD_ORIENT.PORTRAIT
    assert first_section.page_width == expected_width
    assert first_section.page_height == expected_height


@then('the last paragraph contains only a page break')
def then_last_paragraph_contains_only_a_page_break(context):
    document = context.document
    paragraph = document.paragraphs[-1]
    assert len(paragraph.runs) == 1
    assert len(paragraph.runs[0]._r) == 1
    assert paragraph.runs[0]._r[0].type == 'page'


@then('the last paragraph contains the heading text')
def then_last_p_contains_heading_text(context):
    document = context.document
    text = context.heading_text
    paragraph = document.paragraphs[-1]
    assert paragraph.text == text


@then('the second section is landscape')
def then_the_second_section_is_landscape(context):
    new_section = context.document.sections[-1]
    expected_height, expected_width = context.original_dimensions
    assert new_section.orientation == WD_ORIENT.LANDSCAPE
    assert new_section.page_width == expected_width
    assert new_section.page_height == expected_height


@then('the style of the last paragraph is \'{style_name}\'')
def then_the_style_of_the_last_paragraph_is_style(context, style_name):
    document = context.document
    paragraph = document.paragraphs[-1]
    assert paragraph.style.name == style_name, (
        'got %s' % paragraph.style.name
    )
