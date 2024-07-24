# pyright: reportPrivateUsage=false

"""Step implementations for table-related features."""

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.table import (
    WD_ALIGN_VERTICAL,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
    WD_TABLE_DIRECTION,
)
from docx.shared import Inches
from docx.table import Table, _Cell, _Column, _Columns, _Row, _Rows

from helpers import test_docx

# given ===================================================


@given("a 2 x 2 table")
def given_a_2x2_table(context: Context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given("a 3x3 table having {span_state}")
def given_a_3x3_table_having_span_state(context: Context, span_state: str):
    table_idx = {
        "only uniform cells": 0,
        "a horizontal span": 1,
        "a vertical span": 2,
        "a combined span": 3,
    }[span_state]
    document = Document(test_docx("tbl-cell-access"))
    context.table_ = document.tables[table_idx]


@given("a _Cell object spanning {count} layout-grid cells")
def given_a_Cell_object_spanning_count_layout_grid_cells(context: Context, count: str):
    document = Document(test_docx("tbl-cell-props"))
    table = document.tables[0]
    context.cell = _Cell(table._tbl.tr_lst[int(count)].tc_lst[0], table)


@given("a _Cell object with {state} vertical alignment as cell")
def given_a_Cell_object_with_vertical_alignment_as_cell(context: Context, state: str):
    table_idx = {
        "inherited": 0,
        "bottom": 1,
        "center": 2,
        "top": 3,
    }[state]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    context.cell = table.cell(0, 0)


@given("a column collection having two columns")
def given_a_column_collection_having_two_columns(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    context.columns = document.tables[0].columns


@given("a row collection having two rows")
def given_a_row_collection_having_two_rows(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    context.rows = document.tables[0].rows


@given("a table")
def given_a_table(context: Context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given("a table cell")
def given_a_table_cell(context: Context):
    table = Document(test_docx("tbl-2x2-table")).tables[0]
    context.cell = table.cell(0, 0)


@given("a table cell having a width of {width}")
def given_a_table_cell_having_a_width_of_width(context: Context, width: str):
    table_idx = {"no explicit setting": 0, "1 inch": 1, "2 inches": 2}[width]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    cell = table.cell(0, 0)
    context.cell = cell


@given("a table column having a width of {width_desc}")
def given_a_table_having_a_width_of_width_desc(context: Context, width_desc: str):
    col_idx = {
        "no explicit setting": 0,
        "1440": 1,
    }[width_desc]
    docx_path = test_docx("tbl-col-props")
    document = Document(docx_path)
    context.column = document.tables[0].columns[col_idx]


@given("a table having {alignment} alignment")
def given_a_table_having_alignment_alignment(context: Context, alignment: str):
    table_idx = {
        "inherited": 3,
        "left": 4,
        "right": 5,
        "center": 6,
    }[alignment]
    docx_path = test_docx("tbl-props")
    document = Document(docx_path)
    context.table_ = document.tables[table_idx]


@given("a table having an autofit layout of {autofit}")
def given_a_table_having_an_autofit_layout_of_autofit(context: Context, autofit: str):
    tbl_idx = {
        "no explicit setting": 0,
        "autofit": 1,
        "fixed": 2,
    }[autofit]
    document = Document(test_docx("tbl-props"))
    context.table_ = document.tables[tbl_idx]


@given("a table having {style} style")
def given_a_table_having_style(context: Context, style: str):
    table_idx = {
        "no explicit": 0,
        "Table Grid": 1,
        "Light Shading - Accent 1": 2,
    }[style]
    document = Document(test_docx("tbl-having-applied-style"))
    context.document = document
    context.table_ = document.tables[table_idx]


@given("a table having table direction set {setting}")
def given_a_table_having_table_direction_setting(context: Context, setting: str):
    table_idx = ["to inherit", "right-to-left", "left-to-right"].index(setting)
    document = Document(test_docx("tbl-on-off-props"))
    context.table_ = document.tables[table_idx]


@given("a table having two columns")
def given_a_table_having_two_columns(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    # context.table is used internally by behave, underscore added
    # to distinguish this one
    context.table_ = document.tables[0]


@given("a table having two rows")
def given_a_table_having_two_rows(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    context.table_ = document.tables[0]


@given("a table row ending with {count} empty grid columns")
def given_a_table_row_ending_with_count_empty_grid_columns(context: Context, count: str):
    document = Document(test_docx("tbl-props"))
    table = document.tables[8]
    context.row = table.rows[int(count)]


@given("a table row having height of {state}")
def given_a_table_row_having_height_of_state(context: Context, state: str):
    table_idx = {"no explicit setting": 0, "2 inches": 2, "3 inches": 3}[state]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    context.row = table.rows[0]


@given("a table row having height rule {state}")
def given_a_table_row_having_height_rule_state(context: Context, state: str):
    table_idx = {"no explicit setting": 0, "automatic": 1, "at least": 2, "exactly": 3}[state]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    context.row = table.rows[0]


@given("a table row starting with {count} empty grid columns")
def given_a_table_row_starting_with_count_empty_grid_columns(context: Context, count: str):
    document = Document(test_docx("tbl-props"))
    table = document.tables[7]
    context.row = table.rows[int(count)]


# when =====================================================


@when("I add a 1.0 inch column to the table")
def when_I_add_a_1_inch_column_to_table(context: Context):
    context.column = context.table_.add_column(Inches(1.0))


@when("I add a 2 x 2 table into the first cell")
def when_I_add_a_2x2_table_into_the_first_cell(context: Context):
    context.table_ = context.cell.add_table(2, 2)


@when("I add a row to the table")
def when_add_row_to_table(context: Context):
    table = context.table_
    context.row = table.add_row()


@when("I assign a string to the cell text attribute")
def when_assign_string_to_cell_text_attribute(context: Context):
    cell = context.cell
    text = "foobar"
    cell.text = text
    context.expected_text = text


@when("I assign {value} to cell.vertical_alignment")
def when_I_assign_value_to_cell_vertical_alignment(context: Context, value: str):
    context.cell.vertical_alignment = eval(value)


@when("I assign {value} to row.height")
def when_I_assign_value_to_row_height(context: Context, value: str):
    new_value = None if value == "None" else int(value)
    context.row.height = new_value


@when("I assign {value} to row.height_rule")
def when_I_assign_value_to_row_height_rule(context: Context, value: str):
    new_value = None if value == "None" else getattr(WD_ROW_HEIGHT_RULE, value)
    context.row.height_rule = new_value


@when("I assign {value_str} to table.alignment")
def when_I_assign_value_to_table_alignment(context: Context, value_str: str):
    value = {
        "None": None,
        "WD_TABLE_ALIGNMENT.LEFT": WD_TABLE_ALIGNMENT.LEFT,
        "WD_TABLE_ALIGNMENT.RIGHT": WD_TABLE_ALIGNMENT.RIGHT,
        "WD_TABLE_ALIGNMENT.CENTER": WD_TABLE_ALIGNMENT.CENTER,
    }[value_str]
    table = context.table_
    table.alignment = value


@when("I assign {value} to table.style")
def when_apply_value_to_table_style(context: Context, value: str):
    table, styles = context.table_, context.document.styles
    if value == "None":
        new_value = None
    elif value.startswith("styles["):
        new_value = styles[value.split("'")[1]]
    else:
        new_value = styles[value]
    table.style = new_value


@when("I assign {value} to table.table_direction")
def when_assign_value_to_table_table_direction(context: Context, value: str):
    new_value = None if value == "None" else getattr(WD_TABLE_DIRECTION, value)
    context.table_.table_direction = new_value


@when("I merge from cell {origin} to cell {other}")
def when_I_merge_from_cell_origin_to_cell_other(context: Context, origin: str, other: str):
    def cell(table: Table, idx: int):
        row, col = idx // 3, idx % 3
        return table.cell(row, col)

    a_idx, b_idx = int(origin) - 1, int(other) - 1
    table = context.table_
    a, b = cell(table, a_idx), cell(table, b_idx)
    a.merge(b)


@when("I set the cell width to {width}")
def when_I_set_the_cell_width_to_width(context: Context, width: str):
    new_value = {"1 inch": Inches(1)}[width]
    context.cell.width = new_value


@when("I set the column width to {width_emu}")
def when_I_set_the_column_width_to_width_emu(context: Context, width_emu: str):
    new_value = None if width_emu == "None" else int(width_emu)
    context.column.width = new_value


@when("I set the table autofit to {setting}")
def when_I_set_the_table_autofit_to_setting(context: Context, setting: str):
    new_value = {"autofit": True, "fixed": False}[setting]
    table = context.table_
    table.autofit = new_value


# then =====================================================


@then("cell.grid_span is {count}")
def then_cell_grid_span_is_count(context: Context, count: str):
    expected = int(count)
    actual = context.cell.grid_span
    assert actual == expected, f"expected {expected}, got {actual}"


@then("cell.tables[0] is a 2 x 2 table")
def then_cell_tables_0_is_a_2x2_table(context: Context):
    cell = context.cell
    table = cell.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2


@then("cell.vertical_alignment is {value}")
def then_cell_vertical_alignment_is_value(context: Context, value: str):
    expected_value = {
        "None": None,
        "WD_ALIGN_VERTICAL.BOTTOM": WD_ALIGN_VERTICAL.BOTTOM,
        "WD_ALIGN_VERTICAL.CENTER": WD_ALIGN_VERTICAL.CENTER,
    }[value]
    actual_value = context.cell.vertical_alignment
    assert actual_value is expected_value, "cell.vertical_alignment is %s" % actual_value


@then("I can access a collection column by index")
def then_can_access_collection_column_by_index(context: Context):
    columns = context.columns
    for idx in range(2):
        column = columns[idx]
        assert isinstance(column, _Column)


@then("I can access a collection row by index")
def then_can_access_collection_row_by_index(context: Context):
    rows = context.rows
    for idx in range(2):
        row = rows[idx]
        assert isinstance(row, _Row)


@then("I can access the column collection of the table")
def then_can_access_column_collection_of_table(context: Context):
    table = context.table_
    columns = table.columns
    assert isinstance(columns, _Columns)


@then("I can access the row collection of the table")
def then_can_access_row_collection_of_table(context: Context):
    table = context.table_
    rows = table.rows
    assert isinstance(rows, _Rows)


@then("I can iterate over the column collection")
def then_can_iterate_over_column_collection(context: Context):
    columns = context.columns
    actual_count = 0
    for column in columns:
        actual_count += 1
        assert isinstance(column, _Column)
    assert actual_count == 2


@then("I can iterate over the row collection")
def then_can_iterate_over_row_collection(context: Context):
    rows = context.rows
    actual_count = 0
    for row in rows:
        actual_count += 1
        assert isinstance(row, _Row)
    assert actual_count == 2


@then("row.grid_cols_after is {value}")
def then_row_grid_cols_after_is_value(context: Context, value: str):
    expected = int(value)
    actual = context.row.grid_cols_after
    assert actual == expected, "expected %s, got %s" % (expected, actual)


@then("row.grid_cols_before is {value}")
def then_row_grid_cols_before_is_value(context: Context, value: str):
    expected = int(value)
    actual = context.row.grid_cols_before
    assert actual == expected, "expected %s, got %s" % (expected, actual)


@then("row.height is {value}")
def then_row_height_is_value(context: Context, value: str):
    expected_height = None if value == "None" else int(value)
    actual_height = context.row.height
    assert actual_height == expected_height, "expected %s, got %s" % (
        expected_height,
        actual_height,
    )


@then("row.height_rule is {value}")
def then_row_height_rule_is_value(context: Context, value: str):
    expected_rule = None if value == "None" else getattr(WD_ROW_HEIGHT_RULE, value)
    actual_rule = context.row.height_rule
    assert actual_rule == expected_rule, "expected %s, got %s" % (
        expected_rule,
        actual_rule,
    )


@then("table.alignment is {value_str}")
def then_table_alignment_is_value(context: Context, value_str: str):
    value = {
        "None": None,
        "WD_TABLE_ALIGNMENT.LEFT": WD_TABLE_ALIGNMENT.LEFT,
        "WD_TABLE_ALIGNMENT.RIGHT": WD_TABLE_ALIGNMENT.RIGHT,
        "WD_TABLE_ALIGNMENT.CENTER": WD_TABLE_ALIGNMENT.CENTER,
    }[value_str]
    table = context.table_
    assert table.alignment == value, "got %s" % table.alignment


@then("table.cell({row}, {col}).text is {expected_text}")
def then_table_cell_row_col_text_is_text(context: Context, row: str, col: str, expected_text: str):
    table = context.table_
    row_idx, col_idx = int(row), int(col)
    cell_text = table.cell(row_idx, col_idx).text
    assert cell_text == expected_text, "got %s" % cell_text


@then("table.style is styles['{style_name}']")
def then_table_style_is_styles_style_name(context: Context, style_name: str):
    table, styles = context.table_, context.document.styles
    expected_style = styles[style_name]
    assert table.style == expected_style, "got '%s'" % table.style


@then("table.table_direction is {value}")
def then_table_table_direction_is_value(context: Context, value: str):
    expected_value = None if value == "None" else getattr(WD_TABLE_DIRECTION, value)
    actual_value = context.table_.table_direction
    assert actual_value == expected_value, "got '%s'" % actual_value


@then("the cell contains the string I assigned")
def then_cell_contains_string_assigned(context: Context):
    cell, expected_text = context.cell, context.expected_text
    text = cell.paragraphs[0].runs[0].text
    msg = "expected '%s', got '%s'" % (expected_text, text)
    assert text == expected_text, msg


@then("the column cells text is {expected_text}")
def then_the_column_cells_text_is_expected_text(context: Context, expected_text: str):
    table = context.table_
    cells_text = " ".join(c.text for col in table.columns for c in col.cells)
    assert cells_text == expected_text, "got %s" % cells_text


@then("the length of the column collection is 2")
def then_len_of_column_collection_is_2(context: Context):
    columns = context.table_.columns
    assert len(columns) == 2


@then("the length of the row collection is 2")
def then_len_of_row_collection_is_2(context: Context):
    rows = context.table_.rows
    assert len(rows) == 2


@then("the new column has 2 cells")
def then_new_column_has_2_cells(context: Context):
    assert len(context.column.cells) == 2


@then("the new column is 1.0 inches wide")
def then_new_column_is_1_inches_wide(context: Context):
    assert context.column.width == Inches(1)


@then("the new row has 2 cells")
def then_new_row_has_2_cells(context: Context):
    assert len(context.row.cells) == 2


@then("the reported autofit setting is {autofit}")
def then_the_reported_autofit_setting_is_autofit(context: Context, autofit: str):
    expected_value = {"autofit": True, "fixed": False}[autofit]
    table = context.table_
    assert table.autofit is expected_value


@then("the reported column width is {width_emu}")
def then_the_reported_column_width_is_width_emu(context: Context, width_emu: str):
    expected_value = None if width_emu == "None" else int(width_emu)
    assert context.column.width == expected_value, "got %s" % context.column.width


@then("the reported width of the cell is {width}")
def then_the_reported_width_of_the_cell_is_width(context: Context, width: str):
    expected_width = {"None": None, "1 inch": Inches(1)}[width]
    actual_width = context.cell.width
    assert actual_width == expected_width, "expected %s, got %s" % (
        expected_width,
        actual_width,
    )


@then("the row cells text is {encoded_text}")
def then_the_row_cells_text_is_expected_text(context: Context, encoded_text: str):
    expected_text = encoded_text.replace("\\", "\n")
    table = context.table_
    cells_text = " ".join(c.text for row in table.rows for c in row.cells)
    assert cells_text == expected_text, "got %s" % cells_text


@then("the table has {count} columns")
def then_table_has_count_columns(context: Context, count: str):
    column_count = int(count)
    columns = context.table_.columns
    assert len(columns) == column_count


@then("the table has {count} rows")
def then_table_has_count_rows(context: Context, count: str):
    row_count = int(count)
    rows = context.table_.rows
    assert len(rows) == row_count


@then("the width of cell {n_str} is {inches_str} inches")
def then_the_width_of_cell_n_is_x_inches(context: Context, n_str: str, inches_str: str):
    def _cell(table: Table, idx: int):
        row, col = idx // 3, idx % 3
        return table.cell(row, col)

    idx, inches = int(n_str) - 1, float(inches_str)
    cell = _cell(context.table_, idx)
    assert cell.width is not None
    assert cell.width == Inches(inches), "got %s" % cell.width.inches


@then("the width of each cell is {inches} inches")
def then_the_width_of_each_cell_is_inches(context: Context, inches: str):
    table = context.table_
    expected_width = Inches(float(inches))
    for cell in table._cells:
        assert cell.width == expected_width, "got %s" % cell.width.inches


@then("the width of each column is {inches} inches")
def then_the_width_of_each_column_is_inches(context: Context, inches: str):
    table = context.table_
    expected_width = Inches(float(inches))
    for column in table.columns:
        assert column.width == expected_width, "got %s" % column.width.inches
