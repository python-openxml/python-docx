"""Step implementations for section-related features."""

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.section import Section
from docx.shared import Inches

from helpers import test_docx

# given ====================================================


@given("a Section object as section")
def given_a_Section_object_as_section(context: Context):
    context.section = Document(test_docx("sct-section-props")).sections[-1]


@given("a Section object of a multi-section document as section")
def given_a_Section_object_of_a_multi_section_document_as_section(context: Context):
    context.section = Document(test_docx("sct-inner-content")).sections[1]


@given("a Section object {with_or_without} a distinct first-page header as section")
def given_a_Section_object_with_or_without_first_page_header(
    context: Context, with_or_without: str
):
    section_idx = {"with": 1, "without": 0}[with_or_without]
    context.section = Document(test_docx("sct-first-page-hdrftr")).sections[section_idx]


@given("a section collection containing 3 sections")
def given_a_section_collection_containing_3_sections(context: Context):
    document = Document(test_docx("doc-access-sections"))
    context.sections = document.sections


@given("a section having known page dimension")
def given_a_section_having_known_page_dimension(context: Context):
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[-1]


@given("a section having known page margins")
def given_a_section_having_known_page_margins(context: Context):
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[0]


@given("a section having start type {start_type}")
def given_a_section_having_start_type(context: Context, start_type: str):
    section_idx = {
        "CONTINUOUS": 0,
        "NEW_PAGE": 1,
        "ODD_PAGE": 2,
        "EVEN_PAGE": 3,
        "NEW_COLUMN": 4,
    }[start_type]
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[section_idx]


@given("a section known to have {orientation} orientation")
def given_a_section_having_known_orientation(context: Context, orientation: str):
    section_idx = {"landscape": 0, "portrait": 1}[orientation]
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[section_idx]


# when =====================================================


@when("I assign {bool_val} to section.different_first_page_header_footer")
def when_I_assign_value_to_section_different_first_page_hdrftr(
    context: Context, bool_val: str
):
    context.section.different_first_page_header_footer = eval(bool_val)


@when("I set the {margin_side} margin to {inches} inches")
def when_I_set_the_margin_side_length(context: Context, margin_side: str, inches: str):
    prop_name = {
        "left": "left_margin",
        "right": "right_margin",
        "top": "top_margin",
        "bottom": "bottom_margin",
        "gutter": "gutter",
        "header": "header_distance",
        "footer": "footer_distance",
    }[margin_side]
    new_value = Inches(float(inches))
    setattr(context.section, prop_name, new_value)


@when("I set the section orientation to {orientation}")
def when_I_set_the_section_orientation(context: Context, orientation: str):
    new_orientation = {
        "WD_ORIENT.PORTRAIT": WD_ORIENT.PORTRAIT,
        "WD_ORIENT.LANDSCAPE": WD_ORIENT.LANDSCAPE,
        "None": None,
    }[orientation]
    context.section.orientation = new_orientation


@when("I set the section page height to {y} inches")
def when_I_set_the_section_page_height_to_y_inches(context: Context, y: str):
    context.section.page_height = Inches(float(y))


@when("I set the section page width to {x} inches")
def when_I_set_the_section_page_width_to_x_inches(context: Context, x: str):
    context.section.page_width = Inches(float(x))


@when("I set the section start type to {start_type}")
def when_I_set_the_section_start_type_to_start_type(context: Context, start_type: str):
    new_start_type = {
        "None": None,
        "CONTINUOUS": WD_SECTION.CONTINUOUS,
        "EVEN_PAGE": WD_SECTION.EVEN_PAGE,
        "NEW_COLUMN": WD_SECTION.NEW_COLUMN,
        "NEW_PAGE": WD_SECTION.NEW_PAGE,
        "ODD_PAGE": WD_SECTION.ODD_PAGE,
    }[start_type]
    context.section.start_type = new_start_type


# then =====================================================


@then("I can access a section by index")
def then_I_can_access_a_section_by_index(context: Context):
    sections = context.sections
    for idx in range(3):
        section = sections[idx]
        assert isinstance(section, Section)


@then("I can iterate over the sections")
def then_I_can_iterate_over_the_sections(context: Context):
    sections = context.sections
    actual_count = 0
    for section in sections:
        actual_count += 1
        assert isinstance(section, Section)
    assert actual_count == 3


@then("len(sections) is 3")
def then_len_sections_is_3(context: Context):
    sections = context.sections
    assert len(sections) == 3, "expected len(sections) of 3, got %s" % len(sections)


@then("section.different_first_page_header_footer is {bool_val}")
def then_section_different_first_page_header_footer_is(context: Context, bool_val: str):
    actual = context.section.different_first_page_header_footer
    expected = eval(bool_val)
    assert actual == expected, (
        "section.different_first_page_header_footer is %s" % actual
    )


@then("section.even_page_footer is a _Footer object")
def then_section_even_page_footer_is_a_Footer_object(context: Context):
    actual = type(context.section.even_page_footer).__name__
    expected = "_Footer"
    assert actual == expected, "section.even_page_footer is a %s object" % actual


@then("section.even_page_header is a _Header object")
def then_section_even_page_header_is_a_Header_object(context: Context):
    actual = type(context.section.even_page_header).__name__
    expected = "_Header"
    assert actual == expected, "section.even_page_header is a %s object" % actual


@then("section.first_page_footer is a _Footer object")
def then_section_first_page_footer_is_a_Footer_object(context: Context):
    actual = type(context.section.first_page_footer).__name__
    expected = "_Footer"
    assert actual == expected, "section.first_page_footer is a %s object" % actual


@then("section.first_page_header is a _Header object")
def then_section_first_page_header_is_a_Header_object(context: Context):
    actual = type(context.section.first_page_header).__name__
    expected = "_Header"
    assert actual == expected, "section.first_page_header is a %s object" % actual


@then("section.footer is a _Footer object")
def then_section_footer_is_a_Footer_object(context: Context):
    actual = type(context.section.footer).__name__
    expected = "_Footer"
    assert actual == expected, "section.footer is a %s object" % actual


@then("section.header is a _Header object")
def then_section_header_is_a_Header_object(context: Context):
    actual = type(context.section.header).__name__
    expected = "_Header"
    assert actual == expected, "section.header is a %s object" % actual


@then("section.iter_inner_content() produces the paragraphs and tables in section")
def step_impl(context: Context):
    actual = [type(item).__name__ for item in context.section.iter_inner_content()]
    expected = ["Table", "Paragraph", "Paragraph"]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("section.{propname}.is_linked_to_previous is True")
def then_section_hdrftr_prop_is_linked_to_previous_is_True(
    context: Context, propname: str
):
    actual = getattr(context.section, propname).is_linked_to_previous
    expected = True
    assert actual == expected, "section.%s.is_linked_to_previous is %s" % (
        propname,
        actual,
    )


@then("the reported {margin_side} margin is {inches} inches")
def then_the_reported_margin_is_inches(context: Context, margin_side: str, inches: str):
    prop_name = {
        "left": "left_margin",
        "right": "right_margin",
        "top": "top_margin",
        "bottom": "bottom_margin",
        "gutter": "gutter",
        "header": "header_distance",
        "footer": "footer_distance",
    }[margin_side]
    expected_value = Inches(float(inches))
    actual_value = getattr(context.section, prop_name)
    assert actual_value == expected_value


@then("the reported page orientation is {orientation}")
def then_the_reported_page_orientation_is_orientation(
    context: Context, orientation: str
):
    expected_value = {
        "WD_ORIENT.LANDSCAPE": WD_ORIENT.LANDSCAPE,
        "WD_ORIENT.PORTRAIT": WD_ORIENT.PORTRAIT,
    }[orientation]
    assert context.section.orientation == expected_value


@then("the reported page width is {x} inches")
def then_the_reported_page_width_is_width(context: Context, x: str):
    assert context.section.page_width == Inches(float(x))


@then("the reported page height is {y} inches")
def then_the_reported_page_height_is_11_inches(context: Context, y: str):
    assert context.section.page_height == Inches(float(y))


@then("the reported section start type is {start_type}")
def then_the_reported_section_start_type_is_type(context: Context, start_type: str):
    expected_start_type = {
        "CONTINUOUS": WD_SECTION.CONTINUOUS,
        "EVEN_PAGE": WD_SECTION.EVEN_PAGE,
        "NEW_COLUMN": WD_SECTION.NEW_COLUMN,
        "NEW_PAGE": WD_SECTION.NEW_PAGE,
        "ODD_PAGE": WD_SECTION.ODD_PAGE,
    }[start_type]
    assert context.section.start_type == expected_start_type
