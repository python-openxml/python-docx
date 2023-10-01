"""Step implementations for rendered page-break related features."""

from __future__ import annotations

from behave import given, then
from behave.runner import Context

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from helpers import test_docx

# given ===================================================


@given("a rendered_page_break in a hyperlink")
def given_a_rendered_page_break_in_a_hyperlink(context: Context):
    document = Document(test_docx("par-rendered-page-breaks"))
    paragraph = document.paragraphs[2]
    context.rendered_page_break = paragraph.rendered_page_breaks[0]


@given("a rendered_page_break in a paragraph")
def given_a_rendered_page_break_in_a_paragraph(context: Context):
    document = Document(test_docx("par-rendered-page-breaks"))
    paragraph = document.paragraphs[1]
    context.rendered_page_break = paragraph.rendered_page_breaks[0]


# then =====================================================


@then("rendered_page_break.preceding_paragraph_fragment includes the hyperlink")
def then_rendered_page_break_preceding_paragraph_fragment_includes_the_hyperlink(
    context: Context,
):
    para_frag = context.rendered_page_break.preceding_paragraph_fragment

    actual_value = type(para_frag).__name__
    expected_value = "Paragraph"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.text
    expected_value = "Page break in>><<this hyperlink"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.alignment
    expected_value = WD_PARAGRAPH_ALIGNMENT.RIGHT  # pyright: ignore
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.hyperlinks[0].runs[0].style.name
    expected_value = "Hyperlink"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.hyperlinks[0].address
    expected_value = "http://google.com/"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"


@then("rendered_page_break.preceding_paragraph_fragment is the content before break")
def then_rendered_page_break_preceding_paragraph_fragment_is_the_content_before_break(
    context: Context,
):
    para_frag = context.rendered_page_break.preceding_paragraph_fragment

    actual_value = type(para_frag).__name__
    expected_value = "Paragraph"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.text
    expected_value = "Page break here>>"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.alignment
    expected_value = WD_PARAGRAPH_ALIGNMENT.CENTER  # pyright: ignore
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.runs[0].style.name
    expected_value = "Default Paragraph Font"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"


@then("rendered_page_break.following_paragraph_fragment excludes the hyperlink")
def then_rendered_page_break_following_paragraph_fragment_excludes_the_hyperlink(
    context: Context,
):
    para_frag = context.rendered_page_break.following_paragraph_fragment

    # -- paragraph fragment is a Paragraph object --
    actual_value = type(para_frag).__name__
    expected_value = "Paragraph"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    # -- paragraph text is only the fragment after the page-break --
    actual_value = para_frag.text
    expected_value = " and another one here>><<with text following"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    # -- paragraph properties are preserved --
    actual_value = para_frag.alignment
    expected_value = WD_PARAGRAPH_ALIGNMENT.RIGHT  # pyright: ignore
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    # -- paragraph has no hyperlinks --
    actual_value = para_frag.hyperlinks
    expected_value = []
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    # -- following paragraph fragment retains any remaining page-breaks --
    actual_value = [type(rpb).__name__ for rpb in para_frag.rendered_page_breaks]
    expected_value = ["RenderedPageBreak"]
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"


@then("rendered_page_break.following_paragraph_fragment is the content after break")
def then_rendered_page_break_following_paragraph_fragment_is_the_content_after_break(
    context: Context,
):
    para_frag = context.rendered_page_break.following_paragraph_fragment

    actual_value = type(para_frag).__name__
    expected_value = "Paragraph"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.text
    expected_value = "<<followed by more text."
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.alignment
    expected_value = WD_PARAGRAPH_ALIGNMENT.CENTER  # pyright: ignore
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"

    actual_value = para_frag.runs[0].style.name
    expected_value = "Default Paragraph Font"
    assert (
        actual_value == expected_value
    ), f"expected: '{expected_value}', got: '{actual_value}'"
