# -*- coding: utf-8 -*-
#
# docx_steps.py
#
# Copyright (C) 2013 Steve Canny scanny@cisco.com
#
# This module is part of python-docx and is released under the MIT License:
# http://www.opensource.org/licenses/mit-license.php

"""Step implementations for python-docx acceptance tests"""

import os

from behave import given, then, when

from docx import Document


def absjoin(*paths):
    return os.path.abspath(os.path.join(*paths))

thisdir = os.path.split(__file__)[0]
scratch_dir = absjoin(thisdir, '../_scratch')
saved_docx_path = absjoin(scratch_dir, 'test_out.docx')

test_style = 'Heading1'
test_text = 'python-docx was here!'


# given ===================================================

@given('a new document created from the default template')
def step_given_new_doc_from_def_template(context):
    context.doc = Document()


# when ====================================================

@when('I add a new paragraph to the body')
def step_when_add_new_paragraph_to_body(context):
    body = context.doc.body
    context.p = body.add_paragraph()


@when('I add a new run to the paragraph')
def step_when_add_new_run_to_paragraph(context):
    context.r = context.p.add_run()


@when('I add new text to the run')
def step_when_add_new_text_to_run(context):
    context.r.add_text(test_text)


@when('I save the document')
def step_when_save_document(context):
    if os.path.isfile(saved_docx_path):
        os.remove(saved_docx_path)
    context.doc.save(saved_docx_path)


@when('I set the paragraph style')
def step_when_set_paragraph_style(context):
    context.p.add_run().add_text(test_text)
    context.p.style = test_style


# then =====================================================

@then('the document contains the text I added')
def step_then_document_contains_text_I_added(context):
    doc = Document(saved_docx_path)
    body = doc.body
    paragraphs = body.paragraphs
    p = paragraphs[-1]
    r = p.runs[0]
    assert r.text == test_text


@then('the paragraph has the style I set')
def step_then_paragraph_has_the_style_I_set(context):
    doc = Document(saved_docx_path)
    body = doc.body
    paragraphs = body.paragraphs
    p = paragraphs[-1]
    assert p.style == test_style
