import docx
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


def docx_addpicture(document, path, width_mm, caption):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(path, width=Mm(width_mm))
    # DOCX missing bookmarks
    # DOCX numbering goes next paragraph (WHY?)
    # See Issue: https://github.com/python-openxml/python-docx/issues/138
    if caption != "":
        paragraph = document.add_paragraph(style="Caption")
        p = paragraph._p  # this is the actual lxml element for a paragraph
        run = paragraph.add_run("Figure ")
        fld_xml = ('<w:fldSimple %s w:instr=" SEQ Figure \* ARABIC "/>' %
                   nsdecls('w'))
        fldSimple = parse_xml(fld_xml)  # this goes AFTER the field (!)
        #p.addnext(fldSimple)
        paragraph.add_run(" " + caption)


document = Document()

docx_addpicture(document,"testaddemf.emf", width_mm=100,caption="Groups Implementation")

document.save("testaddemf.docx")