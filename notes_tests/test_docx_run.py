import os
import logging

from types import GeneratorType
from itertools import chain

from nose.tools import *

from docx import api
from docx import text


logger = logging.getLogger('docx_converter.tests.docx_run')

DOC = api.Document(os.path.join(
    os.path.dirname(__file__),
    'data/run.docx'
))


def test_get_elements_type():
    result = DOC.paragraphs[0].runs[0].get_elements()
    assert_equals(type(result), GeneratorType)
    assert_equals([type(e) for e in result], [text.Text])
    
    
def test_element_classes():
    result = set()
    for p in DOC.paragraphs:
        for r in p.runs:
            for el in r.get_elements():
                result.add(type(el))
    expected = set([text.Text, text.LineBreak, text.EndnoteReference, text.FootnoteReference, text.Tab])
    assert_equals(result, expected)

