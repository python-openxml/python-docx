import os
import logging

from nose.tools import *

from docx import api
from docx.oxml.shared import qn
from docx.text import EndnoteReference, FootnoteReference
from docx.parts.notes import NotesPart, Note
from docx.oxml.parts.notes import CT_EndnoteReference, CT_FootnoteReference


logger = logging.getLogger('docx_converter.tests.docx')

DOC = api.Document(os.path.join(
    os.path.dirname(__file__),
    'data/notes.docx'
))


def test_parts():
    assert_equals(type(DOC.endnotes_part), NotesPart)
    assert_equals(type(DOC.footnotes_part), NotesPart)
    

def test_notes():
    part = DOC.endnotes_part
    assert_equals(type(part.notes), list)
    assert_equals(len(part.notes), 5)
    

def test_footnotes():
    part = DOC.footnotes_part
    assert_equals(type(part.notes), list)
    assert_equals(len(part.notes), 3)
    

def test_get_endnote():
    note = DOC.endnotes_part.get_note(0)
    assert_true(type(note), Note)
    
    
def test_get_footnote():
    note = DOC.footnotes_part.get_note(0)
    assert_true(type(note), Note)
    
    
def test_endnote():
    note = DOC.endnotes_part.get_note(3)
    assert_equals(note.id, 3)
    assert_is_none(note.type)
    note_paragraphs = note.paragraphs
    assert_equals(type(note_paragraphs), list)
    assert_equals(len(note_paragraphs), 2)
    

def test_footnote():
    note = DOC.footnotes_part.get_note(2)
    assert_equals(note.id, 2)
    assert_is_none(note.type)
    note_paragraphs = note.paragraphs
    assert_equals(type(note_paragraphs), list)
    assert_equals(len(note_paragraphs), 1)
    
    
def test_style_attributes():
    style = DOC.styles_part.get_style('style1')
    assert_equals(style.id, 'style1')
    assert_equals(style.type, 'paragraph')
    assert_equals(style.name, 'Heading 1')
    
    
def test_style_iterator():
    assert_equals(
        [(s.id, s.type, s.name) for s in DOC.styles_part.styles],
        [
            ('style0', 'paragraph', 'Normal'),
            ('style1', 'paragraph', 'Heading 1'),
            ('style15', 'character', 'Endnote anchor'),
            ('style16', 'character', 'Footnote anchor'),
            ('style17', 'character', 'Endnote Characters'),
            ('style18', 'character', 'Footnote Characters'),
            ('style19', 'paragraph', 'Heading'),
            ('style20', 'paragraph', 'Text body'),
            ('style21', 'paragraph', 'List'),
            ('style22', 'paragraph', 'Caption'),
            ('style23', 'paragraph', 'Index'),
            ('style24', 'paragraph', 'Endnote'),
            ('style25', 'paragraph', 'Footnote')
        ]
    )
    

def test_endnoterefs():
    run = DOC.paragraphs[2].runs[1]
    assert_equals(run.text, '')
    _endnoteref = run._r[1]
    assert_equals(_endnoteref.tag, qn('w:endnoteReference'))
    assert_equals(type(_endnoteref), CT_EndnoteReference)
    assert_equals(_endnoteref.id, 2)
    endnoterefs = run.endnote_references
    assert_true(endnoterefs)
    assert_equals(len(endnoterefs), 1)
    assert_equals(type(endnoterefs[0]), EndnoteReference)
    assert_equals(endnoterefs[0].id, _endnoteref.id)


def test_footnoterefs():
    run = DOC.paragraphs[4].runs[1]
    assert_equals(run.text, '')
    _footnoteref = run._r[1]
    assert_equals(_footnoteref.tag, qn('w:footnoteReference'))
    assert_equals(type(_footnoteref), CT_FootnoteReference)
    assert_equals(_footnoteref.id, 2)
    footnoterefs = run.footnote_references
    assert_true(footnoterefs)
    assert_equals(len(footnoterefs), 1)
    assert_equals(type(footnoterefs[0]), FootnoteReference)
    assert_equals(footnoterefs[0].id, _footnoteref.id)
        

