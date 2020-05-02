import docx

print(docx.__version__)

d= docx.Document()

p = d.add_paragraph('r1')

r = p.add_run('r2')

r.add_comment('hola comm')

d.save('d.docx')