# encoding: utf-8

"""
Test suite for the docx.text module
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml.text import CT_P, CT_R
from docx.parts.document import InlineShapes
from docx.shape import InlineShape
from docx.text import Paragraph, Run

import pytest

from .unitutil.cxml import element, xml
from .unitutil.mock import call, class_mock, instance_mock


class DescribeParagraph(object):

    def it_provides_access_to_the_runs_it_contains(self, runs_fixture):
        paragraph, Run_, r_, r_2_, run_, run_2_ = runs_fixture
        runs = paragraph.runs
        assert Run_.mock_calls == [
            call(r_, paragraph), call(r_2_, paragraph)
        ]
        assert runs == [run_, run_2_]

    def it_can_add_a_run_to_itself(self, add_run_fixture):
        paragraph, text, style, expected_xml = add_run_fixture
        run = paragraph.add_run(text, style)
        assert paragraph._p.xml == expected_xml
        assert isinstance(run, Run)
        assert run._r is paragraph._p.r_lst[0]

    def it_knows_its_alignment_value(self, alignment_get_fixture):
        paragraph, expected_value = alignment_get_fixture
        assert paragraph.alignment == expected_value

    def it_can_change_its_alignment_value(self, alignment_set_fixture):
        paragraph, value, expected_xml = alignment_set_fixture
        paragraph.alignment = value
        assert paragraph._p.xml == expected_xml

    def it_knows_its_paragraph_style(self, style_get_fixture):
        paragraph, expected_style = style_get_fixture
        assert paragraph.style == expected_style

    def it_can_change_its_paragraph_style(self, style_set_fixture):
        paragraph, value, expected_xml = style_set_fixture
        paragraph.style = value
        assert paragraph._p.xml == expected_xml

    def it_knows_the_text_it_contains(self, text_get_fixture):
        paragraph, expected_text = text_get_fixture
        assert paragraph.text == expected_text

    def it_can_replace_the_text_it_contains(self, text_set_fixture):
        paragraph, text, expected_text = text_set_fixture
        paragraph.text = text
        assert paragraph.text == expected_text

    def it_can_insert_a_paragraph_before_itself(self, insert_before_fixture):
        paragraph, text, style, body, expected_xml = insert_before_fixture
        new_paragraph = paragraph.insert_paragraph_before(text, style)
        assert isinstance(new_paragraph, Paragraph)
        assert new_paragraph.text == text
        assert new_paragraph.style == style
        assert body.xml == expected_xml

    def it_can_remove_its_content_while_preserving_formatting(
            self, clear_fixture):
        paragraph, expected_xml = clear_fixture
        _paragraph = paragraph.clear()
        assert paragraph._p.xml == expected_xml
        assert _paragraph is paragraph

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        ('w:p', None, None,
         'w:p/w:r'),
        ('w:p', 'foobar', None,
         'w:p/w:r/w:t"foobar"'),
        ('w:p', None, 'Strong',
         'w:p/w:r/w:rPr/w:rStyle{w:val=Strong}'),
        ('w:p', 'foobar', 'Strong',
         'w:p/w:r/(w:rPr/w:rStyle{w:val=Strong}, w:t"foobar")'),
    ])
    def add_run_fixture(self, request):
        before_cxml, text, style, after_cxml = request.param
        paragraph = Paragraph(element(before_cxml), None)
        expected_xml = xml(after_cxml)
        return paragraph, text, style, expected_xml

    @pytest.fixture(params=[
        ('w:p/w:pPr/w:jc{w:val=center}', WD_ALIGN_PARAGRAPH.CENTER),
        ('w:p', None),
    ])
    def alignment_get_fixture(self, request):
        cxml, expected_alignment_value = request.param
        paragraph = Paragraph(element(cxml), None)
        return paragraph, expected_alignment_value

    @pytest.fixture(params=[
        ('w:p', WD_ALIGN_PARAGRAPH.LEFT,
         'w:p/w:pPr/w:jc{w:val=left}'),
        ('w:p/w:pPr/w:jc{w:val=left}', WD_ALIGN_PARAGRAPH.CENTER,
         'w:p/w:pPr/w:jc{w:val=center}'),
        ('w:p/w:pPr/w:jc{w:val=left}', None,
         'w:p/w:pPr'),
        ('w:p', None, 'w:p/w:pPr'),
    ])
    def alignment_set_fixture(self, request):
        initial_cxml, new_alignment_value, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, new_alignment_value, expected_xml

    @pytest.fixture(params=[
        ('w:p', 'w:p'),
        ('w:p/w:pPr', 'w:p/w:pPr'),
        ('w:p/w:r/w:t"foobar"', 'w:p'),
        ('w:p/(w:pPr, w:r/w:t"foobar")', 'w:p/w:pPr'),
    ])
    def clear_fixture(self, request):
        initial_cxml, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, expected_xml

    @pytest.fixture(params=[
        ('w:body/w:p', 'foobar', 'Heading1',
         'w:body/(w:p/(w:pPr/w:pStyle{w:val=Heading1},w:r/w:t"foobar"),w:p)')
    ])
    def insert_before_fixture(self, request):
        body_cxml, text, style, expected_cxml = request.param
        body = element(body_cxml)
        paragraph = Paragraph(body.find(qn('w:p')), None)
        expected_xml = xml(expected_cxml)
        return paragraph, text, style, body, expected_xml

    @pytest.fixture
    def runs_fixture(self, p_, Run_, r_, r_2_, runs_):
        paragraph = Paragraph(p_, None)
        run_, run_2_ = runs_
        return paragraph, Run_, r_, r_2_, run_, run_2_

    @pytest.fixture(params=[
        ('w:p', 'Normal'),
        ('w:p/w:pPr', 'Normal'),
        ('w:p/w:pPr/w:pStyle{w:val=Heading1}', 'Heading1'),
    ])
    def style_get_fixture(self, request):
        p_cxml, expected_style = request.param
        paragraph = Paragraph(element(p_cxml), None)
        return paragraph, expected_style

    @pytest.fixture(params=[
        ('w:p',                                'Heading1',
         'w:p/w:pPr/w:pStyle{w:val=Heading1}'),
        ('w:p/w:pPr',                          'Heading1',
         'w:p/w:pPr/w:pStyle{w:val=Heading1}'),
        ('w:p/w:pPr/w:pStyle{w:val=Heading1}', 'Heading2',
         'w:p/w:pPr/w:pStyle{w:val=Heading2}'),
        ('w:p/w:pPr/w:pStyle{w:val=Heading1}', None,
         'w:p/w:pPr'),
        ('w:p',                                None,
         'w:p/w:pPr'),
    ])
    def style_set_fixture(self, request):
        p_cxml, new_style_value, expected_cxml = request.param
        paragraph = Paragraph(element(p_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, new_style_value, expected_xml

    @pytest.fixture(params=[
        ('w:p', ''),
        ('w:p/w:r', ''),
        ('w:p/w:r/w:t', ''),
        ('w:p/w:r/w:t"foo"', 'foo'),
        ('w:p/w:r/(w:t"foo", w:t"bar")', 'foobar'),
        ('w:p/w:r/(w:t"fo ", w:t"bar")', 'fo bar'),
        ('w:p/w:r/(w:t"foo", w:tab, w:t"bar")', 'foo\tbar'),
        ('w:p/w:r/(w:t"foo", w:br,  w:t"bar")', 'foo\nbar'),
        ('w:p/w:r/(w:t"foo", w:cr,  w:t"bar")', 'foo\nbar'),
    ])
    def text_get_fixture(self, request):
        p_cxml, expected_text_value = request.param
        paragraph = Paragraph(element(p_cxml), None)
        return paragraph, expected_text_value

    @pytest.fixture
    def text_set_fixture(self):
        paragraph = Paragraph(element('w:p'), None)
        paragraph.add_run('must not appear in result')
        new_text_value = 'foo\tbar\rbaz\n'
        expected_text_value = 'foo\tbar\nbaz\n'
        return paragraph, new_text_value, expected_text_value

    # fixture components ---------------------------------------------

    @pytest.fixture
    def p_(self, request, r_, r_2_):
        return instance_mock(request, CT_P, r_lst=(r_, r_2_))

    @pytest.fixture
    def Run_(self, request, runs_):
        run_, run_2_ = runs_
        return class_mock(
            request, 'docx.text.Run', side_effect=[run_, run_2_]
        )

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def r_2_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def runs_(self, request):
        run_ = instance_mock(request, Run, name='run_')
        run_2_ = instance_mock(request, Run, name='run_2_')
        return run_, run_2_


class DescribeRun(object):

    def it_knows_its_bool_prop_states(self, bool_prop_get_fixture):
        run, prop_name, expected_state = bool_prop_get_fixture
        assert getattr(run, prop_name) == expected_state

    def it_can_change_its_bool_prop_settings(self, bool_prop_set_fixture):
        run, prop_name, value, expected_xml = bool_prop_set_fixture
        setattr(run, prop_name, value)
        assert run._r.xml == expected_xml

    def it_knows_its_character_style(self, style_get_fixture):
        run, expected_style = style_get_fixture
        assert run.style == expected_style

    def it_can_change_its_character_style(self, style_set_fixture):
        run, style, expected_xml = style_set_fixture
        run.style = style
        assert run._r.xml == expected_xml

    def it_knows_its_underline_type(self, underline_get_fixture):
        run, expected_value = underline_get_fixture
        assert run.underline is expected_value

    def it_can_change_its_underline_type(self, underline_set_fixture):
        run, underline, expected_xml = underline_set_fixture
        run.underline = underline
        assert run._r.xml == expected_xml

    def it_raises_on_assign_invalid_underline_type(
            self, underline_raise_fixture):
        run, underline = underline_raise_fixture
        with pytest.raises(ValueError):
            run.underline = underline

    def it_can_add_text(self, add_text_fixture):
        run, text_str, expected_xml, Text_ = add_text_fixture
        _text = run.add_text(text_str)
        assert run._r.xml == expected_xml
        assert _text is Text_.return_value

    def it_can_add_a_break(self, add_break_fixture):
        run, break_type, expected_xml = add_break_fixture
        run.add_break(break_type)
        assert run._r.xml == expected_xml

    def it_can_add_a_tab(self, add_tab_fixture):
        run, expected_xml = add_tab_fixture
        run.add_tab()
        assert run._r.xml == expected_xml

    def it_can_add_a_picture(self, add_picture_fixture):
        (run, image_descriptor_, width, height, inline_shapes_,
         expected_width, expected_height, picture_) = add_picture_fixture
        picture = run.add_picture(image_descriptor_, width, height)
        inline_shapes_.add_picture.assert_called_once_with(
            image_descriptor_, run
        )
        assert picture is picture_
        assert picture.width == expected_width
        assert picture.height == expected_height

    def it_can_remove_its_content_but_keep_formatting(self, clear_fixture):
        run, expected_xml = clear_fixture
        _run = run.clear()
        assert run._r.xml == expected_xml
        assert _run is run

    def it_knows_the_text_it_contains(self, text_get_fixture):
        run, expected_text = text_get_fixture
        assert run.text == expected_text

    def it_can_replace_the_text_it_contains(self, text_set_fixture):
        run, text, expected_xml = text_set_fixture
        run.text = text
        assert run._r.xml == expected_xml

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        (WD_BREAK.LINE,   'w:r/w:br'),
        (WD_BREAK.PAGE,   'w:r/w:br{w:type=page}'),
        (WD_BREAK.COLUMN, 'w:r/w:br{w:type=column}'),
        (WD_BREAK.LINE_CLEAR_LEFT,
         'w:r/w:br{w:type=textWrapping, w:clear=left}'),
        (WD_BREAK.LINE_CLEAR_RIGHT,
         'w:r/w:br{w:type=textWrapping, w:clear=right}'),
        (WD_BREAK.LINE_CLEAR_ALL,
         'w:r/w:br{w:type=textWrapping, w:clear=all}'),
    ])
    def add_break_fixture(self, request):
        break_type, expected_cxml = request.param
        run = Run(element('w:r'), None)
        expected_xml = xml(expected_cxml)
        return run, break_type, expected_xml

    @pytest.fixture(params=[
        (None, None,  200,  100),
        (1000, 500,  1000,  500),
        (2000, None, 2000, 1000),
        (None, 2000, 4000, 2000),
    ])
    def add_picture_fixture(
            self, request, paragraph_, inline_shapes_, picture_):
        width, height, expected_width, expected_height = request.param
        paragraph_.part.inline_shapes = inline_shapes_
        run = Run(None, paragraph_)
        image_descriptor_ = 'image_descriptor_'
        picture_.width, picture_.height = 200, 100
        return (
            run, image_descriptor_, width, height, inline_shapes_,
            expected_width, expected_height, picture_
        )

    @pytest.fixture(params=[
        ('w:r/w:t"foo"', 'w:r/(w:t"foo", w:tab)'),
    ])
    def add_tab_fixture(self, request):
        r_cxml, expected_cxml = request.param
        run = Run(element(r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, expected_xml

    @pytest.fixture(params=[
        ('w:r',          'foo', 'w:r/w:t"foo"'),
        ('w:r/w:t"foo"', 'bar', 'w:r/(w:t"foo", w:t"bar")'),
        ('w:r',          'fo ', 'w:r/w:t{xml:space=preserve}"fo "'),
        ('w:r',          'f o', 'w:r/w:t"f o"'),
    ])
    def add_text_fixture(self, request, Text_):
        r_cxml, text, expected_cxml = request.param
        run = Run(element(r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, text, expected_xml, Text_

    @pytest.fixture(params=[
        ('w:r/w:rPr',                          'all_caps',       None),
        ('w:r/w:rPr/w:caps',                   'all_caps',       True),
        ('w:r/w:rPr/w:caps{w:val=on}',         'all_caps',       True),
        ('w:r/w:rPr/w:caps{w:val=off}',        'all_caps',       False),
        ('w:r/w:rPr/w:b{w:val=1}',             'bold',           True),
        ('w:r/w:rPr/w:i{w:val=0}',             'italic',         False),
        ('w:r/w:rPr/w:cs{w:val=true}',         'complex_script', True),
        ('w:r/w:rPr/w:bCs{w:val=false}',       'cs_bold',        False),
        ('w:r/w:rPr/w:iCs{w:val=on}',          'cs_italic',      True),
        ('w:r/w:rPr/w:dstrike{w:val=off}',     'double_strike',  False),
        ('w:r/w:rPr/w:emboss{w:val=1}',        'emboss',         True),
        ('w:r/w:rPr/w:vanish{w:val=0}',        'hidden',         False),
        ('w:r/w:rPr/w:i{w:val=true}',          'italic',         True),
        ('w:r/w:rPr/w:imprint{w:val=false}',   'imprint',        False),
        ('w:r/w:rPr/w:oMath{w:val=on}',        'math',           True),
        ('w:r/w:rPr/w:noProof{w:val=off}',     'no_proof',       False),
        ('w:r/w:rPr/w:outline{w:val=1}',       'outline',        True),
        ('w:r/w:rPr/w:rtl{w:val=0}',           'rtl',            False),
        ('w:r/w:rPr/w:shadow{w:val=true}',     'shadow',         True),
        ('w:r/w:rPr/w:smallCaps{w:val=false}', 'small_caps',     False),
        ('w:r/w:rPr/w:snapToGrid{w:val=on}',   'snap_to_grid',   True),
        ('w:r/w:rPr/w:specVanish{w:val=off}',  'spec_vanish',    False),
        ('w:r/w:rPr/w:strike{w:val=1}',        'strike',         True),
        ('w:r/w:rPr/w:webHidden{w:val=0}',     'web_hidden',     False),
    ])
    def bool_prop_get_fixture(self, request):
        r_cxml, bool_prop_name, expected_value = request.param
        run = Run(element(r_cxml), None)
        return run, bool_prop_name, expected_value

    @pytest.fixture(params=[
        # nothing to True, False, and None ---------------------------
        ('w:r',                             'all_caps',       True,
         'w:r/w:rPr/w:caps'),
        ('w:r',                             'bold',           False,
         'w:r/w:rPr/w:b{w:val=0}'),
        ('w:r',                             'italic',         None,
         'w:r/w:rPr'),
        # default to True, False, and None ---------------------------
        ('w:r/w:rPr/w:cs',                  'complex_script', True,
         'w:r/w:rPr/w:cs'),
        ('w:r/w:rPr/w:bCs',                 'cs_bold',        False,
         'w:r/w:rPr/w:bCs{w:val=0}'),
        ('w:r/w:rPr/w:iCs',                 'cs_italic',      None,
         'w:r/w:rPr'),
        # True to True, False, and None ------------------------------
        ('w:r/w:rPr/w:dstrike{w:val=1}',    'double_strike',  True,
         'w:r/w:rPr/w:dstrike'),
        ('w:r/w:rPr/w:emboss{w:val=on}',    'emboss',         False,
         'w:r/w:rPr/w:emboss{w:val=0}'),
        ('w:r/w:rPr/w:vanish{w:val=1}',     'hidden',         None,
         'w:r/w:rPr'),
        # False to True, False, and None -----------------------------
        ('w:r/w:rPr/w:i{w:val=false}',      'italic',         True,
         'w:r/w:rPr/w:i'),
        ('w:r/w:rPr/w:imprint{w:val=0}',    'imprint',        False,
         'w:r/w:rPr/w:imprint{w:val=0}'),
        ('w:r/w:rPr/w:oMath{w:val=off}',    'math',           None,
         'w:r/w:rPr'),
        # random mix -------------------------------------------------
        ('w:r/w:rPr/w:noProof{w:val=1}',    'no_proof',       False,
         'w:r/w:rPr/w:noProof{w:val=0}'),
        ('w:r/w:rPr',                       'outline',        True,
         'w:r/w:rPr/w:outline'),
        ('w:r/w:rPr/w:rtl{w:val=true}',     'rtl',            False,
         'w:r/w:rPr/w:rtl{w:val=0}'),
        ('w:r/w:rPr/w:shadow{w:val=on}',    'shadow',         True,
         'w:r/w:rPr/w:shadow'),
        ('w:r/w:rPr/w:smallCaps',           'small_caps',     False,
         'w:r/w:rPr/w:smallCaps{w:val=0}'),
        ('w:r/w:rPr/w:snapToGrid',          'snap_to_grid',   True,
         'w:r/w:rPr/w:snapToGrid'),
        ('w:r/w:rPr/w:specVanish',          'spec_vanish',    None,
         'w:r/w:rPr'),
        ('w:r/w:rPr/w:strike{w:val=foo}',   'strike',         True,
         'w:r/w:rPr/w:strike'),
        ('w:r/w:rPr/w:webHidden',           'web_hidden',     False,
         'w:r/w:rPr/w:webHidden{w:val=0}'),
    ])
    def bool_prop_set_fixture(self, request):
        initial_r_cxml, bool_prop_name, value, expected_cxml = request.param
        run = Run(element(initial_r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, bool_prop_name, value, expected_xml

    @pytest.fixture(params=[
        ('w:r',                   'w:r'),
        ('w:r/w:t"foo"',          'w:r'),
        ('w:r/w:br',              'w:r'),
        ('w:r/w:rPr',             'w:r/w:rPr'),
        ('w:r/(w:rPr, w:t"foo")', 'w:r/w:rPr'),
        ('w:r/(w:rPr/(w:b, w:i), w:t"foo", w:cr, w:t"bar")',
         'w:r/w:rPr/(w:b, w:i)'),
    ])
    def clear_fixture(self, request):
        initial_r_cxml, expected_cxml = request.param
        run = Run(element(initial_r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, expected_xml

    @pytest.fixture(params=[
        ('w:r',                              None),
        ('w:r/w:rPr/w:rStyle{w:val=Foobar}', 'Foobar'),
    ])
    def style_get_fixture(self, request):
        r_cxml, expected_style = request.param
        run = Run(element(r_cxml), None)
        return run, expected_style

    @pytest.fixture(params=[
        ('w:r',                            None,
         'w:r/w:rPr'),
        ('w:r',                            'Foo',
         'w:r/w:rPr/w:rStyle{w:val=Foo}'),
        ('w:r/w:rPr/w:rStyle{w:val=Foo}',  None,
         'w:r/w:rPr'),
        ('w:r/w:rPr/w:rStyle{w:val=Foo}',  'Bar',
         'w:r/w:rPr/w:rStyle{w:val=Bar}'),
    ])
    def style_set_fixture(self, request):
        initial_r_cxml, new_style, expected_cxml = request.param
        run = Run(element(initial_r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, new_style, expected_xml

    @pytest.fixture(params=[
        ('w:r', ''),
        ('w:r/w:t"foobar"', 'foobar'),
        ('w:r/(w:t"abc", w:tab, w:t"def", w:cr)', 'abc\tdef\n'),
        ('w:r/(w:br{w:type=page}, w:t"abc", w:t"def", w:tab)', '\nabcdef\t'),
    ])
    def text_get_fixture(self, request):
        r_cxml, expected_text = request.param
        run = Run(element(r_cxml), None)
        return run, expected_text

    @pytest.fixture(params=[
        ('abc  def', 'w:r/w:t"abc  def"'),
        ('abc\tdef', 'w:r/(w:t"abc", w:tab, w:t"def")'),
        ('abc\ndef', 'w:r/(w:t"abc", w:br,  w:t"def")'),
        ('abc\rdef', 'w:r/(w:t"abc", w:br,  w:t"def")'),
    ])
    def text_set_fixture(self, request):
        new_text, expected_cxml = request.param
        initial_r_cxml = 'w:r/w:t"should get deleted"'
        run = Run(element(initial_r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, new_text, expected_xml

    @pytest.fixture(params=[
        ('w:r',                         None),
        ('w:r/w:rPr/w:u',               None),
        ('w:r/w:rPr/w:u{w:val=single}', True),
        ('w:r/w:rPr/w:u{w:val=none}',   False),
        ('w:r/w:rPr/w:u{w:val=double}', WD_UNDERLINE.DOUBLE),
        ('w:r/w:rPr/w:u{w:val=wave}',   WD_UNDERLINE.WAVY),
    ])
    def underline_get_fixture(self, request):
        r_cxml, expected_underline = request.param
        run = Run(element(r_cxml), None)
        return run, expected_underline

    @pytest.fixture(params=[
        ('w:r', True,                'w:r/w:rPr/w:u{w:val=single}'),
        ('w:r', False,               'w:r/w:rPr/w:u{w:val=none}'),
        ('w:r', None,                'w:r/w:rPr'),
        ('w:r', WD_UNDERLINE.SINGLE, 'w:r/w:rPr/w:u{w:val=single}'),
        ('w:r', WD_UNDERLINE.THICK,  'w:r/w:rPr/w:u{w:val=thick}'),
        ('w:r/w:rPr/w:u{w:val=single}', True,
         'w:r/w:rPr/w:u{w:val=single}'),
        ('w:r/w:rPr/w:u{w:val=single}', False,
         'w:r/w:rPr/w:u{w:val=none}'),
        ('w:r/w:rPr/w:u{w:val=single}', None,
         'w:r/w:rPr'),
        ('w:r/w:rPr/w:u{w:val=single}', WD_UNDERLINE.SINGLE,
         'w:r/w:rPr/w:u{w:val=single}'),
        ('w:r/w:rPr/w:u{w:val=single}', WD_UNDERLINE.DOTTED,
         'w:r/w:rPr/w:u{w:val=dotted}'),
    ])
    def underline_set_fixture(self, request):
        initial_r_cxml, new_underline, expected_cxml = request.param
        run = Run(element(initial_r_cxml), None)
        expected_xml = xml(expected_cxml)
        return run, new_underline, expected_xml

    @pytest.fixture(params=['foobar', 42, 'single'])
    def underline_raise_fixture(self, request):
        invalid_underline_setting = request.param
        run = Run(element('w:r/w:rPr'), None)
        return run, invalid_underline_setting

    # fixture components ---------------------------------------------

    @pytest.fixture
    def inline_shapes_(self, request, picture_):
        inline_shapes_ = instance_mock(request, InlineShapes)
        inline_shapes_.add_picture.return_value = picture_
        return inline_shapes_

    @pytest.fixture
    def paragraph_(self, request):
        return instance_mock(request, Paragraph)

    @pytest.fixture
    def picture_(self, request):
        return instance_mock(request, InlineShape)

    @pytest.fixture
    def Text_(self, request):
        return class_mock(request, 'docx.text.Text')
