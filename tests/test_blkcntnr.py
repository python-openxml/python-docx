# pyright: reportPrivateUsage=false

"""Test suite for the docx.blkcntnr (block item container) module."""

from __future__ import annotations

from typing import cast

import pytest

import docx
from docx.blkcntnr import BlockItemContainer
from docx.document import Document
from docx.oxml.document import CT_Body
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element, xml
from .unitutil.file import snippet_seq, test_file
from .unitutil.mock import FixtureRequest, Mock, call, instance_mock, method_mock


class DescribeBlockItemContainer:
    """Unit-test suite for `docx.blkcntnr.BlockItemContainer`."""

    @pytest.mark.parametrize(
        ("text", "style"), [("", None), ("Foo", None), ("", "Bar"), ("Foo", "Bar")]
    )
    def it_can_add_a_paragraph(
        self,
        text: str,
        style: str | None,
        blkcntnr: BlockItemContainer,
        _add_paragraph_: Mock,
        paragraph_: Mock,
    ):
        paragraph_.style = None
        _add_paragraph_.return_value = paragraph_

        paragraph = blkcntnr.add_paragraph(text, style)

        _add_paragraph_.assert_called_once_with(blkcntnr)
        assert paragraph_.add_run.call_args_list == ([call(text)] if text else [])
        assert paragraph.style == style
        assert paragraph is paragraph_

    def it_can_add_a_table(self, blkcntnr: BlockItemContainer):
        rows, cols, width = 2, 2, Inches(2)

        table = blkcntnr.add_table(rows, cols, width)

        assert isinstance(table, Table)
        assert table._element.xml == snippet_seq("new-tbl")[0]
        assert table._parent is blkcntnr

    def it_can_iterate_its_inner_content(self):
        document = docx.Document(test_file("blk-inner-content.docx"))

        inner_content = document.iter_inner_content()

        para = next(inner_content)
        assert isinstance(para, Paragraph)
        assert para.text == "P1"
        # --
        t = next(inner_content)
        assert isinstance(t, Table)
        assert t.rows[0].cells[0].text == "T2"
        # --
        para = next(inner_content)
        assert isinstance(para, Paragraph)
        assert para.text == "P3"
        # --
        with pytest.raises(StopIteration):
            next(inner_content)

    @pytest.mark.parametrize(
        ("blkcntnr_cxml", "expected_count"),
        [
            ("w:body", 0),
            ("w:body/w:p", 1),
            ("w:body/(w:p,w:p)", 2),
            ("w:body/(w:p,w:tbl)", 1),
            ("w:body/(w:p,w:tbl,w:p)", 2),
        ],
    )
    def it_provides_access_to_the_paragraphs_it_contains(
        self, blkcntnr_cxml: str, expected_count: int, document_: Mock
    ):
        blkcntnr = BlockItemContainer(cast(CT_Body, element(blkcntnr_cxml)), document_)

        paragraphs = blkcntnr.paragraphs

        # -- supports len() --
        assert len(paragraphs) == expected_count
        # -- is iterable --
        assert all(isinstance(p, Paragraph) for p in paragraphs)
        # -- is indexable --
        assert all(p is paragraphs[idx] for idx, p in enumerate(paragraphs))

    @pytest.mark.parametrize(
        ("blkcntnr_cxml", "expected_count"),
        [
            ("w:body", 0),
            ("w:body/w:tbl", 1),
            ("w:body/(w:tbl,w:tbl)", 2),
            ("w:body/(w:p,w:tbl)", 1),
            ("w:body/(w:tbl,w:tbl,w:p)", 2),
        ],
    )
    def it_provides_access_to_the_tables_it_contains(
        self, blkcntnr_cxml: str, expected_count: int, document_: Mock
    ):
        blkcntnr = BlockItemContainer(cast(CT_Body, element(blkcntnr_cxml)), document_)

        tables = blkcntnr.tables

        # -- supports len() --
        assert len(tables) == expected_count
        # -- is iterable --
        assert all(isinstance(t, Table) for t in tables)
        # -- is indexable --
        assert all(t is tables[idx] for idx, t in enumerate(tables))

    def it_adds_a_paragraph_to_help(self, document_: Mock):
        blkcntnr = BlockItemContainer(cast(CT_Body, element("w:body")), document_)

        new_paragraph = blkcntnr._add_paragraph()

        assert isinstance(new_paragraph, Paragraph)
        assert new_paragraph._parent == blkcntnr
        assert blkcntnr._element.xml == xml("w:body/w:p")

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def _add_paragraph_(self, request: FixtureRequest):
        return method_mock(request, BlockItemContainer, "_add_paragraph")

    @pytest.fixture
    def blkcntnr(self, document_: Mock):
        blkcntnr_elm = cast(CT_Body, element("w:body"))
        return BlockItemContainer(blkcntnr_elm, document_)

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)

    @pytest.fixture
    def paragraph_(self, request: FixtureRequest):
        return instance_mock(request, Paragraph)
