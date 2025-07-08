"""Regression test for issues #1497 and #1494 – ZeroDivisionError when adding a
JPEG whose header reports 0 × 0 DPI.
"""

from pathlib import Path

from docx import Document

FIX = Path(__file__).with_name("test_files") / "zero_dpi.jpg"


class DescribeZeroDensityJPEG:
    """Suite covering `Document.add_picture()` with 0-DPI images."""

    def it_handles_zero_dpi(self):
        doc = Document()
        doc.add_picture(str(FIX))