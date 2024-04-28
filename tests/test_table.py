# pyright: reportPrivateUsage=false

"""Test suite for the docx.table module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import (
    WD_ALIGN_VERTICAL,
    WD_ROW_HEIGHT,
    WD_TABLE_ALIGNMENT,
    WD_TABLE_DIRECTION,
)
from docx.oxml.parser import parse_xml
from docx.oxml.table import CT_Row, CT_Tbl, CT_TblGridCol, CT_Tc
from docx.parts.document import DocumentPart
from docx.shared import Emu, Inches, Length
from docx.table import Table, _Cell, _Column, _Columns, _Row, _Rows
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element, xml
from .unitutil.file import snippet_seq
from .unitutil.mock import FixtureRequest, Mock, instance_mock, property_mock


class DescribeTable:
    """Unit-test suite for `docx.table._Rows` objects."""

    def it_can_add_a_row(self, document_: Mock):
        snippets = snippet_seq("add-row-col")
        tbl = cast(CT_Tbl, parse_xml(snippets[0]))
        table = Table(tbl, document_)

        row = table.add_row()

        assert table._tbl.xml == snippets[1]
        assert isinstance(row, _Row)
        assert row._tr is table._tbl.tr_lst[-1]
        assert row._parent is table

    def it_can_add_a_column(self, document_: Mock):
        snippets = snippet_seq("add-row-col")
        tbl = cast(CT_Tbl, parse_xml(snippets[0]))
        table = Table(tbl, document_)

        column = table.add_column(Inches(1.5))

        assert table._tbl.xml == snippets[2]
        assert isinstance(column, _Column)
        assert column._gridCol is table._tbl.tblGrid.gridCol_lst[-1]
        assert column._parent is table

    def it_provides_access_to_a_cell_by_row_and_col_indices(self, table: Table):
        for row_idx in range(2):
            for col_idx in range(2):
                cell = table.cell(row_idx, col_idx)
                assert isinstance(cell, _Cell)
                tr = table._tbl.tr_lst[row_idx]
                tc = tr.tc_lst[col_idx]
                assert tc is cell._tc

    def it_provides_access_to_the_table_rows(self, table: Table):
        rows = table.rows
        assert isinstance(rows, _Rows)

    def it_provides_access_to_the_table_columns(self, table: Table):
        columns = table.columns
        assert isinstance(columns, _Columns)

    def it_provides_access_to_the_cells_in_a_column(
        self, _cells_: Mock, _column_count_: Mock, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element("w:tbl")), document_)
        _cells_.return_value = [0, 1, 2, 3, 4, 5, 6, 7, 8]
        _column_count_.return_value = 3
        column_idx = 1

        column_cells = table.column_cells(column_idx)

        assert column_cells == [1, 4, 7]

    def it_provides_access_to_the_cells_in_a_row(
        self, _cells_: Mock, _column_count_: Mock, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element("w:tbl")), document_)
        _cells_.return_value = [0, 1, 2, 3, 4, 5, 6, 7, 8]
        _column_count_.return_value = 3

        row_cells = table.row_cells(1)

        assert row_cells == [3, 4, 5]

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:jc{w:val=center}", WD_TABLE_ALIGNMENT.CENTER),
            ("w:tbl/w:tblPr/w:jc{w:val=right}", WD_TABLE_ALIGNMENT.RIGHT),
            ("w:tbl/w:tblPr/w:jc{w:val=left}", WD_TABLE_ALIGNMENT.LEFT),
        ],
    )
    def it_knows_its_alignment_setting(
        self, tbl_cxml: str, expected_value: WD_TABLE_ALIGNMENT | None, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.alignment == expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", WD_TABLE_ALIGNMENT.LEFT, "w:tbl/w:tblPr/w:jc{w:val=left}"),
            (
                "w:tbl/w:tblPr/w:jc{w:val=left}",
                WD_TABLE_ALIGNMENT.RIGHT,
                "w:tbl/w:tblPr/w:jc{w:val=right}",
            ),
            ("w:tbl/w:tblPr/w:jc{w:val=right}", None, "w:tbl/w:tblPr"),
        ],
    )
    def it_can_change_its_alignment_setting(
        self,
        tbl_cxml: str,
        new_value: WD_TABLE_ALIGNMENT | None,
        expected_cxml: str,
        document_: Mock,
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.alignment = new_value
        assert table._tbl.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", True),
            ("w:tbl/w:tblPr/w:tblLayout", True),
            ("w:tbl/w:tblPr/w:tblLayout{w:type=autofit}", True),
            ("w:tbl/w:tblPr/w:tblLayout{w:type=fixed}", False),
        ],
    )
    def it_knows_whether_it_should_autofit(
        self, tbl_cxml: str, expected_value: bool, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.autofit is expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", True, "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}"),
            ("w:tbl/w:tblPr", False, "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}"),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
                True,
                "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}",
            ),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}",
                False,
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
            ),
        ],
    )
    def it_can_change_its_autofit_setting(
        self, tbl_cxml: str, new_value: bool, expected_cxml: str, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.autofit = new_value
        assert table._tbl.xml == xml(expected_cxml)

    def it_knows_it_is_the_table_its_children_belong_to(self, table: Table):
        assert table.table is table

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:bidiVisual", WD_TABLE_DIRECTION.RTL),
            ("w:tbl/w:tblPr/w:bidiVisual{w:val=0}", WD_TABLE_DIRECTION.LTR),
            ("w:tbl/w:tblPr/w:bidiVisual{w:val=on}", WD_TABLE_DIRECTION.RTL),
        ],
    )
    def it_knows_its_direction(
        self, tbl_cxml: str, expected_value: WD_TABLE_DIRECTION | None, document_: Mock
    ):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        assert Table(tbl, document_).table_direction == expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", WD_TABLE_DIRECTION.RTL, "w:tbl/w:tblPr/w:bidiVisual"),
            (
                "w:tbl/w:tblPr/w:bidiVisual",
                WD_TABLE_DIRECTION.LTR,
                "w:tbl/w:tblPr/w:bidiVisual{w:val=0}",
            ),
            (
                "w:tbl/w:tblPr/w:bidiVisual{w:val=0}",
                WD_TABLE_DIRECTION.RTL,
                "w:tbl/w:tblPr/w:bidiVisual",
            ),
            ("w:tbl/w:tblPr/w:bidiVisual{w:val=1}", None, "w:tbl/w:tblPr"),
        ],
    )
    def it_can_change_its_direction(
        self, tbl_cxml: str, new_value: WD_TABLE_DIRECTION, expected_cxml: str, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.table_direction = new_value
        assert table._element.xml == xml(expected_cxml)

    def it_knows_its_table_style(self, part_prop_: Mock, document_part_: Mock, document_: Mock):
        part_prop_.return_value = document_part_
        style_ = document_part_.get_style.return_value
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblStyle{w:val=BarBaz}")), document_)

        style = table.style

        document_part_.get_style.assert_called_once_with("BarBaz", WD_STYLE_TYPE.TABLE)
        assert style is style_

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "style_id", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", "Tbl A", "TblA", "w:tbl/w:tblPr/w:tblStyle{w:val=TblA}"),
            (
                "w:tbl/w:tblPr/w:tblStyle{w:val=TblA}",
                "Tbl B",
                "TblB",
                "w:tbl/w:tblPr/w:tblStyle{w:val=TblB}",
            ),
            ("w:tbl/w:tblPr/w:tblStyle{w:val=TblB}", None, None, "w:tbl/w:tblPr"),
        ],
    )
    def it_can_change_its_table_style(
        self,
        tbl_cxml: str,
        new_value: str | None,
        style_id: str | None,
        expected_cxml: str,
        document_: Mock,
        part_prop_: Mock,
        document_part_: Mock,
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        part_prop_.return_value = document_part_
        document_part_.get_style_id.return_value = style_id

        table.style = new_value

        document_part_.get_style_id.assert_called_once_with(new_value, WD_STYLE_TYPE.TABLE)
        assert table._tbl.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("snippet_idx", "cell_count", "unique_count", "matches"),
        [
            (0, 9, 9, ()),
            (1, 9, 8, ((0, 1),)),
            (2, 9, 8, ((1, 4),)),
            (3, 9, 6, ((0, 1, 3, 4),)),
            (4, 9, 4, ((0, 1), (3, 6), (4, 5, 7, 8))),
        ],
    )
    def it_provides_access_to_its_cells_to_help(
        self,
        snippet_idx: int,
        cell_count: int,
        unique_count: int,
        matches: tuple[tuple[int, ...]],
        document_: Mock,
    ):
        tbl_xml = snippet_seq("tbl-cells")[snippet_idx]
        table = Table(cast(CT_Tbl, parse_xml(tbl_xml)), document_)

        cells = table._cells

        assert len(cells) == cell_count
        assert len(set(cells)) == unique_count
        for matching_idxs in matches:
            comparator_idx = matching_idxs[0]
            for idx in matching_idxs[1:]:
                assert cells[idx] is cells[comparator_idx]

    def it_knows_its_column_count_to_help(self, document_: Mock):
        tbl_cxml = "w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)"
        expected_value = 3
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        column_count = table._column_count

        assert column_count == expected_value

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def _cells_(self, request: FixtureRequest):
        return property_mock(request, Table, "_cells")

    @pytest.fixture
    def _column_count_(self, request: FixtureRequest):
        return property_mock(request, Table, "_column_count")

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def part_prop_(self, request: FixtureRequest):
        return property_mock(request, Table, "part")

    @pytest.fixture
    def table(self, document_: Mock):
        tbl_cxml = "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol),w:tr/(w:tc,w:tc),w:tr/(w:tc,w:tc))"
        return Table(cast(CT_Tbl, element(tbl_cxml)), document_)


class Describe_Cell:
    """Unit-test suite for `docx.table._Cell` objects."""

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", 1),
            ("w:tc/w:tcPr", 1),
            ("w:tc/w:tcPr/w:gridSpan{w:val=1}", 1),
            ("w:tc/w:tcPr/w:gridSpan{w:val=4}", 4),
        ],
    )
    def it_knows_its_grid_span(self, tc_cxml: str, expected_value: int, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.grid_span == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_text"),
        [
            ("w:tc", ""),
            ('w:tc/w:p/w:r/w:t"foobar"', "foobar"),
            ('w:tc/(w:p/w:r/w:t"foo",w:p/w:r/w:t"bar")', "foo\nbar"),
            ('w:tc/(w:tcPr,w:p/w:r/w:t"foobar")', "foobar"),
            ('w:tc/w:p/w:r/(w:t"fo",w:tab,w:t"ob",w:br,w:t"ar",w:br)', "fo\tob\nar\n"),
        ],
    )
    def it_knows_what_text_it_contains(self, tc_cxml: str, expected_text: str, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        text = cell.text
        assert text == expected_text

    @pytest.mark.parametrize(
        ("tc_cxml", "new_text", "expected_cxml"),
        [
            ("w:tc/w:p", "foobar", 'w:tc/w:p/w:r/w:t"foobar"'),
            (
                "w:tc/w:p",
                "fo\tob\rar\n",
                'w:tc/w:p/w:r/(w:t"fo",w:tab,w:t"ob",w:br,w:t"ar",w:br)',
            ),
            (
                "w:tc/(w:tcPr, w:p, w:tbl, w:p)",
                "foobar",
                'w:tc/(w:tcPr, w:p/w:r/w:t"foobar")',
            ),
        ],
    )
    def it_can_replace_its_content_with_a_string_of_text(
        self, tc_cxml: str, new_text: str, expected_cxml: str, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.text = new_text
        assert cell._tc.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:vAlign{w:val=bottom}", WD_ALIGN_VERTICAL.BOTTOM),
            ("w:tc/w:tcPr/w:vAlign{w:val=top}", WD_ALIGN_VERTICAL.TOP),
        ],
    )
    def it_knows_its_vertical_alignment(
        self, tc_cxml: str, expected_value: WD_ALIGN_VERTICAL | None, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.vertical_alignment == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "new_value", "expected_cxml"),
        [
            ("w:tc", WD_ALIGN_VERTICAL.TOP, "w:tc/w:tcPr/w:vAlign{w:val=top}"),
            (
                "w:tc/w:tcPr",
                WD_ALIGN_VERTICAL.CENTER,
                "w:tc/w:tcPr/w:vAlign{w:val=center}",
            ),
            (
                "w:tc/w:tcPr/w:vAlign{w:val=center}",
                WD_ALIGN_VERTICAL.BOTTOM,
                "w:tc/w:tcPr/w:vAlign{w:val=bottom}",
            ),
            ("w:tc/w:tcPr/w:vAlign{w:val=center}", None, "w:tc/w:tcPr"),
            ("w:tc", None, "w:tc/w:tcPr"),
            ("w:tc/w:tcPr", None, "w:tc/w:tcPr"),
        ],
    )
    def it_can_change_its_vertical_alignment(
        self, tc_cxml: str, new_value: WD_ALIGN_VERTICAL | None, expected_cxml: str, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.vertical_alignment = new_value
        assert cell._element.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:tcW{w:w=25%,w:type=pct}", None),
            ("w:tc/w:tcPr/w:tcW{w:w=1440,w:type=dxa}", 914400),
        ],
    )
    def it_knows_its_width_in_EMU(self, tc_cxml: str, expected_value: int | None, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.width == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "new_value", "expected_cxml"),
        [
            ("w:tc", Inches(1), "w:tc/w:tcPr/w:tcW{w:w=1440,w:type=dxa}"),
            (
                "w:tc/w:tcPr/w:tcW{w:w=25%,w:type=pct}",
                Inches(2),
                "w:tc/w:tcPr/w:tcW{w:w=2880,w:type=dxa}",
            ),
        ],
    )
    def it_can_change_its_width(
        self, tc_cxml: str, new_value: Length, expected_cxml: str, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.width = new_value
        assert cell.width == new_value
        assert cell._tc.xml == xml(expected_cxml)

    def it_provides_access_to_the_paragraphs_it_contains(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/(w:p, w:p)")), parent_)

        paragraphs = cell.paragraphs

        # -- every w:p produces a Paragraph instance --
        assert len(paragraphs) == 2
        assert all(isinstance(p, Paragraph) for p in paragraphs)
        # -- the return value is iterable and indexable --
        assert all(p is paragraphs[idx] for idx, p in enumerate(paragraphs))

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_table_count"),
        [
            ("w:tc", 0),
            ("w:tc/w:tbl", 1),
            ("w:tc/(w:tbl,w:tbl)", 2),
            ("w:tc/(w:p,w:tbl)", 1),
            ("w:tc/(w:tbl,w:tbl,w:p)", 2),
        ],
    )
    def it_provides_access_to_the_tables_it_contains(
        self, tc_cxml: str, expected_table_count: int, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)

        tables = cell.tables

        # --- test len(), iterable, and indexed access
        assert len(tables) == expected_table_count
        assert all(isinstance(t, Table) for t in tables)
        assert all(t is tables[idx] for idx, t in enumerate(tables))

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_cxml"),
        [
            ("w:tc", "w:tc/w:p"),
            ("w:tc/w:p", "w:tc/(w:p, w:p)"),
            ("w:tc/w:tbl", "w:tc/(w:tbl, w:p)"),
        ],
    )
    def it_can_add_a_paragraph(self, tc_cxml: str, expected_cxml: str, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)

        p = cell.add_paragraph()

        assert isinstance(p, Paragraph)
        assert cell._tc.xml == xml(expected_cxml)

    def it_can_add_a_table(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/w:p")), parent_)

        table = cell.add_table(rows=2, cols=2)

        assert isinstance(table, Table)
        assert cell._element.xml == snippet_seq("new-tbl")[1]

    def it_can_merge_itself_with_other_cells(
        self, tc_: Mock, tc_2_: Mock, parent_: Mock, merged_tc_: Mock
    ):
        cell, other_cell = _Cell(tc_, parent_), _Cell(tc_2_, parent_)
        tc_.merge.return_value = merged_tc_

        merged_cell = cell.merge(other_cell)

        assert isinstance(merged_cell, _Cell)
        tc_.merge.assert_called_once_with(other_cell._tc)
        assert merged_cell._tc is merged_tc_
        assert merged_cell._parent is cell._parent

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def merged_tc_(self, request: FixtureRequest):
        return instance_mock(request, CT_Tc)

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def tc_(self, request: FixtureRequest):
        return instance_mock(request, CT_Tc)

    @pytest.fixture
    def tc_2_(self, request: FixtureRequest):
        return instance_mock(request, CT_Tc)


class Describe_Column:
    """Unit-test suite for `docx.table._Cell` objects."""

    def it_provides_access_to_its_cells(self, _index_prop_: Mock, table_prop_: Mock, table_: Mock):
        table_prop_.return_value = table_
        _index_prop_.return_value = 4
        column = _Column(cast(CT_TblGridCol, element("w:gridCol{w:w=500}")), table_)
        table_.column_cells.return_value = [3, 2, 1]

        cells = column.cells

        table_.column_cells.assert_called_once_with(4)
        assert cells == (3, 2, 1)

    def it_provides_access_to_the_table_it_belongs_to(self, table_: Mock):
        table_.table = table_
        column = _Column(cast(CT_TblGridCol, element("w:gridCol{w:w=500}")), table_)

        assert column.table is table_

    @pytest.mark.parametrize(
        ("gridCol_cxml", "expected_width"),
        [
            ("w:gridCol{w:w=4242}", 2693670),
            ("w:gridCol{w:w=1440}", 914400),
            ("w:gridCol{w:w=2.54cm}", 914400),
            ("w:gridCol{w:w=54mm}", 1944000),
            ("w:gridCol{w:w=12.5pt}", 158750),
            ("w:gridCol", None),
        ],
    )
    def it_knows_its_width_in_EMU(
        self, gridCol_cxml: str, expected_width: int | None, table_: Mock
    ):
        column = _Column(cast(CT_TblGridCol, element(gridCol_cxml)), table_)
        assert column.width == expected_width

    @pytest.mark.parametrize(
        ("gridCol_cxml", "new_value", "expected_cxml"),
        [
            ("w:gridCol", Emu(914400), "w:gridCol{w:w=1440}"),
            ("w:gridCol{w:w=4242}", Inches(0.5), "w:gridCol{w:w=720}"),
            ("w:gridCol{w:w=4242}", None, "w:gridCol"),
            ("w:gridCol", None, "w:gridCol"),
        ],
    )
    def it_can_change_its_width(
        self, gridCol_cxml: str, new_value: Length | None, expected_cxml: str, table_: Mock
    ):
        column = _Column(cast(CT_TblGridCol, element(gridCol_cxml)), table_)

        column.width = new_value

        assert column.width == new_value
        assert column._gridCol.xml == xml(expected_cxml)

    def it_knows_its_index_in_table_to_help(self, table_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)"))
        gridCol = tbl.tblGrid.gridCol_lst[1]
        column = _Column(gridCol, table_)
        assert column._index == 1

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def _index_prop_(self, request: FixtureRequest):
        return property_mock(request, _Column, "_index")

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_prop_(self, request: FixtureRequest):
        return property_mock(request, _Column, "table")


class Describe_Columns:
    """Unit-test suite for `docx.table._Columns` objects."""

    def it_has_sequence_behaviors(self, table_: Mock):
        columns = _Columns(cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol)")), table_)

        # -- it supports len() --
        assert len(columns) == 2
        # -- it is iterable --
        assert len(tuple(c for c in columns)) == 2
        assert all(type(c) is _Column for c in columns)
        # -- it is indexable --
        assert all(type(columns[i]) is _Column for i in range(2))

    def it_raises_on_indexed_access_out_of_range(self, table_: Mock):
        columns = _Columns(cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol)")), table_)

        with pytest.raises(IndexError):
            columns[2]
        with pytest.raises(IndexError):
            columns[-3]

    def it_provides_access_to_the_table_it_belongs_to(self, table_: Mock):
        columns = _Columns(cast(CT_Tbl, element("w:tbl")), table_)
        table_.table = table_

        assert columns.table is table_

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)


class Describe_Row:
    """Unit-test suite for `docx.table._Row` objects."""

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", 0),
            ("w:tr/w:trPr", 0),
            ("w:tr/w:trPr/w:gridAfter{w:val=0}", 0),
            ("w:tr/w:trPr/w:gridAfter{w:val=4}", 4),
        ],
    )
    def it_knows_its_grid_cols_after(self, tr_cxml: str, expected_value: int | None, parent_: Mock):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.grid_cols_after == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", 0),
            ("w:tr/w:trPr", 0),
            ("w:tr/w:trPr/w:gridBefore{w:val=0}", 0),
            ("w:tr/w:trPr/w:gridBefore{w:val=3}", 3),
        ],
    )
    def it_knows_its_grid_cols_before(
        self, tr_cxml: str, expected_value: int | None, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.grid_cols_before == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", None),
            ("w:tr/w:trPr", None),
            ("w:tr/w:trPr/w:trHeight", None),
            ("w:tr/w:trPr/w:trHeight{w:val=0}", 0),
            ("w:tr/w:trPr/w:trHeight{w:val=1440}", 914400),
        ],
    )
    def it_knows_its_height(self, tr_cxml: str, expected_value: int | None, parent_: Mock):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.height == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "new_value", "expected_cxml"),
        [
            ("w:tr", Inches(1), "w:tr/w:trPr/w:trHeight{w:val=1440}"),
            ("w:tr/w:trPr", Inches(1), "w:tr/w:trPr/w:trHeight{w:val=1440}"),
            ("w:tr/w:trPr/w:trHeight", Inches(1), "w:tr/w:trPr/w:trHeight{w:val=1440}"),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440}",
                Inches(2),
                "w:tr/w:trPr/w:trHeight{w:val=2880}",
            ),
            ("w:tr/w:trPr/w:trHeight{w:val=2880}", None, "w:tr/w:trPr/w:trHeight"),
            ("w:tr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr/w:trHeight", None, "w:tr/w:trPr/w:trHeight"),
        ],
    )
    def it_can_change_its_height(
        self, tr_cxml: str, new_value: Length | None, expected_cxml: str, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        row.height = new_value
        assert row._tr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", None),
            ("w:tr/w:trPr", None),
            ("w:tr/w:trPr/w:trHeight{w:val=0, w:hRule=auto}", WD_ROW_HEIGHT.AUTO),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=atLeast}",
                WD_ROW_HEIGHT.AT_LEAST,
            ),
            (
                "w:tr/w:trPr/w:trHeight{w:val=2880, w:hRule=exact}",
                WD_ROW_HEIGHT.EXACTLY,
            ),
        ],
    )
    def it_knows_its_height_rule(
        self, tr_cxml: str, expected_value: WD_ROW_HEIGHT | None, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.height_rule == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "new_value", "expected_cxml"),
        [
            ("w:tr", WD_ROW_HEIGHT.AUTO, "w:tr/w:trPr/w:trHeight{w:hRule=auto}"),
            (
                "w:tr/w:trPr",
                WD_ROW_HEIGHT.AT_LEAST,
                "w:tr/w:trPr/w:trHeight{w:hRule=atLeast}",
            ),
            (
                "w:tr/w:trPr/w:trHeight",
                WD_ROW_HEIGHT.EXACTLY,
                "w:tr/w:trPr/w:trHeight{w:hRule=exact}",
            ),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=exact}",
                WD_ROW_HEIGHT.AUTO,
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=auto}",
            ),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=auto}",
                None,
                "w:tr/w:trPr/w:trHeight{w:val=1440}",
            ),
            ("w:tr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr/w:trHeight", None, "w:tr/w:trPr/w:trHeight"),
        ],
    )
    def it_can_change_its_height_rule(
        self, tr_cxml: str, new_value: WD_ROW_HEIGHT | None, expected_cxml: str, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        row.height_rule = new_value
        assert row._tr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tbl_cxml", "row_idx", "expected_len"),
        [
            # -- cell corresponds to single layout-grid cell --
            ("w:tbl/w:tr/w:tc/w:p", 0, 1),
            # -- cell has a horizontal span --
            ("w:tbl/w:tr/w:tc/(w:tcPr/w:gridSpan{w:val=2},w:p)", 0, 2),
            # -- cell is in latter row of vertical span --
            (
                "w:tbl/(w:tr/w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))",
                1,
                1,
            ),
            # -- cell both has horizontal span and is latter row of vertical span --
            (
                "w:tbl/(w:tr/w:tc/(w:tcPr/(w:gridSpan{w:val=2},w:vMerge{w:val=restart}),w:p),"
                "w:tr/w:tc/(w:tcPr/(w:gridSpan{w:val=2},w:vMerge),w:p))",
                1,
                2,
            ),
        ],
    )
    def it_provides_access_to_its_cells(
        self, tbl_cxml: str, row_idx: int, expected_len: int, parent_: Mock
    ):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        tr = tbl.tr_lst[row_idx]
        table = Table(tbl, parent_)
        row = _Row(tr, table)

        cells = row.cells

        assert len(cells) == expected_len
        assert all(type(c) is _Cell for c in cells)

    def it_provides_access_to_the_table_it_belongs_to(self, parent_: Mock, table_: Mock):
        parent_.table = table_
        row = _Row(cast(CT_Row, element("w:tr")), parent_)
        assert row.table is table_

    def it_knows_its_index_in_table_to_help(self, parent_: Mock):
        tbl = element("w:tbl/(w:tr,w:tr,w:tr)")
        row = _Row(cast(CT_Row, tbl[1]), parent_)
        assert row._index == 1

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def _index_prop_(self, request: FixtureRequest):
        return property_mock(request, _Row, "_index")

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_prop_(self, request: FixtureRequest, table_: Mock):
        return property_mock(request, _Row, "table")


class Describe_Rows:
    """Unit-test suite for `docx.table._Rows` objects."""

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_len"),
        [
            ("w:tbl", 0),
            ("w:tbl/w:tr", 1),
            ("w:tbl/(w:tr,w:tr)", 2),
            ("w:tbl/(w:tr,w:tr,w:tr)", 3),
        ],
    )
    def it_has_sequence_behaviors(self, tbl_cxml: str, expected_len: int, parent_: Mock):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        table = Table(tbl, parent_)
        rows = _Rows(tbl, table)

        # -- it supports len() --
        assert len(rows) == expected_len
        # -- it is iterable --
        assert len(tuple(r for r in rows)) == expected_len
        assert all(type(r) is _Row for r in rows)
        # -- it is indexable --
        assert all(type(rows[i]) is _Row for i in range(expected_len))

    @pytest.mark.parametrize(
        ("tbl_cxml", "out_of_range_idx"),
        [
            ("w:tbl", 0),
            ("w:tbl", 1),
            ("w:tbl", -1),
            ("w:tbl/w:tr", 1),
            ("w:tbl/w:tr", -2),
            ("w:tbl/(w:tr,w:tr,w:tr)", 3),
            ("w:tbl/(w:tr,w:tr,w:tr)", -4),
        ],
    )
    def it_raises_on_indexed_access_out_of_range(
        self, tbl_cxml: str, out_of_range_idx: int, parent_: Mock
    ):
        rows = _Rows(cast(CT_Tbl, element(tbl_cxml)), parent_)

        with pytest.raises(IndexError, match="list index out of range"):
            rows[out_of_range_idx]

    @pytest.mark.parametrize(("start", "end", "expected_len"), [(1, 3, 2), (0, -1, 2)])
    def it_provides_sliced_access_to_rows(
        self, start: int, end: int, expected_len: int, parent_: Mock
    ):
        tbl = cast(CT_Tbl, element("w:tbl/(w:tr,w:tr,w:tr)"))
        rows = _Rows(tbl, parent_)

        slice_of_rows = rows[start:end]

        assert len(slice_of_rows) == expected_len
        for idx, row in enumerate(slice_of_rows):
            assert tbl.tr_lst.index(row._tr) == start + idx
            assert isinstance(row, _Row)

    def it_provides_access_to_the_table_it_belongs_to(self, parent_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl"))
        table = Table(tbl, parent_)
        rows = _Rows(tbl, table)

        assert rows.table is table

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Document)
