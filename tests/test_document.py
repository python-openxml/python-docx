# pyright: reportPrivateUsage=false
# pyright: reportUnknownMemberType=false

"""Unit test suite for the docx.document module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.comments import Comment, Comments
from docx.document import Document, _Body
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.opc.coreprops import CoreProperties
from docx.oxml.document import CT_Body, CT_Document
from docx.parts.document import DocumentPart
from docx.section import Section, Sections
from docx.settings import Settings
from docx.shape import InlineShape, InlineShapes
from docx.shared import Length
from docx.styles.styles import Styles
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .unitutil.cxml import element, xml
from .unitutil.mock import (
    FixtureRequest,
    Mock,
    class_mock,
    instance_mock,
    method_mock,
    property_mock,
)


class DescribeDocument:
    """Unit-test suite for `docx.document.Document`."""

    def it_can_add_a_comment(
        self,
        document_part_: Mock,
        comments_prop_: Mock,
        comments_: Mock,
        comment_: Mock,
        run_mark_comment_range_: Mock,
    ):
        comment_.comment_id = 42
        comments_.add_comment.return_value = comment_
        comments_prop_.return_value = comments_
        document = Document(cast(CT_Document, element("w:document/w:body/w:p/w:r")), document_part_)
        run = document.paragraphs[0].runs[0]

        comment = document.add_comment(run, "Comment text.")

        comments_.add_comment.assert_called_once_with("Comment text.", "", "")
        run_mark_comment_range_.assert_called_once_with(run, run, 42)
        assert comment is comment_

    @pytest.mark.parametrize(
        ("level", "style"), [(0, "Title"), (1, "Heading 1"), (2, "Heading 2"), (9, "Heading 9")]
    )
    def it_can_add_a_heading(
        self, level: int, style: str, document: Document, add_paragraph_: Mock, paragraph_: Mock
    ):
        add_paragraph_.return_value = paragraph_

        paragraph = document.add_heading("Spam vs. Bacon", level)

        add_paragraph_.assert_called_once_with(document, "Spam vs. Bacon", style)
        assert paragraph is paragraph_

    def it_raises_on_heading_level_out_of_range(self, document: Document):
        with pytest.raises(ValueError, match="level must be in range 0-9, got -1"):
            document.add_heading(level=-1)
        with pytest.raises(ValueError, match="level must be in range 0-9, got 10"):
            document.add_heading(level=10)

    def it_can_add_a_page_break(
        self, document: Document, add_paragraph_: Mock, paragraph_: Mock, run_: Mock
    ):
        add_paragraph_.return_value = paragraph_
        paragraph_.add_run.return_value = run_

        paragraph = document.add_page_break()

        add_paragraph_.assert_called_once_with(document)
        paragraph_.add_run.assert_called_once_with()
        run_.add_break.assert_called_once_with(WD_BREAK.PAGE)
        assert paragraph is paragraph_

    @pytest.mark.parametrize(
        ("text", "style"), [("", None), ("", "Heading 1"), ("foo\rbar", "Body Text")]
    )
    def it_can_add_a_paragraph(
        self,
        text: str,
        style: str | None,
        document: Document,
        body_: Mock,
        body_prop_: Mock,
        paragraph_: Mock,
    ):
        body_prop_.return_value = body_
        body_.add_paragraph.return_value = paragraph_

        paragraph = document.add_paragraph(text, style)

        body_.add_paragraph.assert_called_once_with(text, style)
        assert paragraph is paragraph_

    def it_can_add_a_picture(
        self, document: Document, add_paragraph_: Mock, run_: Mock, picture_: Mock
    ):
        path, width, height = "foobar.png", 100, 200
        add_paragraph_.return_value.add_run.return_value = run_
        run_.add_picture.return_value = picture_

        picture = document.add_picture(path, width, height)

        run_.add_picture.assert_called_once_with(path, width, height)
        assert picture is picture_

    @pytest.mark.parametrize(
        ("sentinel_cxml", "start_type", "new_sentinel_cxml"),
        [
            ("w:sectPr", WD_SECTION.EVEN_PAGE, "w:sectPr/w:type{w:val=evenPage}"),
            (
                "w:sectPr/w:type{w:val=evenPage}",
                WD_SECTION.ODD_PAGE,
                "w:sectPr/w:type{w:val=oddPage}",
            ),
            ("w:sectPr/w:type{w:val=oddPage}", WD_SECTION.NEW_PAGE, "w:sectPr"),
        ],
    )
    def it_can_add_a_section(
        self,
        sentinel_cxml: str,
        start_type: WD_SECTION,
        new_sentinel_cxml: str,
        Section_: Mock,
        section_: Mock,
        document_part_: Mock,
    ):
        Section_.return_value = section_
        document = Document(
            cast(CT_Document, element("w:document/w:body/(w:p,%s)" % sentinel_cxml)),
            document_part_,
        )

        section = document.add_section(start_type)

        assert document.element.xml == xml(
            "w:document/w:body/(w:p,w:p/w:pPr/%s,%s)" % (sentinel_cxml, new_sentinel_cxml)
        )
        sectPr = document.element.xpath("w:body/w:sectPr")[0]
        Section_.assert_called_once_with(sectPr, document_part_)
        assert section is section_

    def it_can_add_a_table(
        self,
        document: Document,
        _block_width_prop_: Mock,
        body_prop_: Mock,
        body_: Mock,
        table_: Mock,
    ):
        rows, cols, style = 4, 2, "Light Shading Accent 1"
        body_prop_.return_value = body_
        body_.add_table.return_value = table_
        _block_width_prop_.return_value = width = 42

        table = document.add_table(rows, cols, style)

        body_.add_table.assert_called_once_with(rows, cols, width)
        assert table == table_
        assert table.style == style

    def it_can_save_the_document_to_a_file(self, document_part_: Mock):
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        document.save("foobar.docx")

        document_part_.save.assert_called_once_with("foobar.docx")

    def it_provides_access_to_the_comments(self, document_part_: Mock, comments_: Mock):
        document_part_.comments = comments_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.comments is comments_

    def it_provides_access_to_its_core_properties(
        self, document_part_: Mock, core_properties_: Mock
    ):
        document_part_.core_properties = core_properties_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        core_properties = document.core_properties

        assert core_properties is core_properties_

    def it_provides_access_to_its_inline_shapes(self, document_part_: Mock, inline_shapes_: Mock):
        document_part_.inline_shapes = inline_shapes_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.inline_shapes is inline_shapes_

    def it_can_iterate_the_inner_content_of_the_document(
        self, body_prop_: Mock, body_: Mock, document_part_: Mock
    ):
        body_prop_.return_value = body_
        body_.iter_inner_content.return_value = iter((1, 2, 3))
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert list(document.iter_inner_content()) == [1, 2, 3]

    def it_provides_access_to_its_paragraphs(
        self, document: Document, body_prop_: Mock, body_: Mock, paragraphs_: Mock
    ):
        body_prop_.return_value = body_
        body_.paragraphs = paragraphs_
        paragraphs = document.paragraphs
        assert paragraphs is paragraphs_

    def it_provides_access_to_its_sections(
        self, document_part_: Mock, Sections_: Mock, sections_: Mock
    ):
        document_elm = cast(CT_Document, element("w:document"))
        Sections_.return_value = sections_
        document = Document(document_elm, document_part_)

        sections = document.sections

        Sections_.assert_called_once_with(document_elm, document_part_)
        assert sections is sections_

    def it_provides_access_to_its_settings(self, document_part_: Mock, settings_: Mock):
        document_part_.settings = settings_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.settings is settings_

    def it_provides_access_to_its_styles(self, document_part_: Mock, styles_: Mock):
        document_part_.styles = styles_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.styles is styles_

    def it_provides_access_to_its_tables(
        self, document: Document, body_prop_: Mock, body_: Mock, tables_: Mock
    ):
        body_prop_.return_value = body_
        body_.tables = tables_

        assert document.tables is tables_

    def it_provides_access_to_the_document_part(self, document_part_: Mock):
        document = Document(cast(CT_Document, element("w:document")), document_part_)
        assert document.part is document_part_

    def it_provides_access_to_the_document_body(
        self, _Body_: Mock, body_: Mock, document_part_: Mock
    ):
        _Body_.return_value = body_
        document_elm = cast(CT_Document, element("w:document/w:body"))
        body_elm = document_elm[0]
        document = Document(document_elm, document_part_)

        body = document._body

        _Body_.assert_called_once_with(body_elm, document)
        assert body is body_

    def it_determines_block_width_to_help(
        self, document: Document, sections_prop_: Mock, section_: Mock
    ):
        sections_prop_.return_value = [None, section_]
        section_.page_width = 6000
        section_.left_margin = 1500
        section_.right_margin = 1000

        width = document._block_width

        assert isinstance(width, Length)
        assert width == 3500

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def add_paragraph_(self, request: FixtureRequest):
        return method_mock(request, Document, "add_paragraph")

    @pytest.fixture
    def _Body_(self, request: FixtureRequest):
        return class_mock(request, "docx.document._Body")

    @pytest.fixture
    def body_(self, request: FixtureRequest):
        return instance_mock(request, _Body)

    @pytest.fixture
    def _block_width_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "_block_width")

    @pytest.fixture
    def body_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "_body")

    @pytest.fixture
    def comment_(self, request: FixtureRequest):
        return instance_mock(request, Comment)

    @pytest.fixture
    def comments_(self, request: FixtureRequest):
        return instance_mock(request, Comments)

    @pytest.fixture
    def comments_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "comments")

    @pytest.fixture
    def core_properties_(self, request: FixtureRequest):
        return instance_mock(request, CoreProperties)

    @pytest.fixture
    def document(self, document_part_: Mock) -> Document:
        document_elm = cast(CT_Document, element("w:document"))
        return Document(document_elm, document_part_)

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def inline_shapes_(self, request: FixtureRequest):
        return instance_mock(request, InlineShapes)

    @pytest.fixture
    def paragraph_(self, request: FixtureRequest):
        return instance_mock(request, Paragraph)

    @pytest.fixture
    def paragraphs_(self, request: FixtureRequest):
        return instance_mock(request, list)

    @pytest.fixture
    def picture_(self, request: FixtureRequest):
        return instance_mock(request, InlineShape)

    @pytest.fixture
    def run_(self, request: FixtureRequest):
        return instance_mock(request, Run)

    @pytest.fixture
    def run_mark_comment_range_(self, request: FixtureRequest):
        return method_mock(request, Run, "mark_comment_range")

    @pytest.fixture
    def Section_(self, request: FixtureRequest):
        return class_mock(request, "docx.document.Section")

    @pytest.fixture
    def section_(self, request: FixtureRequest):
        return instance_mock(request, Section)

    @pytest.fixture
    def Sections_(self, request: FixtureRequest):
        return class_mock(request, "docx.document.Sections")

    @pytest.fixture
    def sections_(self, request: FixtureRequest):
        return instance_mock(request, Sections)

    @pytest.fixture
    def sections_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "sections")

    @pytest.fixture
    def settings_(self, request: FixtureRequest):
        return instance_mock(request, Settings)

    @pytest.fixture
    def styles_(self, request: FixtureRequest):
        return instance_mock(request, Styles)

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def tables_(self, request: FixtureRequest):
        return instance_mock(request, list)


class Describe_Body:
    """Unit-test suite for `docx.document._Body`."""

    @pytest.mark.parametrize(
        ("cxml", "expected_cxml"),
        [
            ("w:body", "w:body"),
            ("w:body/w:p", "w:body"),
            ("w:body/w:sectPr", "w:body/w:sectPr"),
            ("w:body/(w:p, w:sectPr)", "w:body/w:sectPr"),
        ],
    )
    def it_can_clear_itself_of_all_content_it_holds(
        self, cxml: str, expected_cxml: str, document_: Mock
    ):
        body = _Body(cast(CT_Body, element(cxml)), document_)

        _body = body.clear_content()

        assert body._body.xml == xml(expected_cxml)
        assert _body is body

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)
