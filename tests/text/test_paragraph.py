"""Unit test suite for the docx.text.paragraph module."""

from typing import List, cast

import pytest

from docx import types as t
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.footnotes import Footnotes
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.parts.document import DocumentPart
from docx.text.paragraph import Paragraph
from docx.text.parfmt import ParagraphFormat
from docx.text.run import Run

from ..unitutil.cxml import element, xml
from ..unitutil.mock import call, class_mock, instance_mock, method_mock, property_mock


class DescribeParagraph:
    """Unit-test suite for `docx.text.run.Paragraph`."""

    @pytest.mark.parametrize(
        ("p_cxml", "expected_value"),
        [
            ("w:p/w:r", False),
            ('w:p/w:r/w:t"foobar"', False),
            ('w:p/w:hyperlink/w:r/(w:t"abc",w:lastRenderedPageBreak,w:t"def")', True),
            ("w:p/w:r/(w:lastRenderedPageBreak, w:lastRenderedPageBreak)", True),
        ],
    )
    def it_knows_whether_it_contains_a_page_break(
        self, p_cxml: str, expected_value: bool, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        assert paragraph.contains_page_break == expected_value

    def it_provides_access_to_its_footnotes(self, footnotes_fixture):
        paragraph, footnotes_ = footnotes_fixture
        assert paragraph.footnotes == footnotes_

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ("w:p/w:r", 0),
            ("w:p/w:hyperlink", 1),
            ("w:p/(w:r,w:hyperlink,w:r)", 1),
            ("w:p/(w:r,w:hyperlink,w:r,w:hyperlink)", 2),
            ("w:p/(w:hyperlink,w:r,w:hyperlink,w:r)", 2),
        ],
    )
    def it_provides_access_to_the_hyperlinks_it_contains(
        self, p_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        hyperlinks = paragraph.hyperlinks

        actual = [type(item).__name__ for item in hyperlinks]
        expected = ["Hyperlink" for _ in range(count)]
        assert actual == expected, f"expected: {expected}, got: {actual}"

    @pytest.mark.parametrize(
        ("p_cxml", "expected"),
        [
            ("w:p", []),
            ("w:p/w:r", ["Run"]),
            ("w:p/w:hyperlink", ["Hyperlink"]),
            ("w:p/(w:r,w:hyperlink,w:r)", ["Run", "Hyperlink", "Run"]),
            ("w:p/(w:hyperlink,w:r,w:hyperlink)", ["Hyperlink", "Run", "Hyperlink"]),
        ],
    )
    def it_can_iterate_its_inner_content_items(
        self, p_cxml: str, expected: List[str], fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        inner_content = paragraph.iter_inner_content()

        actual = [type(item).__name__ for item in inner_content]
        assert actual == expected, f"expected: {expected}, got: {actual}"

    def it_knows_its_paragraph_style(self, style_get_fixture):
        paragraph, style_id_, style_ = style_get_fixture
        style = paragraph.style
        paragraph.part.get_style.assert_called_once_with(
            style_id_, WD_STYLE_TYPE.PARAGRAPH
        )
        assert style is style_

    def it_can_change_its_paragraph_style(self, style_set_fixture):
        paragraph, value, expected_xml = style_set_fixture

        paragraph.style = value

        paragraph.part.get_style_id.assert_called_once_with(
            value, WD_STYLE_TYPE.PARAGRAPH
        )
        assert paragraph._p.xml == expected_xml

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ("w:p/w:r", 0),
            ("w:p/w:r/w:lastRenderedPageBreak", 1),
            ("w:p/w:hyperlink/w:r/w:lastRenderedPageBreak", 1),
            (
                "w:p/(w:r/w:lastRenderedPageBreak,"
                "w:hyperlink/w:r/w:lastRenderedPageBreak)",
                2,
            ),
            (
                "w:p/(w:hyperlink/w:r/w:lastRenderedPageBreak,w:r,"
                "w:r/w:lastRenderedPageBreak,w:r,w:hyperlink)",
                2,
            ),
        ],
    )
    def it_provides_access_to_the_rendered_page_breaks_it_contains(
        self, p_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        rendered_page_breaks = paragraph.rendered_page_breaks

        actual = [type(item).__name__ for item in rendered_page_breaks]
        expected = ["RenderedPageBreak" for _ in range(count)]
        assert actual == expected, f"expected: {expected}, got: {actual}"

    @pytest.mark.parametrize(
        ("p_cxml", "expected_value"),
        [
            ("w:p", ""),
            ("w:p/w:r", ""),
            ("w:p/w:r/w:t", ""),
            ('w:p/w:r/w:t"foo"', "foo"),
            ('w:p/w:r/(w:t"foo", w:t"bar")', "foobar"),
            ('w:p/w:r/(w:t"fo ", w:t"bar")', "fo bar"),
            ('w:p/w:r/(w:t"foo", w:tab, w:t"bar")', "foo\tbar"),
            ('w:p/w:r/(w:t"foo", w:br,  w:t"bar")', "foo\nbar"),
            ('w:p/w:r/(w:t"foo", w:cr,  w:t"bar")', "foo\nbar"),
            (
                'w:p/(w:r/w:t"click ",w:hyperlink{r:id=rId6}/w:r/w:t"here",'
                'w:r/w:t" for more")',
                "click here for more",
            ),
        ],
    )
    def it_knows_the_text_it_contains(self, p_cxml: str, expected_value: str):
        """Including the text of embedded hyperlinks."""
        paragraph = Paragraph(element(p_cxml), None)
        assert paragraph.text == expected_value

    def it_can_replace_the_text_it_contains(self, text_set_fixture):
        paragraph, text, expected_text = text_set_fixture
        paragraph.text = text
        assert paragraph.text == expected_text

    def it_knows_its_alignment_value(self, alignment_get_fixture):
        paragraph, expected_value = alignment_get_fixture
        assert paragraph.alignment == expected_value

    def it_can_change_its_alignment_value(self, alignment_set_fixture):
        paragraph, value, expected_xml = alignment_set_fixture
        paragraph.alignment = value
        assert paragraph._p.xml == expected_xml

    def it_provides_access_to_its_paragraph_format(self, parfmt_fixture):
        paragraph, ParagraphFormat_, paragraph_format_ = parfmt_fixture
        paragraph_format = paragraph.paragraph_format
        ParagraphFormat_.assert_called_once_with(paragraph._element)
        assert paragraph_format is paragraph_format_

    def it_provides_access_to_the_runs_it_contains(self, runs_fixture):
        paragraph, Run_, r_, r_2_, run_, run_2_ = runs_fixture
        runs = paragraph.runs
        assert Run_.mock_calls == [call(r_, paragraph), call(r_2_, paragraph)]
        assert runs == [run_, run_2_]

    def it_can_add_a_run_to_itself(self, add_run_fixture):
        paragraph, text, style, style_prop_, expected_xml = add_run_fixture
        run = paragraph.add_run(text, style)
        assert paragraph._p.xml == expected_xml
        assert isinstance(run, Run)
        assert run._r is paragraph._p.r_lst[0]
        if style:
            style_prop_.assert_called_once_with(style)

    def it_can_insert_a_paragraph_before_itself(self, insert_before_fixture):
        text, style, paragraph_, add_run_calls = insert_before_fixture
        paragraph = Paragraph(None, None)

        new_paragraph = paragraph.insert_paragraph_before(text, style)

        paragraph._insert_paragraph_before.assert_called_once_with(paragraph)
        assert new_paragraph.add_run.call_args_list == add_run_calls
        assert new_paragraph.style == style
        assert new_paragraph is paragraph_

    def it_can_remove_its_content_while_preserving_formatting(self, clear_fixture):
        paragraph, expected_xml = clear_fixture
        _paragraph = paragraph.clear()
        assert paragraph._p.xml == expected_xml
        assert _paragraph is paragraph

    def it_inserts_a_paragraph_before_to_help(self, _insert_before_fixture):
        paragraph, body, expected_xml = _insert_before_fixture
        new_paragraph = paragraph._insert_paragraph_before()
        assert isinstance(new_paragraph, Paragraph)
        assert body.xml == expected_xml

    # fixtures -------------------------------------------------------

    @pytest.fixture(
        params=[
            (
                "w:p/w:r/w:footnoteReference{w:id=2}",
                'w:footnotes/(w:footnote{w:id=1}/w:p/w:r/w:t"first note", w:footnote{w:id=2}/w:p/w:r/w:t"second note")',
                [2],
            ),
            (
                "w:p/w:r/w:footnoteReference{w:id=1}",
                'w:footnotes/(w:footnote{w:id=1}/w:p/w:r/w:t"first note", w:footnote{w:id=2}/w:p/w:r/w:t"second note")',
                [1],
            ),
            (
                "w:p/w:r/(w:footnoteReference{w:id=1}, w:footnoteReference{w:id=2})",
                'w:footnotes/(w:footnote{w:id=1}/w:p/w:r/w:t"first note", w:footnote{w:id=2}/w:p/w:r/w:t"second note")',
                [1, 2],
            ),
            (
                "w:p/w:r/(w:footnoteReference{w:id=3}, w:footnoteReference{w:id=2})",
                'w:footnotes/(w:footnote{w:id=1}/w:p/w:r/w:t"first note", w:footnote{w:id=2}/w:p/w:r/w:t"second note", w:footnote{w:id=3}/w:p/w:r/w:t"third note")',
                [3, 2],
            ),
        ]
    )
    def footnotes_fixture(self, request, document_part_):
        paragraph_cxml, footnotes_cxml, footnote_ids_in_p = request.param
        document_elm = element("w:document/w:body")
        document = Document(document_elm, document_part_)
        paragraph = Paragraph(element(paragraph_cxml), document._body)
        footnotes = Footnotes(element(footnotes_cxml), None)
        footnotes_in_p = [footnotes[id] for id in footnote_ids_in_p]
        document_part_.footnotes = footnotes
        return paragraph, footnotes_in_p

    @pytest.fixture(
        params=[
            ("w:p", None, None, "w:p/w:r"),
            ("w:p", "foobar", None, 'w:p/w:r/w:t"foobar"'),
            ("w:p", None, "Strong", "w:p/w:r"),
            ("w:p", "foobar", "Strong", 'w:p/w:r/w:t"foobar"'),
        ]
    )
    def add_run_fixture(self, request, run_style_prop_):
        before_cxml, text, style, after_cxml = request.param
        paragraph = Paragraph(element(before_cxml), None)
        expected_xml = xml(after_cxml)
        return paragraph, text, style, run_style_prop_, expected_xml

    @pytest.fixture(
        params=[
            ("w:p/w:pPr/w:jc{w:val=center}", WD_ALIGN_PARAGRAPH.CENTER),
            ("w:p", None),
        ]
    )
    def alignment_get_fixture(self, request):
        cxml, expected_alignment_value = request.param
        paragraph = Paragraph(element(cxml), None)
        return paragraph, expected_alignment_value

    @pytest.fixture(
        params=[
            ("w:p", WD_ALIGN_PARAGRAPH.LEFT, "w:p/w:pPr/w:jc{w:val=left}"),
            (
                "w:p/w:pPr/w:jc{w:val=left}",
                WD_ALIGN_PARAGRAPH.CENTER,
                "w:p/w:pPr/w:jc{w:val=center}",
            ),
            ("w:p/w:pPr/w:jc{w:val=left}", None, "w:p/w:pPr"),
            ("w:p", None, "w:p/w:pPr"),
        ]
    )
    def alignment_set_fixture(self, request):
        initial_cxml, new_alignment_value, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, new_alignment_value, expected_xml

    @pytest.fixture(
        params=[
            ("w:p", "w:p"),
            ("w:p/w:pPr", "w:p/w:pPr"),
            ('w:p/w:r/w:t"foobar"', "w:p"),
            ('w:p/(w:pPr, w:r/w:t"foobar")', "w:p/w:pPr"),
        ]
    )
    def clear_fixture(self, request):
        initial_cxml, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, expected_xml

    @pytest.fixture(
        params=[
            (None, None),
            ("Foo", None),
            (None, "Bar"),
            ("Foo", "Bar"),
        ]
    )
    def insert_before_fixture(self, request, _insert_paragraph_before_, add_run_):
        text, style = request.param
        paragraph_ = _insert_paragraph_before_.return_value
        add_run_calls = [] if text is None else [call(text)]
        paragraph_.style = None
        return text, style, paragraph_, add_run_calls

    @pytest.fixture(params=[("w:body/w:p{id=42}", "w:body/(w:p,w:p{id=42})")])
    def _insert_before_fixture(self, request):
        body_cxml, expected_cxml = request.param
        body = element(body_cxml)
        paragraph = Paragraph(body[0], None)
        expected_xml = xml(expected_cxml)
        return paragraph, body, expected_xml

    @pytest.fixture
    def parfmt_fixture(self, ParagraphFormat_, paragraph_format_):
        paragraph = Paragraph(element("w:p"), None)
        return paragraph, ParagraphFormat_, paragraph_format_

    @pytest.fixture
    def runs_fixture(self, p_, Run_, r_, r_2_, runs_):
        paragraph = Paragraph(p_, None)
        run_, run_2_ = runs_
        return paragraph, Run_, r_, r_2_, run_, run_2_

    @pytest.fixture
    def style_get_fixture(self, part_prop_):
        style_id = "Foobar"
        p_cxml = "w:p/w:pPr/w:pStyle{w:val=%s}" % style_id
        paragraph = Paragraph(element(p_cxml), None)
        style_ = part_prop_.return_value.get_style.return_value
        return paragraph, style_id, style_

    @pytest.fixture(
        params=[
            ("w:p", "Heading 1", "Heading1", "w:p/w:pPr/w:pStyle{w:val=Heading1}"),
            (
                "w:p/w:pPr",
                "Heading 1",
                "Heading1",
                "w:p/w:pPr/w:pStyle{w:val=Heading1}",
            ),
            (
                "w:p/w:pPr/w:pStyle{w:val=Heading1}",
                "Heading 2",
                "Heading2",
                "w:p/w:pPr/w:pStyle{w:val=Heading2}",
            ),
            ("w:p/w:pPr/w:pStyle{w:val=Heading1}", "Normal", None, "w:p/w:pPr"),
            ("w:p", None, None, "w:p/w:pPr"),
        ]
    )
    def style_set_fixture(self, request, part_prop_):
        p_cxml, value, style_id, expected_cxml = request.param
        paragraph = Paragraph(element(p_cxml), None)
        part_prop_.return_value.get_style_id.return_value = style_id
        expected_xml = xml(expected_cxml)
        return paragraph, value, expected_xml

    @pytest.fixture
    def text_set_fixture(self):
        paragraph = Paragraph(element("w:p"), None)
        paragraph.add_run("must not appear in result")
        new_text_value = "foo\tbar\rbaz\n"
        expected_text_value = "foo\tbar\nbaz\n"
        return paragraph, new_text_value, expected_text_value

    # fixture components ---------------------------------------------

    @pytest.fixture
    def add_run_(self, request):
        return method_mock(request, Paragraph, "add_run")

    @pytest.fixture
    def document_part_(self, request):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def _insert_paragraph_before_(self, request):
        return method_mock(request, Paragraph, "_insert_paragraph_before")

    @pytest.fixture
    def p_(self, request, r_, r_2_):
        return instance_mock(request, CT_P, r_lst=(r_, r_2_))

    @pytest.fixture
    def ParagraphFormat_(self, request, paragraph_format_):
        return class_mock(
            request,
            "docx.text.paragraph.ParagraphFormat",
            return_value=paragraph_format_,
        )

    @pytest.fixture
    def paragraph_format_(self, request):
        return instance_mock(request, ParagraphFormat)

    @pytest.fixture
    def part_prop_(self, request, document_part_):
        return property_mock(request, Paragraph, "part", return_value=document_part_)

    @pytest.fixture
    def Run_(self, request, runs_):
        run_, run_2_ = runs_
        return class_mock(
            request, "docx.text.paragraph.Run", side_effect=[run_, run_2_]
        )

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def r_2_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def run_style_prop_(self, request):
        return property_mock(request, Run, "style")

    @pytest.fixture
    def runs_(self, request):
        run_ = instance_mock(request, Run, name="run_")
        run_2_ = instance_mock(request, Run, name="run_2_")
        return run_, run_2_
