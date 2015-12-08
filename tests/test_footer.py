from .unitutil.file import absjoin, test_file_dir
from docx.api import Document
from docx.oxml.footer import CT_Ftr
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.part import XmlPart


dir_pkg_path = absjoin(test_file_dir, 'expanded_docx')


class DescribeFooterLoad(object):
    def it_has_part_as_footer_part(self):
        document = Document(dir_pkg_path)
        footer_part_exists = False
        for rel_id, part in document.part.related_parts.items():
            if part.content_type == CT.WML_FOOTER:
                footer_part_exists = True
                assert isinstance(part, XmlPart)

        assert footer_part_exists

    def it_has_rel_as_footer_rel(self):
        document = Document(dir_pkg_path)
        footer_rel_exists = False
        for rel_id, rel in document.part.rels.items():
            if rel.reltype == RT.FOOTER:
                footer_rel_exists = True

        assert footer_rel_exists


class DescribeRemoveFooter(object):
    def it_removes_footer_part(self):
        document = Document(dir_pkg_path)
        document.remove_footers()

        for rel_id, part in document.part.related_parts.items():
            assert part.content_type != CT.WML_FOOTER

        footer_elm_tag = 'w:footerReference'
        sentinel_sectPr = document._body._body.get_or_add_sectPr()
        footer_elms = sentinel_sectPr.findall(qn(footer_elm_tag))
        assert len(footer_elms) == 0


class DescribeAddFooter(object):
    def it_adds_to_doc_without_footer(self):
        document = Document(dir_pkg_path)
        document.remove_footers()

        footer = document.add_footer()
        footer_elm_tag = 'w:footerReference'
        sentinel_sectPr = document._body._body.get_or_add_sectPr()
        footer_elms = sentinel_sectPr.findall(qn(footer_elm_tag))
        assert len(footer_elms) == 1

        assert footer
        assert len(footer.paragraphs) == 0

        footer.add_paragraph('foobar')
        assert len(footer.paragraphs) == 1
        # import uuid
        # random_name = uuid.uuid4().hex
        # finish_path = '{}.docx'.format(random_name)
        # document.save(finish_path)
        # print 'file {} footer added!'.format(finish_path)


class DescribeCTHdr(object):
    def it_creates_an_element_of_type_w_hdr(self):
        footer = CT_Ftr.new()
        assert footer.tag.endswith('ftr')
