from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

# ایجاد یک سند جدید
doc = Document()

# افزودن یک جدول با 1 ردیف و 2 ستون برای سربرگ
table = doc.add_table(rows=1, cols=2)

# تنظیم عرض ستون‌ها
table.autofit = False
table.columns[0].width = Inches(2)  # عرض ستون لوگو
table.columns[1].width = Inches(4)  # عرض ستون عنوان

# افزودن لوگوی شرکت (در اینجا فقط یک متن برای نمایش قرار می‌دهیم)
cell_logo = table.cell(0, 0)
cell_logo.text = 'لوگوی شرکت'

# افزودن عنوان شرکت
cell_info = table.cell(0, 1)
cell_info.text = 'فروهر سالار ایرانیان کهن بران'

# تنظیم فونت عنوان
run = cell_info.paragraphs[0].runs[0]
run.font.name = 'B Nazanin'  # استفاده از فونت فارسی
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')  # تنظیم فونت برای متن فارسی
run.font.size = Pt(18)
run.bold = True

# افزودن خط زیر عنوان
paragraph = cell_info.add_paragraph()
paragraph.add_run().add_break()  # اضافه کردن یک خط خالی
paragraph.add_run('_________________________________________')

# افزودن اطلاعات شرکت
info_paragraph = doc.add_paragraph()
info_paragraph.add_run('کد اقتصادی: 411656516457\n')
info_paragraph.add_run('شناسه ملی: 14008991180\n')
info_paragraph.add_run('آدرس: تهران، خیابان آزادی، کوچه بهار، پلاک ۱۲۳\n')
info_paragraph.add_run('تلفن: ۰۲۱-۱۲۳۴۵۶۷۸\n')
info_paragraph.add_run('ایمیل: info@furouhar.com\n')

# تنظیم فونت اطلاعات شرکت
for run in info_paragraph.runs:
    run.font.name = 'B Nazanin'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    run.font.size = Pt(12)

# افزودن یک جدول برای اطلاعات بیشتر
more_info_table = doc.add_table(rows=4, cols=2)
more_info_table.autofit = False
more_info_table.columns[0].width = Inches(2)
more_info_table.columns[1].width = Inches(4)

# پر کردن جدول با اطلاعات
more_info_table.cell(0, 0).text = 'مدیر عامل:'
more_info_table.cell(0, 1).text = 'آقای علی رضایی'
more_info_table.cell(1, 0).text = 'تاریخ تأسیس:'
more_info_table.cell(1, 1).text = '۱۳۸۵/۰۵/۱۵'
more_info_table.cell(2, 0).text = 'شماره حساب:'
more_info_table.cell(2, 1).text = '۱۲۳۴۵۶۷۸۹۰'
more_info_table.cell(3, 0).text = 'شماره شبا:'
more_info_table.cell(3, 1).text = 'IR123456789012345678901234'

# تنظیم فونت جدول اطلاعات بیشتر
for row in more_info_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(12)

# ذخیره فایل ورد
file_path = 'sarbargh_furouh_design_v2.docx'
doc.save(file_path)

print(f"فایل ورد با موفقیت ایجاد شد و در مسیر '{file_path}' ذخیره شد.")
